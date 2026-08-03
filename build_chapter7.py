from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK

SRC = "/Users/mariomartins2/Downloads/7. Ambiente do estudante - submissão, acompanhamento e revisão.docx"
OUT = "/Users/mariomartins2/Library/CloudStorage/GoogleDrive-mario.martins@ufersa.edu.br/Meu Drive/Disciplinas_em_R/ANETE/7. Ambiente do estudante no Textopia_ - versão corrigida.docx"

doc = Document(SRC)
body = doc._element.body
sect = body.sectPr
for child in list(body):
    if child is not sect:
        body.remove(child)

def p(text="", style=None):
    return doc.add_paragraph(text, style=style)

def h(text, level=2):
    return doc.add_heading(text, level=level)

def bullets(items):
    for item in items:
        par = p()
        par.paragraph_format.left_indent = Pt(18)
        par.paragraph_format.first_line_indent = Pt(-12)
        par.add_run("• ")
        par.add_run(item)

title = p("7. Ambiente do estudante no Textopia_: submissão, acompanhamento e revisão")
title.style = doc.styles["Heading 1"]

h("7.1 Finalidade e escopo")
p("Este capítulo especifica a experiência do estudante desde o acesso à atividade até o encerramento de um ciclo de reescrita. No Textopia_, o estudante é o autor do texto e conserva a responsabilidade por todas as decisões de escrita. O sistema apresenta evidências, diagnósticos e orientações, mas não altera o texto nem substitui o julgamento do autor. O ambiente deve tornar visíveis o estado de cada processo, a próxima ação possível, o responsável por ela e as limitações da análise.")
p("A especificação articula requisitos pedagógicos, linguísticos, funcionais e técnicos. Os princípios pedagógicos orientam a interpretação e a interação; campos, códigos, contratos, estados e versões são regidos pelos anexos técnicos canônicos vigentes.")

h("7.2 Princípios da experiência do estudante")
bullets([
    "Agência humana: aceitar uma orientação significa concordar com ela como apoio à decisão; nenhuma alteração é aplicada automaticamente. Uma correção somente existe quando o estudante incorpora explicitamente uma alteração a uma nova versão do texto.",
    "Contestabilidade: o estudante pode aceitar, adaptar, rejeitar ou contestar uma orientação. Essas decisões não reduzem automaticamente a nota e permanecem vinculadas ao achado e à versão analisada.",
    "Finalidade formativa: o valor pedagógico da revisão se realiza quando o estudante interpreta os achados, decide como agir e produz uma nova versão.",
    "Seletividade: a revisão é priorizada, adequada ao momento formativo, limitada por ciclo e capaz de agrupar ocorrências recorrentes.",
    "Transparência: o ambiente diferencia revisão, validação, avaliação e nota, informa limitações, estados e responsabilidades e não apresenta a IA como autoridade final.",
    "Diversidade linguística: diferenças de variedade não são tratadas automaticamente como deficiência; prescrições normativas são contextualizadas pelo gênero, pela situação comunicativa, pelos objetivos e pelos interlocutores.",
    "Acessibilidade e inclusão: todas as operações devem ser utilizáveis por teclado e tecnologias assistivas, possuir alternativas não visuais e evitar pressupostos capacitistas.",
    "Privacidade e segurança: o estudante recebe informação clara sobre finalidade, acesso, retenção, serviços externos e canais para exercer direitos, sem presumir que consentimento seja a única base aplicável ao tratamento.",
])

h("7.3 Jornada do estudante")
h("7.3.1 Acesso e painel", 3)
p("O painel apresenta atividades abertas, futuras e encerradas, prazos, pendências, notificações pertinentes e o estado atual de cada trabalho. Cada cartão informa a próxima ação disponível e quem deve realizá-la. Estados vazios distinguem ausência de atividades, ausência de submissões e indisponibilidade temporária. Informações críticas não dependem apenas de cor, ícone ou posição.")
h("7.3.2 Leitura da atividade", 3)
p("Antes da submissão, o estudante acessa proposta, objetivos, gênero, interlocutores, finalidade comunicativa, critérios habilitados, limites de intervenção, política de validação, uso eventual de exemplos e existência de avaliação independente. A interface informa a publicação da atividade aplicável à submissão, de modo que alterações posteriores da configuração não modifiquem retroativamente as condições do trabalho.")
p("O conjunto de critérios mostrado resulta da interseção entre o mapeamento canônico do gênero e a seleção docente expressa por enabledCriterionCodes e, quando utilizado, disabledCriterionCodes. Códigos internos podem ser acompanhados de rótulos claros, mas não devem ser inventados pela interface.")
h("7.3.3 Produção, salvamento e submissão", 3)
p("O editor distingue rascunho, salvamento e submissão. O salvamento automático informa sucesso, pendência ou falha e permite recuperação após perda de conexão. A interface previne submissões duplicadas por operação idempotente e alerta sobre edições concorrentes em outras abas, sessões ou dispositivos.")
p("A atividade define as modalidades aceitas — texto digitado, texto colado ou arquivo —, os formatos, a codificação, o tamanho máximo e o tratamento de conteúdo incompatível. A confirmação de submissão apresenta data e hora, identificador, versão do texto e resumo das condições aplicáveis. A submissão cria uma TextVersion imutável; versões anteriores nunca são sobrescritas.")
h("7.3.4 Acompanhamento do processamento", 3)
p("O estudante acompanha separadamente os estados de submissão, revisão, validação, avaliação e nota. Para cada processo, a interface informa o estado, a próxima ação, seu responsável e, quando cabível, estimativa ou orientação de continuidade. Reprocessamento técnico, nova submissão, nova versão e novo ciclo de revisão são operações distintas.")
p("Mensagens diferenciam: falha recuperável; falha definitiva; saída incompatível com o contrato; retenção para validação humana; abstenção por conflito ou evidência insuficiente; revisão concluída sem achados autorizados; e revisão ainda incompleta. Uma nova tentativa técnica preserva a idempotência e não cria versões ou decisões duplicadas.")

h("7.4 Apresentação da revisão")
p("A revisão apresenta três unidades complementares e não intercambiáveis:")
bullets([
    "Intervenção sentencial: vinculada a exatamente um sentenceId e apresentada com trecho literal, critério, categoria, problema observável, diagnóstico, orientação, prioridade e estado de validação.",
    "Achado global estruturado: diagnóstico suprassentencial registrado em achadosGlobais, com ao menos id, criterioCodigo, categoriaCodigo, sentenceIds, diagnostico, orientacao, prioridade e confianca.",
    "Feedback global sintetizador: síntese do conjunto da revisão, sem funcionar como recipiente de diagnósticos suprassentenciais.",
])
p("Cada achado separa evidência textual, diagnóstico e orientação. A orientação deve ser coerente com o problema, executável e limitada ao que pode ser sustentado pelo texto e pela atividade. Exemplos são opcionais, explicitamente não obrigatórios e apresentados somente quando ampliam a compreensão.")
p("A interface distingue erro verificável, inadequação ao gênero ou à atividade, ambiguidade, limitação de desenvolvimento e possibilidade opcional de aprimoramento. Também apresenta evidências positivas e informa que a revisão é seletiva, não exaustiva. Ocorrências recorrentes podem ser agrupadas, com acesso aos trechos relacionados, sem transformar repetição em contagem punitiva.")
p("A confiança é um dado do processamento, não um juízo sobre o estudante. O backend a compara ao limiar configurado na atividade e aplica a política vigente de validação, retenção ou abstenção. Conflitos não resolvidos produzem abstenção, nunca interpretação improvisada pelo modelo.")

h("7.5 Decisões, contestação e trilha de aprendizagem")
p("Para cada intervenção ou achado global validamente apresentado, o estudante pode aceitar a orientação, adaptá-la, rejeitá-la ou contestá-la. A decisão é explícita, imutável e vinculada ao identificador do achado, à TextVersion e ao autor da decisão. Aceitar não modifica o texto; adaptar registra que a solução adotada difere da orientação; rejeitar conserva a autoria; contestar abre solicitação de análise humana com justificativa opcional.")
p("O sistema preserva uma trilha auditável sem confundir decisão pedagógica, permissão técnica e alteração textual. O estudante pode consultar suas decisões e o resultado da contestação. Permissões de visualizar, editar, reenviar, excluir ou compartilhar dependem do estado e da política institucional e são apresentadas antes da ação.")
p("Indicadores de processo podem registrar leitura dos achados, decisões, justificativas, padrões de reescrita e evolução entre versões. Esses indicadores apoiam a autorregulação e não reduzem aprendizagem à nota, ao número de correções ou ao número de cliques.")

h("7.6 Reescrita e versionamento")
p("A reescrita parte de uma versão identificada e produz outra TextVersion imutável mediante decisão do autor. A comparação entre versões indica inclusões, exclusões e alterações, preserva o acesso ao texto literal anterior e relaciona, quando possível, decisões e achados às mudanças realizadas. O sistema não presume que toda mudança decorra de uma orientação.")
p("A segmentação segue o contrato canônico: hash do texto, paragraphId, sentenceId, offsets de início e fim e conteúdo literal. Achados permanecem vinculados à versão analisada. Quando uma edição altera offsets ou segmentação, o sistema não reaproveita silenciosamente a localização anterior; ele mantém o histórico e exige novo processamento quando necessário.")
p("Um ciclo formativo é concluído quando a nova versão é registrada e as operações exigidas pela atividade atingem estados finais válidos. Conclusão do ciclo, encerramento da atividade, término da avaliação e publicação da nota são eventos distintos.")

h("7.7 Avaliação e nota")
p("Revisão, validação da revisão, orientação pedagógica, avaliação rubricada e decisão sobre a nota são operações separadas. A avaliação considera o texto completo, evidências positivas, achados validados, recorrência e limitações, sem converter a quantidade de problemas em desconto automático.")
p("A nota calculada é o resultado determinístico produzido pelo backend conforme a rubrica, os níveis, os pesos e a fórmula canônica. A decisão docente confirma, modifica ou rejeita esse resultado. A nota publicada é o resultado liberado ao estudante após decisão docente válida. O ambiente mostra cada etapa separadamente e nunca apresenta cálculo provisório como nota final.")
p("A contestação de uma orientação, por si só, não afeta a nota. Quando a atividade não possui avaliação, o ambiente não exibe componentes de nota. Quando possui, a avaliação utiliza fluxo e permissões próprios, relacionados à submissão e à versão avaliadas.")

h("7.8 Situações excepcionais")
bullets([
    "Prazo encerrado com rascunho: o sistema preserva o conteúdo salvo, bloqueia ou permite a submissão conforme a política e informa a consequência antes de qualquer ação.",
    "Processamento interrompido: o estudante recebe estado específico, preservação dos dados e opção de nova tentativa técnica quando autorizada.",
    "Baixa confiança ou conflito: a saída é retida, encaminhada à validação humana ou substituída por abstenção conforme a política da atividade.",
    "Indícios de integridade acadêmica: a comunicação é cautelosa e não acusatória; nenhum resultado automático constitui decisão, e o caso depende de análise humana autorizada.",
    "Resultado invalidado: a interface informa a natureza da invalidação sem expor dados internos sensíveis, preserva o histórico e não o apresenta como orientação válida.",
    "Ausência de achados: a interface diferencia revisão concluída sem problemas autorizados de análise limitada, abstenção ou falha.",
])

h("7.9 Acessibilidade, inclusão e linguagem")
p("Todo o fluxo — proposta, editor, submissão, acompanhamento, revisão, decisões, comparação, contestação e nota — deve operar por teclado, possuir foco visível, ordem de navegação coerente, rótulos acessíveis, compatibilidade com leitores de tela, ampliação sem perda de conteúdo e contraste suficiente. Tabelas, gráficos, ícones, cores e comparações visuais possuem alternativas textuais equivalentes.")
p("Mensagens e orientações empregam linguagem clara, respeitosa, específica e não capacitista. O sistema não pressupõe velocidade, modalidade sensorial ou padrão único de interação. O estudante pode solicitar análise humana e contestar orientações inacessíveis, inadequadas à sua variedade linguística ou incompatíveis com a situação comunicativa.")

h("7.10 Privacidade, notificações e transparência")
p("Antes do tratamento, o ambiente informa finalidade, categorias de dados, destinatários autorizados, serviços externos, período ou critério de retenção, medidas pertinentes e canais para acesso, correção, contestação e demais direitos. Controles de privacidade são compatíveis com a base jurídica e com a política institucional aplicáveis.")
p("Notificações são pertinentes ao fluxo, configuráveis quando possível e não revelam conteúdo sensível em canais externos. O estudante pode consultar histórico de eventos relevantes, versões, decisões e responsáveis sem receber dados restritos de outros participantes.")

h("7.11 Requisitos funcionais e critérios de aceitação")
bullets([
    "Cada tela apresenta estado atual, próxima ação permitida e responsável pela ação.",
    "Nenhuma orientação altera automaticamente o texto; toda nova versão depende de ação explícita do estudante.",
    "Intervenções sentenciais, achados globais e feedback global são exibidos como unidades distintas e vinculados à versão correta.",
    "Estados da interface correspondem às máquinas canônicas de atividade, submissão, revisão, validação, avaliação e nota.",
    "Operações repetidas de salvamento, submissão ou reprocessamento não criam duplicidades.",
    "A interface distingue falha, retenção, abstenção, ausência de achados e processamento em curso.",
    "Decisões e contestações permanecem auditáveis e não alteram automaticamente a avaliação.",
    "Todos os fluxos críticos passam por testes de teclado, leitor de tela, contraste, ampliação e alternativas não visuais.",
    "A versão e o hash dos contratos canônicos aplicáveis são verificados na publicação e no processamento.",
])

h("7.12 Artefatos de implementação")
p("A implementação deve manter, como artefatos derivados desta especificação: mapa da jornada; wireframes anotados; matriz de permissões por estado; catálogo de mensagens, falhas e estados vazios; matriz de notificações; casos de uso; critérios de aceitação; testes de acessibilidade; e rastreabilidade entre componentes da interface, entidades do domínio, máquinas de estados e contratos canônicos.")

h("7.13 Relação com as fontes canônicas")
p("Este capítulo descreve a experiência do estudante e não duplica contratos técnicos. O modelo de domínio, a segmentação, os esquemas JSON, as rubricas, as máquinas de estados, os códigos de critérios e categorias e as regras de versionamento pertencem aos anexos canônicos vigentes. Em caso de conflito, prevalecem esses anexos; conflitos não resolvidos impedem o processamento e produzem abstenção ou encaminhamento humano conforme a política aplicável.")

doc.save(OUT)
print(OUT)
