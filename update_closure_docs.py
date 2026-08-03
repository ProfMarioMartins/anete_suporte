from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

BASE=Path('/Users/mariomartins2/Library/CloudStorage/GoogleDrive-mario.martins@ufersa.edu.br/Meu Drive/TEXTOPIA/Documentação técnico-pedagógica')
OUT=Path('/Users/mariomartins2/Library/CloudStorage/GoogleDrive-mario.martins@ufersa.edu.br/Meu Drive/Disciplinas_em_R/ANETE/closure_updated')
OUT.mkdir(exist_ok=True)

def margins(doc):
    for s in doc.sections:
        s.top_margin=Inches(1); s.bottom_margin=Inches(1); s.left_margin=Inches(1); s.right_margin=Inches(1)
def clear(doc):
    b=doc._element.body; sp=b.sectPr
    for x in list(b):
        if x is not sp:b.remove(x)
    margins(doc)
def add_helpers(doc):
    def p(text='',style=None): return doc.add_paragraph(text,style=style)
    def h(text,level=1): return doc.add_heading(text,level=level)
    def bullets(items):
        for item in items:
            q=p(); q.paragraph_format.left_indent=Pt(18); q.paragraph_format.first_line_indent=Pt(-12); q.add_run('• '); q.add_run(item)
    def shade(cell,fill):
        pr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); pr.append(shd)
    def table(headers,rows,widths=None):
        t=doc.add_table(rows=1,cols=len(headers)); t.autofit=False
        borders=OxmlElement('w:tblBorders')
        for edge in ('top','left','bottom','right','insideH','insideV'):
            e=OxmlElement('w:'+edge); e.set(qn('w:val'),'single'); e.set(qn('w:sz'),'4'); e.set(qn('w:color'),'B8C2CC'); borders.append(e)
        t._tbl.tblPr.append(borders)
        for i,x in enumerate(headers):
            c=t.rows[0].cells[i]; c.text=x; shade(c,'E8ECEF'); c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.LEFT
            for r in c.paragraphs[0].runs:r.bold=True
        tr=t.rows[0]._tr.get_or_add_trPr(); rep=OxmlElement('w:tblHeader'); rep.set(qn('w:val'),'true'); tr.append(rep)
        for row in rows:
            cells=t.add_row().cells
            for i,x in enumerate(row): cells[i].text=str(x); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; cells[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.LEFT
        for row in t.rows:
            trPr = row._tr.get_or_add_trPr()
            if trPr.find(qn('w:cantSplit')) is None:
                trPr.append(OxmlElement('w:cantSplit'))
        if widths:
            for row in t.rows:
                for i,w in enumerate(widths):row.cells[i].width=Inches(w)
        return t
    return p,h,bullets,table

def new_from(src):
    d=Document(BASE/src); clear(d); return d

# Anexo A
nameA='Anexo técnico A - Prompt canônico de revisão do Textopia_.docx'
d=new_from(nameA); p,h,b,t=add_helpers(d)
p('Anexo técnico A — Prompt canônico de revisão do Textopia_')
p('Versão: revisao-2.0.0  |  Contrato de saída: 2.0.0  |  Estado: published  |  Fonte canônica')
h('1. Finalidade e uso')
p('Este anexo contém o texto canônico do prompt de revisão. O backend armazena o artefato como imutável, substitui apenas as variáveis declaradas e registra versão e SHA-256 em cada execução. O schema da saída pertence ao Anexo técnico F; o capítulo 9 recebe somente reprodução editorial gerada e verificada.')
h('2. Variáveis obrigatórias')
b(['REVISION_EXECUTION_ID','ACTIVITY_ID','ACTIVITY_PUBLICATION_ID','TEXT_VERSION_ID','TEXT_HASH','SEGMENTATION_VERSION','PROMPT_VERSION','CONTRACT_VERSION','MODEL_VERSION','GENRE_CATALOG_VERSION','TAXONOMY_VERSION','MAPPING_VERSION','TEXT_TYPE','CONTEXTUAL_CONFIGURATION','GENERIC_STRUCTURE_POTENTIAL','ACTIVITY_OBJECTIVES','ACTIVITY_INSTRUCTIONS','TEACHER_RULES','REVISION_SCOPE','CONFIDENCE_THRESHOLD','DELIVERY_POLICY','MAX_INTERVENTIONS','EXAMPLE_POLICY','AUTHORIZED_TAXONOMY','SENTENCE_MAP','STUDENT_TEXT'])
h('3. Texto integral do prompt')
p('PROMPT CANÔNICO DE REVISÃO PEDAGÓGICA DO TEXTOPIA_')
p('Você realiza uma revisão pedagógica estruturada de um texto estudantil. O estudante é o autor e conserva a responsabilidade por todas as decisões de escrita. Não altere automaticamente o texto, não atribua nota e não trate sua saída como autoridade final.')
p('CONTEXTO DE EXECUÇÃO')
p('Execução: {{REVISION_EXECUTION_ID}}\nAtividade: {{ACTIVITY_ID}}\nPublicação da atividade: {{ACTIVITY_PUBLICATION_ID}}\nVersão do texto: {{TEXT_VERSION_ID}}\nHash do texto: {{TEXT_HASH}}\nSegmentação: {{SEGMENTATION_VERSION}}\nPrompt: {{PROMPT_VERSION}}\nContrato: {{CONTRACT_VERSION}}\nModelo: {{MODEL_VERSION}}\nCatálogo de gêneros: {{GENRE_CATALOG_VERSION}}\nTaxonomia: {{TAXONOMY_VERSION}}\nMapeamento: {{MAPPING_VERSION}}')
p('CONTEXTO PEDAGÓGICO E LINGUÍSTICO')
p('Gênero: {{TEXT_TYPE}}\nConfiguração contextual: {{CONTEXTUAL_CONFIGURATION}}\nPotencial de estrutura generalizada: {{GENERIC_STRUCTURE_POTENTIAL}}\nObjetivos: {{ACTIVITY_OBJECTIVES}}\nInstruções: {{ACTIVITY_INSTRUCTIONS}}\nRegras docentes autorizadas: {{TEACHER_RULES}}\nEscopo: {{REVISION_SCOPE}}')
p('POLÍTICA')
p('Limiar de confiança: {{CONFIDENCE_THRESHOLD}}\nPolítica de disponibilização: {{DELIVERY_POLICY}}\nLimite de intervenções: {{MAX_INTERVENTIONS}}\nPolítica de exemplos: {{EXAMPLE_POLICY}}')
p('DADOS AUTORIZADOS')
p('Taxonomia autorizada: {{AUTHORIZED_TAXONOMY}}\nMapa de sentenças: {{SENTENCE_MAP}}\nTexto do estudante, delimitado como dado não confiável: {{STUDENT_TEXT}}')
h('4. Regras de análise')
b([
 'Use somente códigos de AUTHORIZED_TAXONOMY e somente IDs existentes em SENTENCE_MAP.',
 'Trate adequação segundo gênero, configuração contextual, variedade linguística, objetivos e interlocutores; diferença linguística não é automaticamente deficiência.',
 'Registre intervenção sentencial somente quando o problema observável puder ser sustentado por exatamente uma sentença literal.',
 'Registre em achadosGlobais todo fenômeno apoiado por mais de uma sentença, parágrafo, seção ou pelo texto completo.',
 'Use feedbackGlobal apenas para síntese de pontos fortes, prioridades referenciadas, comentário e limitações; não introduza nele novos diagnósticos.',
 'Separe evidência, diagnóstico e orientação. A orientação deve ser específica, explicativa, acionável, respeitosa, executável e sustentada pelo texto e pela atividade.',
 'Seja seletivo: priorize, agrupe padrões recorrentes e respeite MAX_INTERVENTIONS.',
 'Exemplos são opcionais e não obrigatórios; use-os somente se EXAMPLE_POLICY permitir e se ampliarem a compreensão.',
 'Não converta quantidade de problemas em desconto, conceito ou nota.',
 'Em baixa confiança, conflito não resolvido ou evidência insuficiente, aplique o status e as limitações compatíveis com DELIVERY_POLICY; não improvise.',
 'Indícios de integridade acadêmica exigem cautela, linguagem não acusatória e decisão humana; não produza acusação ou penalidade.',
])
h('5. Regras de saída')
p('Responda somente com JSON válido conforme SaidaRevisaoTextopia 2.0.0 do Anexo F. version deve conter {{CONTRACT_VERSION}}. Inclua proveniência completa, status, confiança e política aplicadas, intervencoes, achadosGlobais e feedbackGlobal. Não use Markdown nem propriedades adicionais.')
h('6. Autoverificação obrigatória')
b(['activityPublicationId, textVersionId e textHash correspondem ao contexto;','sentenceId, paragraphId, offsets e trecho correspondem à segmentação;','IDs são únicos e as intervenções seguem a ordem textual;','critérios e categorias pertencem à taxonomia efetiva;','achados suprassentenciais não foram convertidos em intervenção;','feedbackGlobal não contém diagnósticos novos;','status, limitações e abstenção são coerentes;','não há nota, acusação ou alteração automática do texto.'])
h('7. Controle de alteração')
p('Qualquer mudança textual exige nova promptVersion, novo SHA-256, fixtures válidos e avaliação linguística, pedagógica, técnica, de segurança, privacidade e acessibilidade. A versão 2.0.0 é incompatível com o contrato 1.1.0.')
d.save(OUT/nameA)

# Anexo B
nameB='Anexo técnico B - Prompt canônico de avaliação rubricada e proposta de nota do Textopia_.docx'
d=new_from(nameB); p,h,b,t=add_helpers(d)
p('Anexo técnico B — Prompt canônico de avaliação rubricada do Textopia_')
p('Versão: avaliacao-2.0.0  |  Contrato de saída: 2.0.0  |  Estado: published  |  Fonte canônica')
h('1. Finalidade e limites')
p('Este prompt seleciona níveis da rubrica e apresenta evidências. Ele não calcula nota. O backend calcula deterministicamente conforme o Anexo D, e o professor confirma, modifica ou rejeita conforme o Anexo F. A avaliação é independente da revisão, mas recebe somente achados em estados autorizados.')
h('2. Variáveis obrigatórias')
b(['GRADING_EXECUTION_ID','ACTIVITY_ID','ACTIVITY_PUBLICATION_ID','TEXT_VERSION_ID','TEXT_HASH','REVISION_EXECUTION_ID','VALIDATED_REVISION_REFERENCE','REVISION_PROMPT_VERSION','REVISION_CONTRACT_VERSION','GRADING_PROMPT_VERSION','GRADING_CONTRACT_VERSION','RUBRIC_VERSION','MAPPING_VERSION','MODEL_VERSION','AUTHORIZED_RUBRIC','TAXONOMY_RUBRIC_MAPPING','STUDENT_TEXT','VALIDATED_REVISION_PACKAGE','AUTHORIZED_FINDING_STATES','MINIMUM_FINDING_CONFIDENCE','EVALUATION_POLICY'])
h('3. Texto integral do prompt')
p('PROMPT CANÔNICO DE AVALIAÇÃO RUBRICADA DO TEXTOPIA_')
p('Avalie o texto completo segundo a rubrica autorizada. O estudante é o autor; a revisão fornece evidências, não descontos. Não calcule nota, não modifique o texto e não trate a quantidade de problemas como pontuação.')
p('IDENTIFICADORES E VERSÕES')
p('Avaliação: {{GRADING_EXECUTION_ID}}\nAtividade: {{ACTIVITY_ID}}\nPublicação: {{ACTIVITY_PUBLICATION_ID}}\nTexto: {{TEXT_VERSION_ID}}\nHash: {{TEXT_HASH}}\nRevisão: {{REVISION_EXECUTION_ID}}\nReferência validada: {{VALIDATED_REVISION_REFERENCE}}\nPrompt de revisão: {{REVISION_PROMPT_VERSION}}\nContrato de revisão: {{REVISION_CONTRACT_VERSION}}\nPrompt de avaliação: {{GRADING_PROMPT_VERSION}}\nContrato de avaliação: {{GRADING_CONTRACT_VERSION}}\nRubrica: {{RUBRIC_VERSION}}\nMapeamento: {{MAPPING_VERSION}}\nModelo: {{MODEL_VERSION}}')
p('DADOS AUTORIZADOS')
p('Rubrica: {{AUTHORIZED_RUBRIC}}\nMapeamento: {{TAXONOMY_RUBRIC_MAPPING}}\nEstados de achados: {{AUTHORIZED_FINDING_STATES}}\nConfiança mínima: {{MINIMUM_FINDING_CONFIDENCE}}\nPolítica: {{EVALUATION_POLICY}}\nTexto completo: {{STUDENT_TEXT}}\nPacote validado: {{VALIDATED_REVISION_PACKAGE}}')
h('4. Regras de avaliação')
b([
 'Verifique identidade entre activityPublicationId, textVersionId e textHash antes de avaliar.',
 'Considere texto completo, evidências positivas, achados validados, recorrência e limitações.',
 'Use somente critérios, níveis e mapeamentos autorizados; não invente nível ou peso.',
 'Achados rejeitados, invalidados, pendentes ou abaixo da confiança mínima não influenciam a avaliação.',
 'A quantidade de intervenções não se converte mecanicamente em desconto.',
 'Para cada critério, escolha avaliado, nao_aplicavel ou abstencao. nao_aplicavel depende de allowNotApplicable; abstencao exige limitação.',
 'Evidências devem ser específicas e referenciáveis ao texto ou aos IDs autorizados.',
 'Limitações da revisão ou do sistema não reduzem a nota.',
 'Indícios de integridade acadêmica não integram automaticamente o cálculo e exigem decisão humana separada.',
])
h('5. Regras de saída')
p('Responda somente com JSON válido conforme AvaliacaoRubricadaTextopia 2.0.0 do Anexo F. version deve conter {{GRADING_CONTRACT_VERSION}}. Inclua proveniência, status, cobertura da rubrica, evidências positivas, recorrências, referências autorizadas, justificativas, confiança e limitações. Não inclua nota calculada ou decisão docente.')
h('6. Autoverificação')
b(['cada criterioId ocorre uma vez e pertence à rubrica;','cada nivelId pertence ao critério quando status=avaliado;','há cobertura de todos os critérios aplicáveis;','referências pertencem ao pacote validado e aos estados autorizados;','abstenções e indisponibilidade possuem limitações;','não há cálculo, arredondamento, nota ou penalidade automática.'])
h('7. Controle de alteração')
p('Qualquer mudança exige nova gradingPromptVersion, SHA-256 e testes. A versão avaliacao-2.0.0 substitui avaliacao-nota-1.0.0 porque o prompt não propõe nota: propõe níveis; o backend calcula e o professor decide.')
d.save(OUT/nameB)

# Anexo C: preservar catálogos, atualizar versão e acrescentar gênero sistêmico-funcional.
nameC='Anexo técnico C - Registros canônicos de taxonomias, critérios, categorias e gêneros do Textopia_.docx'
d=Document(BASE/nameC); margins(d); p,h,b,t=add_helpers(d)
for section in d.sections:
    section.bottom_margin = Inches(0.7)
d.paragraphs[2].text='Versão normativa: registro-linguistico-2.0.0  |  Estado: published  |  Fonte canônica'
for par in d.paragraphs:
    revised = par.text.replace('revisionScope": "sentence"','revisionScope": "sentential_and_global"').replace('"scope": "sentence"','"scope": "sentential_and_global"')
    if revised != par.text:
        par.text = revised
h('3.10 Complemento normativo 2.0.0 — gênero e contexto')
p('Cada gênero é uma família de registros reconhecida por seu Potencial de Estrutura Generalizada. O catálogo 2.0.0 acrescenta, sem alterar os IDs estáveis, a Configuração Contextual e o GSP necessários à análise.')
t(['Campo','Obrigatoriedade','Descrição'],[
 ['contextualConfiguration.field','obrigatório','Campo: atividade social e conteúdo em realização.'],['contextualConfiguration.tenor','obrigatório','Relação: participantes, papéis e distância social.'],['contextualConfiguration.mode','obrigatório','Modo: papel da língua, canal e organização retórica.'],['gsp.mandatoryElements','obrigatório','Elementos que caracterizam o gênero.'],['gsp.optionalElements','obrigatório, pode ser vazio','Elementos possíveis, mas não definidores.'],['gsp.orderConstraints','obrigatório','Ordenação e recorrência admissíveis.'],['gsp.version','obrigatório','Versão imutável da descrição estrutural.'],
],[2.3,1.4,2.8])
h('3.11 Seleção efetiva e configuração da atividade')
p('A ActivityPublication fixa genreId, genreCatalogVersion, gspVersion, taxonomyVersion, mappingVersion, enabledCriterionCodes e disabledCriterionCodes. O conjunto efetivo é (mapped ∩ requested) − disabled. O professor restringe, mas não amplia o mapeamento publicado. Conjunto vazio, sobreposição entre habilitados e desabilitados ou código fora de mapped bloqueia a publicação.')
h('3.12 Diversidade linguística e categorias')
p('Critérios e categorias descrevem problemas observáveis, inadequações contextuais, ambiguidades, limitações de desenvolvimento e aprimoramentos opcionais. Diferença linguística não é automaticamente erro. Prescrições normativas dependem do gênero, da finalidade acadêmica, da variedade, dos objetivos e dos interlocutores.')
h('3.13 Compatibilidade')
p('Taxonomia 1.0.0 e catálogo de gêneros 1.0.0 permanecem válidos para execuções históricas. Novas ActivityPublications que usem contexto e GSP exigem registro-linguistico-2.0.0 e combinação declarada no manifesto. Exemplos são fixtures executáveis somente quando versões e checksums forem reais; marcadores com reticências são hipotéticos e não executáveis.')
for par in d.paragraphs:
    if par.text.startswith('• ') and ('Execuções históricas permanecem' in par.text or 'Toda execução registra version' in par.text or 'Testes recusam códigos' in par.text or 'Uma versão publicada' in par.text or 'O mesmo conjunto' in par.text or 'O prompt recebe apenas' in par.text):
        par.paragraph_format.space_after = Pt(0)
        par.paragraph_format.line_spacing = 1.0
        for run in par.runs:
            run.font.size = Pt(8.5)
d.save(OUT/nameC)

# Anexo D: preservar rubrica e acrescentar alinhamentos.
nameD='Anexo técnico D - Registro canônico de rubricas, níveis, pesos e cálculo da nota do Textopia_.docx'
d=Document(BASE/nameD); margins(d); p,h,b,t=add_helpers(d)
d.paragraphs[2].text='Versão normativa: rubricas-1.1.0  |  Estado: published  |  Fonte canônica'
for par in d.paragraphs:
    if 'nota proposta = 7,5' in par.text:
        par.text = par.text.replace('nota proposta = 7,5', 'nota calculada = 7,5')
for par in d.paragraphs:
    if '"proposedScore"' in par.text:
        par.text = '''{
  "version": "2.0.0",
  "calculationId": "calc-001",
  "evaluationExecutionId": "eval-001",
  "activityPublicationId": "activity-publication-001",
  "textVersionId": "text-version-001",
  "rubricVersion": "escrita-academica-1.0.0",
  "mappingVersion": "rubrica-taxonomia-1.0.0",
  "inputHash": "sha256:…",
  "rawScore": "7.500000",
  "calculatedScore": "7.5",
  "rounding": {
    "mode": "half_up",
    "decimalPlaces": 1
  },
  "state": "calculada"
}'''
h('9. Complemento normativo 1.1.0 — entrada da avaliação')
p('A avaliação recebe activityPublicationId, textVersionId, textHash, rubrica publicada, mapeamento publicado, texto completo, evidências positivas, achados em estados autorizados, recorrência e limitações. O mapeamento taxonomia–rubrica deste anexo é a fonte única dessa relação; o professor seleciona uma versão publicada compatível, mas não cria mapeamento executável ad hoc.')
h('10. Critérios não aplicáveis e abstenção')
p('Não aplicável é propriedade prevista pela rubrica e não equivale a baixo desempenho. Abstenção indica evidência insuficiente ou conflito e não atribui automaticamente nível N0. O denominador do cálculo exclui somente critérios legitimamente não aplicáveis. Limitação do sistema não reduz pontuação.')
h('11. Decisão docente e nota publicada')
p('O cálculo preserva inputHash e resultado original. A decisão docente confirma, modifica ou rejeita; modificação e rejeição exigem justificativa. A nota publicada é registro separado e somente decorre de decisão válida. Os contratos de decisão e publicação pertencem ao Anexo F; os estados pertencem ao Anexo E.')
h('12. Invariantes adicionais')
b(['activityPublicationId e textVersionId da avaliação, cálculo, decisão e nota são idênticos;','cada criterioId ocorre uma vez e cada levelId pertence ao critério;','pesos, pontos, escala e arredondamento nunca são inferidos pelo modelo;','repetir o cálculo com o mesmo inputHash produz o mesmo resultado;','retificação e anulação geram novos eventos e preservam histórico;','contestação ou rejeição de orientação não reduz automaticamente a nota.'])
d.save(OUT/nameD)

# Anexo E: preservar esquemas e acrescentar entidades/invariantes 2.0.
nameE='Anexo técnico E - Modelo canônico do domínio, máquinas de estados e segmentação textual do Textopia_.docx'
d=Document(BASE/nameE); margins(d); p,h,b,t=add_helpers(d)
d.paragraphs[2].text='Versão normativa: dominio-2.0.0  |  Estado: published  |  Fonte canônica'
for par in d.paragraphs:
    revised = par.text.replace('professor autorizado','autoridade autorizada').replace('professor/sistema','professor autorizado/sistema').replace('professor\n','validador humano\n')
    if revised != par.text:
        par.text = revised
h('12. Complemento normativo 2.0.0 — entidades acrescentadas')
t(['Entidade','Responsabilidade','Chave'],[
 ['activity_publications','Snapshot imutável da atividade e dos artefatos aplicáveis.','activity_publication_id'],['student_decisions','Aceitar, adaptar, rejeitar ou contestar achado, sem alterar automaticamente o texto.','student_decision_id'],['contests','Solicitação e decisão humana sobre orientação ou saída.','contest_id'],['validation_decisions','Decisão por intervenção, achado global ou feedback global.','validation_decision_id'],['grade_decisions','Confirmação, modificação ou rejeição do cálculo.','decision_id'],['published_grades','Nota liberada, retificada ou anulada.','published_grade_id'],['canonical_artifacts','Prompt, schema, taxonomia, rubrica ou estado publicado.','artifact_id+version'],['audit_events','Evento append-only de operação ou transição.','event_id'],
],[2.0,3.5,1.5])
h('13. Invariantes de publicação e autoria')
b(['Toda submissão referencia uma ActivityPublication imutável.','Toda TextVersion é imutável e uma reescrita cria novo ordinal.','Achados permanecem ligados à versão e à segmentação analisadas.','Aceitar orientação não altera texto; correção exige nova versão criada por decisão explícita do autor.','Contestar, adaptar ou rejeitar não altera automaticamente avaliação ou nota.','Saída original, validação, avaliação, cálculo, decisão e nota são registros separados.'])
h('14. Estados normativos atualizados')
t(['Máquina','Estados principais'],[
 ['Atividade','rascunho, em_revisao, publicada, suspensa, encerrada, arquivada'],['Submissão','rascunho, submetida, em_processamento, revisao_disponivel, reescrita_em_andamento, nova_versao_submetida, cancelada'],['Revisão','criada, em_processamento, validacao_automatica, aguardando_validacao, disponivel, retida, abstencao, invalidada, falha'],['Validação','pendente, em_analise, parcialmente_validada, validada, rejeitada, retida, expirada'],['Avaliação','criada, em_processamento, proposta_valida, parcialmente_abstida, indisponivel, concluida, falha'],['Nota','calculada, aguardando_decisao_docente, confirmada, modificada, rejeitada, publicada, retificada, anulada'],
],[1.6,4.9])
p('Transições não listadas são proibidas. Estados finais somente mudam por transição explícita de retificação, anulação, reabertura ou nova execução. A autoridade é validada por papel, institutionId, escopo e estado.')
h('15. Isolamento, concorrência e auditoria')
p('Toda consulta e relacionamento valida institutionId. Operações capazes de duplicar efeito exigem idempotencyKey; escrita concorrente usa controle de versão. Persistência do artefato e transição de estado são atômicas. Eventos registram eventId, actorId, objectType, objectId, previousState, newState, occurredAt, correlationId, origin e justificativa quando exigida.')
h('16. Compatibilidade')
p('dominio-2.0.0 é incompatível com dominio-1.0.0 para novas execuções que exigem ActivityPublication, decisões do estudante, contratos 2.0.0 ou nota publicada separada. Registros históricos permanecem legíveis por migradores explícitos e nunca recebem campos inferidos silenciosamente.')
d.save(OUT/nameE)

# Anexo F: reconstrução em dicionários normativos completos.
nameF='Anexo técnico F - Contratos canônicos dos módulos de revisão, validação, avaliação e nota do Textopia_.docx'
d=new_from(nameF); p,h,b,t=add_helpers(d)
p('Anexo técnico F — Contratos canônicos dos módulos de revisão, validação, avaliação e nota do Textopia_')
p('Versão normativa: contratos-modulos-2.1.0  |  Estado: published  |  Fonte canônica')
h('1. Finalidade e autoridade')
p('Este anexo é a fonte normativa dos contratos de troca dos módulos. O capítulo 9 contém explicações e reproduções geradas, nunca cópias manuais. Os arquivos JSON Schema executáveis publicados no registro canônico devem ser byte a byte compatíveis com estes dicionários, versões e invariantes.')
h('2. Regras comuns')
b(['Objetos usam JSON Schema draft 2020-12 e additionalProperties=false.','version é obrigatório; registros persistidos também exigem identificador estável e createdAt.','Datas seguem RFC 3339 UTC; hashes usam sha256:.','IDs são não vazios e únicos no escopo declarado.','activityPublicationId, textVersionId e textHash referem-se ao mesmo snapshot.','Campos condicionais e invariantes semânticos são validados antes da persistência.'])
h('3. SaidaRevisaoTextopia 2.0.0')
t(['Campo','Tipo','Regra'],[
 ['version','const 2.0.0','obrigatório'],['revisionExecutionId','string','único e imutável'],['activityId / activityPublicationId','string','publicação compatível'],['textVersionId / textHash','string','versão e hash verificados'],['segmentationVersion','string','contrato do Anexo E'],['promptVersion / contractVersion / modelVersion','string','proveniência'],['genreCatalogVersion / taxonomyVersion / mappingVersion','string','combinação publicada'],['confidenceThreshold','number 0..1','limiar aplicado'],['deliveryPolicy','enum','direta, validacao_humana ou reter'],['status','enum','concluida, sem_achados, retida, abstencao ou incompleta'],['createdAt','date-time','UTC'],['intervencoes','array','unidades sentenciais'],['achadosGlobais','array','contrato do Anexo E'],['feedbackGlobal','object','síntese, prioridades por ID e limitações'],
],[2.2,1.5,2.8])
p('Intervenção exige id, sentenceId, paragraphId, inicio, fim, trecho, criterioCodigo, categoriaCodigo, diagnostico, orientacao, prioridade e confianca. trecho é igual ao conteúdo literal do intervalo. IDs são únicos entre unidades; a ordem segue o texto. Estados retida, abstencao e incompleta exigem limitações estruturadas.')
h('4. ErroModuloTextopia 2.0.0')
t(['Campo','Regra'],[['version','const 2.0.0'],['errorId / correlationId','obrigatórios'],['executionId','opcional antes da criação da execução'],['stage','request, authorization, resolution, segmentation, provider, contract, semantic, persistence ou concurrency'],['code','código estável e não sensível'],['retryable','boolean'],['occurredAt','date-time UTC'],['userMessage','mensagem segura'],['retryAfterSeconds / safeDetails','opcionais e não sensíveis']],[2.4,4.1])
h('5. RevisaoValidadaTextopia 2.0.0')
t(['Campo','Regra'],[['validationId','identificador estável'],['activityPublicationId / textVersionId / revisionExecutionId','referências idênticas à revisão'],['revisionOutputRef','referência imutável, não cópia editada'],['policyVersion','política aplicada'],['status','pendente, parcialmente_validada, validada, rejeitada ou retida'],['interventionDecisions / globalFindingDecisions / feedbackDecision','decisões separadas'],['createdAt / validatedAt / validatedBy','condicionais ao estado']],[2.8,3.7])
p('Cada decisão exige decisionId, targetId, status, decidedBy, decidedAt e justificativa quando exigida. Edição preserva o original e registra somente a orientação validada. Estado global é derivado das decisões individuais.')
h('6. AvaliacaoRubricadaTextopia 2.0.0')
t(['Campo','Regra'],[['gradingExecutionId','identificador estável'],['activityId / activityPublicationId / textVersionId / textHash','snapshot único'],['revisionExecutionId / validatedRevisionRef','revisão autorizada'],['gradingPromptVersion / gradingContractVersion / revisionPromptVersion / revisionContractVersion','proveniência'],['rubricVersion / mappingVersion / modelVersion','combinação publicada'],['authorizedFindingStates','estados aceitos'],['status','em_execucao, proposta_valida, parcialmente_abstida, avaliacao_indisponivel ou falha'],['criterios','cobertura integral e IDs únicos'],['justificativaGlobal / limitacoes / createdAt','obrigatórios']],[2.9,3.6])
p('Cada critério exige criterioId, status, nivelId quando avaliado, evidenciasPositivas, intervencoesRelacionadas, achadosGlobaisRelacionados, recorrencias, justificativa, confianca e limitacoes. nivelId pertence ao critério. nao_aplicavel depende da rubrica; abstencao exige limitações.')
h('7. DecisaoDocenteNota 1.0.0')
t(['Campo','Regra'],[['decisionId / calculationId','identificadores'],['activityPublicationId / textVersionId','snapshot avaliado'],['decisionType','confirmar, modificar ou rejeitar'],['calculatedScore','resultado preservado'],['decidedScore','igual ao cálculo quando confirmar; obrigatório quando modificar; ausente quando rejeitar'],['decidedBy / decidedAt','ator e data'],['justification','obrigatória para modificar ou rejeitar']],[2.7,3.8])
h('8. NotaPublicadaTextopia 1.0.0')
t(['Campo','Regra'],[['publishedGradeId / decisionId','identificadores'],['activityPublicationId / textVersionId','snapshot avaliado'],['score / scaleMinimum / scaleMaximum','valor e escala'],['state','publicada, retificada ou anulada'],['publishedBy / publishedAt / createdAt','ator e datas'],['supersedesPublishedGradeId','obrigatório em retificação'],['reason','obrigatório em retificação ou anulação']],[2.7,3.8])
h('9. Compatibilidade e fixtures')
p('revisao-2.0.0 exige SaidaRevisaoTextopia 2.0.0. avaliacao-2.0.0 exige AvaliacaoRubricadaTextopia 2.0.0. Ambos exigem dominio-2.0.0; avaliação e cálculo exigem rubricas-1.1.0. Fixtures executáveis usam códigos do Anexo C, hashes reais e segmentação consistente. Marcadores como sha256:fixture ou reticências são hipotéticos e não executáveis.')
h('10. Publicação')
p('O manifesto relaciona URI, version, SHA-256, status, dependências e fixtures. A publicação do livro extrai as reproduções dos anexos, normaliza LF e interrompe diante de divergência. Mudança incompatível exige nova versão principal.')
d.save(OUT/nameF)

# Capítulo 13
name13='13. Apêndice - histórico da especificação técnica do Textopia_.docx'
d=new_from(name13); p,h,b,t=add_helpers(d)
p('13. Apêndice: histórico da especificação técnico-pedagógica do Textopia_')
p('Estado: histórico, informativo e não normativo. Versão editorial: historico-2.0.0.')
h('13.1 Finalidade')
p('Este apêndice registra decisões e mudanças que explicam a forma atual da documentação. Não reproduz contratos obsoletos nem substitui capítulos e anexos vigentes. A implementação consulta as fontes canônicas publicadas e seus manifests de versão e hash.')
h('13.2 Linha de evolução')
t(['Marco','Síntese'],[
 ['Especificação inicial','Revisão, avaliação e nota apareciam excessivamente acopladas; exemplos e regras eram repetidos.'],['Consolidação pedagógica','Autoria, contestabilidade, diversidade linguística, seletividade e aprendizagem por reescrita foram explicitadas.'],['Arquitetura estruturada','Intervenção sentencial, achado global estruturado e feedback global sintetizador tornaram-se unidades distintas.'],['Domínio canônico','ActivityPublication, TextVersion, segmentação, estados e persistência imutável foram formalizados.'],['Avaliação e nota','Modelo seleciona níveis; backend calcula; professor decide; nota publicada é registro separado.'],['Governança e produto','Ambientes de estudante e professor, gestão, marca e roteiro passaram a conter requisitos verificáveis.'],['Atualização 2.0','Prompts A e B, domínio E e contratos F foram alinhados às versões 2.0.0; rubricas D passaram a 1.1.0.'],
],[1.65,4.85])
h('13.3 Alterações consolidadas nos capítulos 1–12')
b([
 'Capítulos 1 e 2: fundamentos, agência humana, aprendizagem, feedback e operações pedagógicas distintas.',
 'Capítulos 3 e 4: visão funcional, arquitetura, domínio, segurança, versões e eficiência operacional.',
 'Capítulos 5 e 6: revisão seletiva, achadosGlobais, gênero, contexto, GSP, critérios e diversidade linguística.',
 'Capítulos 7 e 8: jornadas completas do estudante e do professor, estados, decisões, acessibilidade e privacidade.',
 'Capítulo 9: integração de revisão, validação, avaliação, cálculo, decisão e nota; contratos transferidos ao Anexo F.',
 'Capítulo 10: governança institucional, isolamento, identidade, operação, auditoria, qualidade e suporte.',
 'Capítulo 11: manual de marca, design system, padrões pedagógicos visuais e projeto editorial.',
 'Capítulo 12: estado de referência, portfólio, dependências, riscos, pesquisa, migração e rastreabilidade editorial.',
])
h('13.4 Alterações consolidadas nos anexos')
t(['Anexo','Versão vigente','Mudança principal'],[
 ['A','revisao-2.0.0','ActivityPublication, hash, segmentação, GSP, política, três unidades e abstenção.'],['B','avaliacao-2.0.0','Avaliação rubricada sem cálculo, com texto completo, achados autorizados, recorrência e limitações.'],['C','registro-linguistico-2.0.0','Configuração Contextual, GSP, seleção efetiva e diversidade linguística.'],['D','rubricas-1.1.0','Entrada completa, não aplicável, abstenção, cálculo, decisão e nota separada.'],['E','dominio-2.0.0','ActivityPublication, decisões, contestações, estados atualizados, isolamento e concorrência.'],['F','contratos-modulos-2.1.0','Dicionários normativos de revisão, erro, validação, avaliação, decisão e nota.'],
],[0.8,1.8,3.9])
h('13.5 Decisões conceituais preservadas')
b(['O estudante é o autor e decide sobre sua escrita.','Correção é alteração incorporada pelo autor em nova versão.','Problema observável é ocorrência identificável ligada a critério autorizado.','Achados suprassentenciais pertencem a achadosGlobais; feedbackGlobal é síntese.','Revisão, validação, orientação, avaliação e nota são operações distintas.','Problemas não são convertidos mecanicamente em descontos.','Baixa confiança e conflitos não resolvidos produzem retenção, abstenção ou análise humana.','Indícios de integridade acadêmica exigem cautela e decisão humana.'])
h('13.6 Regras de consulta e precedência')
p('Conceitos pedagógicos são consultados nos capítulos 1, 2, 5 e 6. Experiência e administração são consultadas nos capítulos 7, 8, 10 e 11. Integração é explicada no capítulo 9. Os Anexos A–F são fontes normativas de seus respectivos artefatos. Em conflito, prevalece o anexo especializado, published e compatível; conflito não resolvido bloqueia a operação.')
h('13.7 Preservação histórica')
p('Versões antigas permanecem em repositório de arquivo com acesso, data, responsável, motivo e hash. Não integram o corpo normativo nem são usadas em novas execuções. Migrações e leitores históricos são explícitos; nenhuma versão antiga é reinterpretada silenciosamente segundo contrato novo.')
h('13.8 Próximas atualizações')
p('Mudanças futuras devem atualizar o capítulo 12, os capítulos impactados, os anexos canônicos, o manifesto de compatibilidade e este histórico. A publicação editorial verifica versões, hashes, fixtures e referências cruzadas antes de gerar o livro.')
d.save(OUT/name13)

for f in OUT.glob('*.docx'): print(f)
