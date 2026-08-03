from docx import Document
from docx.shared import Pt

SRC='/Users/mariomartins2/Downloads/10. Gestão e administração da plataforma TEXTOPIA_.docx'
OUT='/Users/mariomartins2/Library/CloudStorage/GoogleDrive-mario.martins@ufersa.edu.br/Meu Drive/Disciplinas_em_R/ANETE/10. Gestão e administração da plataforma Textopia_ - versão corrigida.docx'
doc=Document(SRC)
body=doc._element.body; sect=body.sectPr
for x in list(body):
    if x is not sect: body.remove(x)

def p(text='',style=None): return doc.add_paragraph(text,style=style)
def h(text,level=2): return doc.add_heading(text,level=level)
def bullets(items):
    for item in items:
        q=p(); q.paragraph_format.left_indent=Pt(18); q.paragraph_format.first_line_indent=Pt(-12)
        q.add_run('• '); q.add_run(item)

title=p('10. Gestão e administração da plataforma Textopia_'); title.style=doc.styles['Heading 1']

h('10.1 Finalidade e escopo')
p('Este capítulo especifica a governança institucional, a administração pedagógica e a operação técnica do Textopia_. Seu objetivo é assegurar que pessoas, instituições, configurações, dados e artefatos canônicos sejam administrados com responsabilidade definida, menor privilégio, segregação de funções, rastreabilidade, segurança, acessibilidade e continuidade.')
p('A governança institucional define autoridade, políticas e prestação de contas; a administração pedagógica governa gêneros, taxonomias, rubricas e qualidade formativa; a operação técnica mantém identidade, infraestrutura, integrações, observabilidade, segurança e recuperação. O ambiente do professor permanece regido pelo capítulo 8. Contratos, estados, versões e hashes obedecem aos anexos canônicos vigentes.')

h('10.2 Princípios administrativos')
bullets([
 'Agência e finalidade pedagógica: decisões administrativas não transferem ao sistema a autoria do estudante nem a decisão pedagógica do professor.',
 'Menor privilégio e necessidade de saber: cada pessoa acessa somente os objetos e dados necessários ao seu papel e escopo.',
 'Segregação de funções: criação, revisão, aprovação e publicação de mudanças sensíveis não se concentram indevidamente na mesma pessoa.',
 'Imutabilidade e proveniência: publicações, execuções e decisões concluídas não são sobrescritas; correções geram novos registros relacionados.',
 'Privacidade e minimização: dados são tratados para finalidades explícitas, com acesso e retenção proporcionais.',
 'Contestabilidade e supervisão humana: decisões automatizadas relevantes podem ser contestadas e não substituem decisões humanas exigidas.',
 'Acessibilidade e inclusão: operações administrativas e canais de suporte são utilizáveis por pessoas com diferentes necessidades e tecnologias assistivas.',
 'Continuidade e economia: capacidade, custo e dependências externas são controlados sem comprometer segurança ou valor pedagógico.',
])

h('10.3 Estrutura de governança')
h('10.3.1 Papéis centrais',3)
bullets([
 'Proprietário do produto: prioriza evolução, aceita riscos de produto e responde pelo alinhamento institucional.',
 'Responsável pedagógico: aprova princípios, critérios formativos e impacto educacional; não publica sozinho artefato que tenha criado.',
 'Curador linguístico-pedagógico: propõe e mantém gêneros, critérios, categorias, descrições e mapeamentos em rascunho.',
 'Engenheiro de prompts: mantém prompts, fixtures e testes, sem autoridade para alterar taxonomias ou rubricas sem aprovação correspondente.',
 'Responsável técnico: responde por arquitetura, integrações, implantação, confiabilidade e compatibilidade.',
 'Responsável por segurança: governa controles, vulnerabilidades, segredos, acessos privilegiados e incidentes.',
 'Responsável por privacidade: mantém inventário de tratamento, retenção, solicitações de titulares e avaliação de fornecedores.',
 'Responsável por acessibilidade: homologa fluxos críticos e acompanha regressões de inclusão.',
 'Administrador institucional: gerencia sua instituição dentro de políticas globais; não modifica fontes canônicas centrais.',
 'Auditor: consulta evidências e eventos sem permissão para alterar os objetos auditados.',
])
h('10.3.2 Segregação e decisão',3)
p('Mudanças canônicas e configurações de alto impacto seguem, no mínimo, as etapas rascunho, revisão, aprovação, publicação, descontinuação e arquivamento. Quem cria não aprova nem publica sozinho a própria mudança. Conflitos pedagógicos, técnicos, de privacidade ou segurança são registrados e encaminhados ao responsável competente; conflito não resolvido bloqueia a publicação ou produz abstenção, nunca interpretação improvisada pelo modelo.')
p('A matriz RACI identifica, para cada processo, responsável pela execução, autoridade decisória, consultados e informados. A matriz de permissões é distinta: ela determina tecnicamente quem pode executar cada operação em cada escopo e estado.')

h('10.4 Instituições, unidades, turmas e pertencimento')
p('Institution é a fronteira principal de isolamento. Uma instituição pode conter unidades, cursos e turmas; usuários podem possuir vínculos independentes com mais de uma instituição. Toda entidade institucional registra version, institutionId, identificador estável, estado, createdAt e eventos de transição.')
p('Os estados institucionais distinguem rascunho, ativa, suspensa, encerrada e arquivada. Suspensão interrompe novas operações sem apagar histórico; encerramento inicia retenção e desativação de acessos; arquivamento preserva somente o conjunto autorizado. Exclusão física obedece à política de dados e não equivale a encerramento administrativo.')
p('Criação, transferência, fusão, encerramento e reabertura exigem autorização, validação de dependências e auditoria. Turmas encerradas não recebem novas atividades ou submissões, salvo reabertura explícita. Dados nunca são movidos entre instituições apenas pela alteração de um identificador.')
h('10.4.1 Isolamento institucional',3)
p('Toda consulta, comando, fila, arquivo, cache, índice, log e relacionamento valida institutionId no backend. Identificadores previsíveis não concedem acesso. Compartilhamento interinstitucional exige finalidade, autorização de ambas as partes, escopo, prazo e registro próprio. Testes automatizados verificam isolamento e impedem referências cruzadas indevidas.')

h('10.5 Identidade, autenticação e autorização')
p('O ciclo da identidade compreende convite, ativação, vínculo institucional, alteração de papel, suspensão, bloqueio, desligamento e anonimização ou retenção autorizada. Autenticação federada é preferida quando disponível; autenticação local exige política de credenciais, recuperação segura, proteção contra abuso e autenticação multifator para funções privilegiadas.')
p('Sessões possuem expiração, renovação, revogação e registro de dispositivo e risco compatíveis com a política institucional. APIs e contas de serviço usam credenciais próprias, escopo mínimo, rotação e proprietário identificado. Contas compartilhadas são proibidas.')
p('Autorização combina papel, permissão, institutionId, escopo do objeto, estado e prazo. Delegações são explícitas, temporárias e auditáveis. A revisão periódica remove acessos obsoletos; desligamento e mudança de função revogam sessões e permissões sem demora indevida.')
h('10.5.1 Acesso privilegiado',3)
p('Acesso emergencial exige justificativa, duração curta, autenticação reforçada e revisão posterior. Personificação de usuário é desabilitada por padrão; quando indispensável ao suporte, exige consentimento ou fundamento autorizado, aviso visível, escopo limitado e trilha completa. Conteúdo pedagógico e metadados técnicos possuem permissões separadas.')

h('10.6 Configuração institucional')
p('Configurações institucionais são registros versionados com estados de rascunho, revisão, aprovação e publicação. Uma publicação imutável define identidade, idioma, localização, políticas autorizadas, integrações, notificações e limites. Alteração produz nova versão e não modifica retroativamente atividades ou execuções existentes.')
p('A ordem de precedência é: contratos e políticas globais obrigatórios; política institucional publicada; ActivityPublication; parâmetros docentes autorizados. Camada inferior pode restringir, mas não ampliar permissões ou contrariar a superior. Conflitos bloqueiam a operação.')
p('Limites são classificados em pedagógicos, técnicos, financeiros e contratuais e podem ser aplicados por instituição, turma, atividade, usuário, operação ou período. Ao atingir um limite, o sistema preserva o trabalho, informa a condição e aplica bloqueio, fila, degradação segura ou autorização excepcional conforme política explícita.')

h('10.7 Governança dos artefatos canônicos')
p('São governados: prompts, contratos, schemas, gêneros, taxonomias, critérios, categorias, mapeamentos, rubricas, níveis, pesos, fórmulas, máquinas de estados, fixtures, versões de segmentação e combinações de modelos. Os Anexos A a F são fontes normativas do livro; artefatos executáveis correspondentes permanecem versionados e imutáveis no registro canônico.')
p('Cada artefato possui id, version semântica, status, SHA-256, createdAt, autor, revisor, aprovador, publishedAt, compatibilidades e histórico. Somente o estado published pode ser usado em nova execução. Mudança incompatível exige nova versão principal; correção compatível segue as regras do registro especializado.')
h('10.7.1 Fluxo de publicação',3)
bullets([
 'Criar rascunho a partir de versão identificada.',
 'Executar revisão linguística, pedagógica, técnica, de segurança, privacidade e acessibilidade conforme o objeto.',
 'Executar fixtures, testes contratuais, regressão e avaliação de impacto.',
 'Aprovar com segregação de funções e registrar pendências aceitas.',
 'Publicar versão imutável e atualizar o manifesto de compatibilidade.',
 'Implantar gradualmente, observar indicadores e permitir rollback para combinação compatível.',
 'Descontinuar sem alterar execuções históricas; arquivar após o período definido.',
])
p('Professores selecionam apenas opções publicadas e parâmetros autorizados. Eles não editam prompts canônicos nem criam códigos executáveis por configuração local. Alteração de modelo ou provedor exige análise comparativa de qualidade pedagógica, vieses, segurança, privacidade, latência, disponibilidade e custo.')

h('10.8 Privacidade e ciclo de vida dos dados')
p('O inventário de tratamento relaciona finalidade, categoria de dado, titulares, base aplicável, origem, destinatários, localização, controles, retenção e descarte. Cadastro, texto, segmentação, revisão, validação, avaliação, nota, contestação, auditoria, suporte e telemetria são conjuntos distintos.')
p('A minimização é verificada por campo e finalidade. Provedores externos recebem somente o necessário, mediante avaliação e contrato compatíveis. Textos e dados sensíveis não são usados em desenvolvimento, demonstração ou teste sem autorização e proteção adequadas.')
h('10.8.1 Retenção e descarte',3)
p('A política define prazo ou critério por categoria. Encerramento de turma ou conta inicia as ações previstas, mas não implica exclusão imediata indiscriminada. Eliminação, anonimização, pseudonimização, bloqueio e arquivamento são operações diferentes e auditáveis. Cópias de segurança seguem expiração controlada e não restauram permanentemente dados já eliminados.')
h('10.8.2 Direitos e solicitações',3)
p('Solicitações de acesso, correção, contestação, oposição, portabilidade ou eliminação possuem identificador, estado, responsável, prazo e decisão. O sistema verifica identidade, escopo e obrigações aplicáveis, preservando somente a auditoria mínima autorizada. A interface não presume consentimento como única base para todo tratamento.')

h('10.9 Segurança da informação')
bullets([
 'Criptografia em trânsito e em repouso, incluindo bancos, arquivos, backups e artefatos restritos.',
 'Gestão de segredos com cofre, escopo, rotação, revogação e proibição de exposição em código ou logs.',
 'Validação de autorização no backend para toda operação.',
 'Análise de dependências, código, infraestrutura e imagens, com tratamento de vulnerabilidades por severidade.',
 'Validação de tipo real, tamanho e conteúdo de arquivos enviados, com isolamento de processamento.',
 'Limites de taxa, cotas, proteção contra automação abusiva e circuit breaker para integrações.',
 'Separação formal entre instruções do sistema, configurações autorizadas e dados não confiáveis.',
 'Defesas e testes contra injeção de prompt, exfiltração, conteúdo executável e ampliação indevida de taxonomia ou permissão.',
])
h('10.9.1 Ambientes',3)
p('Desenvolvimento, teste, homologação, produção e recuperação são isolados por contas, redes, segredos e dados. Dados reais não são copiados para ambientes não produtivos sem processo autorizado de minimização e proteção. Acesso à produção é limitado, temporário quando possível e integralmente auditado.')

h('10.10 Auditoria e proveniência')
p('Eventos de auditoria são append-only e contêm eventId, version, eventType, actorId, actorType, institutionId, objectType, objectId, action, previousState, newState, occurredAt, correlationId, origin, justification quando exigida e resultado. Integridade, acesso e retenção dos eventos são protegidos e periodicamente verificados.')
p('Auditoria de aplicação, log técnico e telemetria são fluxos distintos. Logs comuns não armazenam texto integral, saída bruta ou segredos. Publicações, alterações de acesso, configurações, exportações, descarte, execuções, validações, decisões docentes, notas, contestações, incidentes e ações corretivas possuem eventos específicos.')

h('10.11 Operação, implantação e confiabilidade')
p('Implantações usam artefatos reproduzíveis, migrações compatíveis, aprovação, observação e rollback. Mudanças de alto risco empregam feature flag, canário ou liberação gradual. Banco e aplicação mantêm compatibilidade durante a janela de implantação; falha parcial não deixa schema ou estado intermediário utilizável.')
p('Filas, tarefas e integrações aplicam idempotencyKey, deduplicação, timeout, espera progressiva, limite de tentativas e fila de falhas. Tarefas travadas, mensagens duplicadas e respostas tardias são detectadas. Indisponibilidade do provedor de IA produz retenção, fila ou abstenção segura, sem perda de texto e sem resultado improvisado.')
h('10.11.1 Objetivos e indicadores',3)
p('A operação mantém objetivos de disponibilidade, latência, durabilidade e recuperação para submissão, salvamento, revisão, validação, avaliação e nota. Indicadores incluem taxa de sucesso, filas, tempo por estado, tentativas, abstenções, custos e consumo de tokens. Orçamentos de erro e limites de custo orientam liberação, degradação e interrupção controlada.')

h('10.12 Continuidade, backup e recuperação')
p('O plano identifica bancos, arquivos, registros canônicos, configurações, auditoria e segredos que exigem proteção. Para cada conjunto define frequência, retenção, criptografia, isolamento, objetivo de ponto de recuperação e objetivo de tempo de recuperação. Backups críticos permanecem protegidos contra exclusão ou comprometimento da conta principal.')
p('Restaurações são testadas periodicamente em ambiente isolado, com responsável, evidência e critério de sucesso. O teste verifica integridade, relacionamentos, hashes, estados e capacidade de retomar processamento idempotente. O plano de continuidade define comunicação, operação degradada e reconciliação das transações realizadas durante a interrupção.')

h('10.13 Gestão de incidentes')
p('Incidentes são classificados por segurança, privacidade, disponibilidade, integridade dos dados, qualidade pedagógica e cálculo ou publicação de nota. Cada incidente possui id, severidade, líder, participantes, início, escopo, usuários afetados, contenção, comunicação, recuperação e encerramento.')
p('A resposta contempla detecção, triagem, contenção, preservação de evidências, erradicação, recuperação, comunicação e análise posterior sem culpabilização. Ações corretivas possuem responsável, prazo e verificação de eficácia. Exceções emergenciais expiram e são revisadas.')

h('10.14 Qualidade e homologação')
p('Qualidade é avaliada nas dimensões técnica, linguística, pedagógica, algorítmica, de segurança, privacidade, acessibilidade e operação. Cada dimensão possui responsáveis e critérios de liberação. Aprovação técnica não substitui aprovação pedagógica, e nenhuma delas elimina controles de segurança e privacidade.')
bullets([
 'Testes unitários, integração, contrato, migração e ponta a ponta.',
 'Fixtures válidos e inválidos para prompts, schemas, hashes e combinações de versões.',
 'Testes de idempotência, concorrência, isolamento institucional, falha e recuperação.',
 'Testes de teclado, leitores de tela, contraste, ampliação e alternativas não visuais.',
 'Avaliação por gênero, variedade linguística, critério e grupo, incluindo falsos positivos e falsos negativos.',
 'Testes de linguagem acusatória, integridade acadêmica, abstenção e supervisão humana.',
 'Testes do cálculo determinístico, decisão docente, publicação, retificação e anulação da nota.',
])
p('Defeitos são classificados por severidade, impacto, dimensão, responsável e prazo. Regressões acima do limite definido bloqueiam a liberação. Exceções exigem aceite de risco, prazo de expiração e plano corretivo.')

h('10.15 Governança pedagógica e linguística')
p('Mudanças que afetem gêneros, critérios, categorias, rubricas, feedback, seletividade, limites por ciclo ou exemplos passam por análise pedagógica e linguística. A homologação verifica especificidade, explicabilidade, acionabilidade, respeito, oportunidade, diversidade linguística, contestabilidade e preservação da autoria.')
p('Indicadores pedagógicos incluem distribuição de achados, evidências positivas, abstenções, contestações, alterações docentes, decisões dos estudantes e evolução entre versões. Nenhum indicador isolado mede aprendizagem; quantidade de problemas, correções ou cliques não é convertida automaticamente em desempenho.')
p('Indícios de integridade acadêmica são tratados com cautela, linguagem não acusatória, acesso restrito e decisão humana. O sistema monitora vieses e não transforma diferença linguística em deficiência.')

h('10.16 Acessibilidade e inclusão')
p('Painéis, formulários, tabelas, alertas, logs consultáveis, relatórios e suporte devem operar por teclado, possuir foco visível, rótulos acessíveis, contraste, ampliação sem perda e alternativas textuais. Informações não dependem exclusivamente de cor, ícone ou posição. Mudanças de interface passam por homologação de acessibilidade e regressões bloqueiam fluxos críticos.')
p('A linguagem administrativa é clara e não capacitista. Canais de suporte oferecem modalidades acessíveis, e prazos e autenticação consideram necessidades legítimas sem reduzir segurança.')

h('10.17 Suporte e comunicação')
p('Suporte ao estudante, professor e administrador possui canais, escopo e base de conhecimento próprios. Chamados registram categoria, severidade, prioridade, instituição, objeto afetado, dados minimizados, responsável, prazo, estado e encaminhamento. Textos estudantis não são copiados para canais inadequados.')
p('Escalonamento distingue dúvidas de uso, orientação pedagógica, contestação, privacidade, segurança, incidente e defeito. Prazos de resposta e resolução variam por severidade. Manutenções, falhas e incidentes são comunicados com clareza, acessibilidade e atualização periódica.')
p('A base de conhecimento é versionada, revisada por responsáveis e vinculada às versões do produto. Conteúdo obsoleto é descontinuado sem apagar o histórico necessário.')

h('10.18 Métricas, custos e capacidade')
p('O catálogo separa métricas operacionais, administrativas, pedagógicas, de segurança, privacidade e suporte. Cada indicador possui definição, fonte, unidade, janela, responsável, limiar, finalidade e restrições de acesso. Painéis agregados aplicam minimização e proteção contra reidentificação.')
p('Custos são acompanhados por instituição, modelo, operação e período. Limites e alertas detectam consumo anormal. Otimizações de custo não podem remover contexto necessário, reduzir a qualidade pedagógica abaixo dos critérios de homologação ou enfraquecer controles.')

h('10.19 Registros administrativos obrigatórios')
bullets([
 'instituições, vínculos, papéis, delegações e revisões de acesso;',
 'configurações publicadas, exceções e alterações de precedência;',
 'artefatos canônicos, versões, hashes, compatibilidades e publicações;',
 'execuções, falhas, reprocessamentos e transições de estado;',
 'exportações, solicitações de titulares, retenção e descarte;',
 'incidentes, comunicações, decisões, ações corretivas e testes de recuperação;',
 'homologações, defeitos, aceitações de risco e liberações;',
 'decisões docentes, notas publicadas, retificações e anulações.',
])

h('10.20 Critérios de aceitação')
bullets([
 'Nenhuma operação acessa objeto de outra instituição sem compartilhamento formal autorizado.',
 'Toda ação privilegiada é autorizada no backend e produz evento de auditoria íntegro.',
 'Artefatos não publicados, incompatíveis ou com hash divergente não entram em execução.',
 'Mudanças sensíveis respeitam segregação de funções e possuem rollback ou plano de recuperação.',
 'Dados e logs respeitam finalidade, minimização, acesso, retenção e descarte.',
 'Backups restauráveis atendem aos objetivos documentados e são testados periodicamente.',
 'Incidentes e chamados seguem classificação, responsabilidade, prazo e escalonamento definidos.',
 'Homologação cobre técnica, pedagogia, linguística, segurança, privacidade e acessibilidade.',
 'Operações repetidas ou concorrentes não duplicam publicações, execuções, decisões ou notas.',
 'A administração não altera texto do estudante nem substitui decisões humanas exigidas.',
])

h('10.21 Artefatos de implementação')
p('A implementação deve manter: matriz RACI; matriz de papéis, permissões e escopos; modelo de pertencimento institucional; catálogo de configurações e precedência; manifesto de compatibilidade; política e inventário de dados; matriz de retenção; modelo de evento de auditoria; catálogo de segredos; plano de segurança; plano de incidentes; catálogo de serviços e níveis; plano de capacidade e custos; plano de backup e recuperação; matriz de testes e homologação; catálogo de suporte; e matriz de rastreabilidade entre requisitos, controles, evidências e responsáveis.')

h('10.22 Relação com as fontes canônicas')
p('Este capítulo descreve governança e operação e não duplica prompts, schemas, taxonomias, rubricas ou máquinas de estados. Os Anexos A a F permanecem fontes canônicas. Em caso de conflito, prevalece o anexo especializado compatível e publicado; conflito não resolvido bloqueia a operação, produz abstenção ou exige decisão humana autorizada.')

doc.save(OUT)
print(OUT)
