from docx import Document
from docx.shared import Pt

SRC = "/Users/mariomartins2/Downloads/8. Ambiente do professor - planejamento, orientação e avaliação.docx"
OUT = "/Users/mariomartins2/Library/CloudStorage/GoogleDrive-mario.martins@ufersa.edu.br/Meu Drive/Disciplinas_em_R/ANETE/8. Ambiente do professor no Textopia_ - versão corrigida.docx"

doc = Document(SRC)
body = doc._element.body
sect = body.sectPr
for child in list(body):
    if child is not sect:
        body.remove(child)

def p(text="", style=None):
    return doc.add_paragraph(text, style=style)

def h(text, level=2):
    return doc.add_heading(text, level=level)

def bullets(items):
    for item in items:
        par = p()
        par.paragraph_format.left_indent = Pt(18)
        par.paragraph_format.first_line_indent = Pt(-12)
        par.add_run("• ")
        par.add_run(item)

title = p("8. Ambiente do professor no Textopia_: planejamento, orientação e avaliação")
title.style = doc.styles["Heading 1"]

h("8.1 Finalidade e escopo")
p("Este capítulo especifica o ambiente no qual o professor planeja e publica atividades, acompanha processos de escrita, valida revisões quando autorizado, oferece orientação pedagógica, conduz avaliações rubricadas e decide sobre a nota. Essas operações são distintas e relacionadas. O ambiente deve preservar a autoria do estudante, a responsabilidade pedagógica do professor e a rastreabilidade das operações do sistema.")
p("O estudante continua sendo o autor do texto. O professor define as condições da atividade, interpreta evidências, orienta o processo e toma decisões pedagógicas; não incorpora alterações ao texto em nome do estudante. A inteligência artificial apoia operações delimitadas, mas não é autoridade final. Princípios pedagógicos orientam a interpretação; formatos, campos, códigos, estados, versões e cálculos obedecem aos anexos técnicos canônicos vigentes.")

h("8.2 Papéis, responsabilidades e permissões")
bullets([
    "Professor responsável: cria, revisa, publica e encerra atividades; acompanha sua turma; valida saídas quando a política exigir; realiza a decisão docente sobre a nota.",
    "Professor colaborador ou substituto: atua apenas nas turmas, atividades e operações explicitamente delegadas, com registro de identidade, período e escopo da delegação.",
    "Equipe pedagógica ou validador humano autorizado: pode analisar saídas retidas ou contestações conforme permissão institucional, sem adquirir automaticamente poderes de avaliação ou publicação de nota.",
    "Administrador: governa registros canônicos, integrações, políticas institucionais e permissões globais; essas funções não são confundidas com a configuração pedagógica cotidiana.",
    "Sistema: valida contratos, controla transições, registra eventos, executa prompts versionados, calcula resultados determinísticos e impede ações incompatíveis com o estado ou a permissão.",
])
p("A matriz de permissões deve indicar, por papel e estado, quem pode visualizar, criar, editar, publicar, validar, reprocessar, avaliar, decidir, exportar, reabrir, cancelar e arquivar. Toda delegação e operação sensível integra a trilha de auditoria.")

h("8.3 Painel docente")
p("O painel organiza turmas, atividades e trabalhos por estado, prazo e responsabilidade. Cada item informa o estado atual, a próxima ação permitida, seu responsável e as pendências relevantes. Alertas distinguem atraso, falha técnica, saída retida, validação pendente, contestação, avaliação incompleta e nota ainda não publicada.")
p("O painel permite filtros acessíveis e não depende exclusivamente de cor, ícone ou posição. Estados vazios distinguem ausência de turmas, atividades, submissões, revisões ou avaliações. Métricas operacionais não são apresentadas como indicadores de aprendizagem.")

h("8.4 Planejamento e publicação da atividade")
h("8.4.1 Rascunho e publicação", 3)
p("A atividade nasce como rascunho editável. A publicação valida a configuração e cria uma ActivityPublication imutável, identificada por activityPublicationId, versão e hash. Cada submissão permanece vinculada à publicação vigente no momento do envio. Alterações posteriores geram nova publicação e não modificam retroativamente submissões existentes.")
p("O professor pode duplicar, arquivar, cancelar ou criar nova publicação conforme as permissões e a máquina de estados. Reabertura, prorrogação e exceção individual são operações explícitas, justificadas quando exigido e auditáveis.")
h("8.4.2 Configuração pedagógica", 3)
p("A configuração reúne proposta, objetivos pedagógicos, objetivos comunicativos, interlocutores, finalidade, gênero, instruções, modalidades de entrega, prazos e número máximo de ciclos. O gênero é selecionado no registro canônico e pode ser apresentado com sua Configuração Contextual — Campo, Relação e Modo — e seu Potencial de Estrutura Generalizada, incluindo elementos obrigatórios, opcionais e ordenação admissível.")
p("O sistema verifica a coerência entre gênero, situação comunicativa, objetivos, instruções, critérios e rubrica. Diferenças linguísticas não são tratadas automaticamente como deficiência, e prescrições normativas devem ser justificadas pela finalidade acadêmica e pelas condições da atividade.")
h("8.4.3 Critérios e categorias da revisão", 3)
p("O professor seleciona critérios pelo código canônico. A configuração inclui enabledCriterionCodes e, quando necessário, disabledCriterionCodes. O conjunto efetivo é a interseção entre o mapeamento canônico do gênero e a seleção docente, descontadas as desativações válidas. Categorias somente podem ser habilitadas quando compatíveis com os critérios efetivos.")
p("A interface não permite criar silenciosamente códigos locais nem usar registros descontinuados ou incompatíveis. Mudanças em taxonomias produzem nova versão canônica e não alteram processamentos anteriores.")
h("8.4.4 Política da revisão", 3)
p("O professor configura prioridade, limites de intervenções por ciclo, agrupamento de padrões, uso opcional de exemplos, limiar de confiança e política de disponibilização. Exemplos são explicitamente não obrigatórios e empregados apenas quando ampliarem a compreensão.")
p("A política determina se a saída pode seguir diretamente ao estudante, se deve ser retida para validação humana ou se exige validação em condições específicas. O backend compara a confiança ao limiar da atividade e aplica abstenção, retenção ou encaminhamento humano. Conflitos não resolvidos segundo a ordem de precedência canônica produzem abstenção.")
h("8.4.5 Privacidade e prazos", 3)
p("O professor visualiza as políticas institucionais de privacidade e retenção aplicáveis, mas não pode substituí-las por opções locais incompatíveis. Pode configurar apenas parâmetros autorizados. Prazos de submissão, reescrita, validação e avaliação são distintos, com regras explícitas para atraso, prorrogação, reabertura e exceções.")

h("8.5 Pré-visualização, validação e publicação")
p("Antes da publicação, o sistema executa validações de completude, compatibilidade de códigos, versões e hashes, coerência entre gênero e critérios, rubrica, política de confiança, prazos e permissões. A pré-visualização permite ao professor conferir a proposta, a experiência do estudante e exemplos hipotéticos não executáveis, sem usar textos reais da turma.")
p("Instruções docentes são tratadas como dados delimitados e não podem modificar o prompt canônico, desativar salvaguardas ou contradizer contratos e princípios superiores. A publicação é bloqueada quando houver incompatibilidade estrutural; divergências pedagógicas não resolvidas são apresentadas para decisão humana autorizada.")

h("8.6 Acompanhamento pedagógico")
p("O professor acompanha submissões, TextVersions, ciclos, revisões, validações, avaliações, contestações e notas por máquinas de estados separadas. Reprocessamento técnico, nova análise, nova submissão e nova versão são operações diferentes. O sistema informa sempre o objeto afetado e preserva o histórico.")
p("Indicadores de processo podem abranger leitura dos achados, decisões autorais, justificativas, reescritas e evolução entre versões. Eles apoiam a interpretação pedagógica e não medem aprendizagem somente pela nota, pela quantidade de problemas, pelo número de correções ou por métricas de interação isoladas.")
p("Padrões individuais são apresentados com cautela, contexto e possibilidade de contestação. Padrões da turma são agregados segundo limiares institucionais de anonimização, sem trechos identificáveis, pequenos grupos reidentificáveis ou exposição de estudantes. Análise coletiva, análise individual e relatório institucional possuem finalidades e permissões distintas.")

h("8.7 Validação da revisão")
p("A validação atua sobre uma saída preservada e vinculada ao texto, à atividade, ao prompt, às taxonomias e às respectivas versões e hashes. O professor ou validador autorizado não substitui a saída original: sua decisão é registrada como novo evento, com autoria, data, objeto, estado anterior, estado posterior e justificativa quando exigida.")
p("A revisão contém três unidades distintas: intervenção sentencial vinculada a exatamente um sentenceId; achado global estruturado em achadosGlobais; e feedbackGlobal sintetizador, que não recebe diagnósticos suprassentenciais. O validador analisa evidência, diagnóstico e orientação separadamente e confirma a correspondência com a versão literal do texto.")
p("Para cada unidade aplicável, a decisão pode validar, invalidar, editar a orientação dentro dos limites autorizados ou solicitar nova análise. Validação parcial é permitida quando prevista pelo contrato. Editar a orientação preserva conteúdo original, autoria da edição e proveniência; não autoriza alterar silenciosamente evidência, critério, categoria ou diagnóstico.")
p("Nova análise cria resultado versionado e não apaga o anterior. Reprocessamento técnico reutiliza a mesma intenção de operação com idempotência. Achados duplicados, contraditórios, incompatíveis com o contrato, associados à versão errada ou com offsets inválidos não são disponibilizados como válidos.")
p("A revisão é seletiva, priorizada e não exaustiva. Orientações devem ser específicas, explicativas, acionáveis, respeitosas, oportunas, coerentes com o problema e sustentadas pelo texto e pela atividade. Evidências positivas também podem ser validadas e apresentadas.")

h("8.8 Orientação pedagógica e contestação")
p("Orientação pedagógica, validação da revisão, avaliação rubricada e decisão sobre a nota são operações distintas. O professor pode complementar a orientação sem transformar sua sugestão em alteração automática. Uma correção somente existe quando o estudante incorpora explicitamente uma mudança a uma nova versão.")
p("O professor acompanha decisões do estudante — aceitar, adaptar, rejeitar ou contestar — sem tratá-las como obediência ou desobediência. Rejeição e contestação não reduzem automaticamente a nota. A contestação abre análise humana vinculada ao achado e à versão; sua resposta registra decisão, justificativa, responsável e encaminhamento, com possibilidade de reabertura segundo a política institucional.")

h("8.9 Avaliação rubricada")
p("A avaliação utiliza fluxo independente da revisão, embora possa receber achados validados como evidências. O professor seleciona uma rubrica canônica compatível com a atividade, contendo critérios, níveis, pesos, escala e fórmula. A rubrica e sua versão permanecem vinculadas à avaliação e à TextVersion avaliada.")
p("O prompt canônico de avaliação recebe a atividade, o texto completo, a rubrica, evidências positivas, achados validados, recorrência e limitações. Não transforma a quantidade de problemas em desconto automático. Abstenções e limitações permanecem visíveis e impedem conclusões não sustentadas.")
p("O professor aciona a operação pedagógica de iniciar avaliação; a interface não expõe a invocação técnica do prompt como se fosse decisão docente. Operações idempotentes e controle de concorrência impedem avaliações duplicadas ou simultaneamente conflitantes.")

h("8.10 Cálculo, decisão docente e publicação da nota")
p("A nota calculada é o resultado determinístico produzido pelo backend conforme a rubrica canônica, os níveis, os pesos e a fórmula. O modelo não calcula livremente a nota final. A decisão docente confirma, modifica ou rejeita o resultado; modificação ou rejeição exige justificativa e preserva o cálculo original.")
p("A nota publicada é o resultado liberado após decisão docente válida. Nota calculada, decisão docente e nota publicada possuem estados, autores e registros próprios. A interface impede apresentar proposta ou cálculo provisório como resultado final e associa a nota à atividade, à rubrica e à versão textual avaliadas.")
p("Reabertura, retificação ou anulação seguem transições autorizadas, com justificativa, auditoria e nova comunicação ao estudante. A publicação de nota não altera revisão, validação ou texto, e a conclusão de uma dessas operações não encerra automaticamente as demais.")

h("8.11 Integridade acadêmica")
p("Indícios de integridade acadêmica são apresentados com cautela, linguagem não acusatória, acesso restrito e indicação explícita de limitações. Nenhuma saída automatizada constitui acusação, prova ou decisão. O professor ou instância humana competente analisa o contexto e registra uma decisão segundo a política institucional. Esses indícios não são incluídos automaticamente no feedback coletivo nem convertidos mecanicamente em desconto.")

h("8.12 Exportação, retenção e auditoria")
p("Exportações são autorizadas por finalidade, papel e escopo. Texto, revisão, avaliação, nota, dados agregados e trilha de auditoria são conjuntos distintos. O sistema aplica minimização, anonimização ou pseudonimização quando cabível, registra responsável, data, finalidade e conteúdo exportado e impede a inclusão de dados não necessários.")
p("Relatórios coletivos não reproduzem textos pessoais nem permitem reidentificação. Arquivamento pedagógico não equivale a retenção indefinida de dados pessoais. Ao final do período aplicável, o sistema executa a política institucional de eliminação, anonimização ou conservação autorizada e registra a operação.")

h("8.13 Acessibilidade e inclusão")
p("Painel, configuração, pré-visualização, validação, comparação de versões, avaliação, decisão de nota, contestação e exportação devem ser utilizáveis por teclado e tecnologias assistivas. O ambiente exige foco visível, ordem de navegação coerente, rótulos acessíveis, contraste, ampliação sem perda de conteúdo, linguagem clara e alternativas textuais equivalentes para recursos visuais.")
p("Alertas e estados não dependem somente de cor ou ícones. A interface evita pressupostos capacitistas, oferece tempo e modalidades de interação compatíveis com políticas inclusivas e permite contestar orientações linguisticamente inadequadas ou inacessíveis.")

h("8.14 Governança de prompts, taxonomias e rubricas")
p("Prompts, taxonomias, gêneros, critérios, categorias e rubricas são registros canônicos versionados. O professor configura somente parâmetros autorizados; não edita diretamente prompts canônicos nem cria códigos executáveis fora do processo de governança. Alterações canônicas pertencem a papéis administrativos ou técnicos autorizados e exigem validação, versão, hash, histórico e publicação.")
p("Cada processamento registra as versões e os hashes da atividade publicada, do prompt, das taxonomias, da rubrica e dos contratos utilizados. Incompatibilidades bloqueiam a operação ou produzem abstenção conforme a política. Conteúdo do estudante e instruções da atividade são tratados como dados, não como comandos capazes de substituir as regras do sistema.")

h("8.15 Requisitos funcionais e critérios de aceitação")
bullets([
    "Toda submissão permanece vinculada a uma publicação imutável da atividade.",
    "O conjunto efetivo de critérios é calculado a partir do gênero e da seleção docente canônica.",
    "A interface impede publicação com códigos, versões, hashes, prazos ou rubricas incompatíveis.",
    "Cada processo mostra estado, próxima ação e responsável conforme a máquina canônica correspondente.",
    "Saídas originais, validações, edições e novas análises permanecem distintas e auditáveis.",
    "Intervenções sentenciais, achadosGlobais e feedbackGlobal são validados e exibidos como unidades diferentes.",
    "Nenhuma ação docente altera automaticamente o texto do estudante.",
    "Avaliação, cálculo, decisão docente e publicação da nota permanecem separados e vinculados à versão correta.",
    "Operações repetidas ou concorrentes não geram publicações, revisões, avaliações ou notas duplicadas.",
    "Exportações respeitam finalidade, minimização, permissões, retenção e auditoria.",
    "Todos os fluxos críticos passam por testes de teclado, leitor de tela, contraste, ampliação e alternativas não visuais.",
])

h("8.16 Artefatos de implementação")
p("A implementação deve manter mapa da jornada docente; fluxo de criação e publicação; matriz de permissões por papel e estado; wireframes anotados; catálogo de mensagens, alertas e estados vazios; matriz de notificações; protótipos de validação e comparação; catálogo de relatórios; casos de uso; critérios de aceitação; testes de acessibilidade e usabilidade; e matriz de rastreabilidade entre interface, entidades, contratos e máquinas de estados.")

h("8.17 Relação com as fontes canônicas")
p("Este capítulo especifica a experiência docente e não mantém cópias manuais de prompts, esquemas, registros, rubricas ou máquinas de estados. Esses elementos pertencem aos anexos técnicos canônicos vigentes. Reproduções editoriais devem ser geradas automaticamente durante a publicação, com verificação de versão e hash. Em caso de conflito, prevalecem os anexos; conflito não resolvido produz bloqueio, abstenção ou encaminhamento humano, nunca interpretação improvisada pelo modelo.")

doc.save(OUT)
print(OUT)
