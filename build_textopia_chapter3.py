import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_run_font(run, size, bold=False):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold


def configure_style(style, size, bold=False, before=0, after=8):
    style.font.name = "Arial"
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)
    fmt = style.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = 1.15


def main():
    if len(sys.argv) != 2:
        raise SystemExit("usage: build_textopia_chapter3.py OUTPUT.docx")

    content = sys.stdin.read().strip()
    blocks = [block.strip() for block in content.split("\n\n") if block.strip()]

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    configure_style(doc.styles["Normal"], 12, False, 0, 8)
    configure_style(doc.styles["Heading 1"], 17, True, 14, 6)
    configure_style(doc.styles["List Bullet"], 12, False, 0, 4)
    doc.styles["List Bullet"].paragraph_format.left_indent = Inches(0.5)
    doc.styles["List Bullet"].paragraph_format.first_line_indent = Inches(-0.25)

    for index, block in enumerate(blocks):
        if index == 0:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.0
            set_run_font(paragraph.add_run(block), 23, True)
        elif block.startswith("3.") and block.split(" ", 1)[0].count(".") == 1:
            paragraph = doc.add_paragraph(block, style="Heading 1")
            paragraph.paragraph_format.keep_with_next = True
        elif block.startswith("• "):
            paragraph = doc.add_paragraph(block[2:], style="List Bullet")
        else:
            paragraph = doc.add_paragraph(block, style="Normal")

    props = doc.core_properties
    props.title = "3. Visão geral e funcionamento da plataforma Textopia_"
    props.subject = "Documentação técnico-pedagógica do Textopia_"
    props.author = "Equipe Textopia_"

    output = Path(sys.argv[1])
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


if __name__ == "__main__":
    main()
