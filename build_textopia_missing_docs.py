from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import json

OUT = Path('/private/tmp/textopia_missing')
OUT.mkdir(parents=True, exist_ok=True)

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def configure(doc):
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Inches(0.8)
    sec.left_margin = sec.right_margin = Inches(0.85)
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Calibri'; normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.18
    for name, size, color, before, after in [
        ('Title',22,'1F4D78',0,12),('Subtitle',11,'5B6573',0,16),
        ('Heading 1',16,'2E74B5',16,8),('Heading 2',13,'2E74B5',12,6),('Heading 3',11.5,'1F4D78',9,4)]:
        st=styles[name]; st.font.name='Calibri'; st.font.size=Pt(size); st.font.color.rgb=RGBColor.from_string(color)
        st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.keep_with_next=True
    for section in doc.sections:
        footer=section.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run('TEXTOPIA_ — documentação técnico-pedagógica').font.size=Pt(8)

def add_code(doc, value):
    p=doc.add_paragraph()
    p.style=doc.styles['Normal']; p.paragraph_format.left_indent=Inches(.2); p.paragraph_format.right_indent=Inches(.2)
    p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(8)
    for line in (value if isinstance(value,str) else json.dumps(value,ensure_ascii=False,indent=2)).splitlines():
        r=p.add_run(line+'\n'); r.font.name='Consolas'; r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string('25364A')
    p._p.get_or_add_pPr().append(OxmlElement('w:keepLines'))

def add_table(doc, headers, rows, widths=None):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
    t.rows[0]._tr.get_or_add_trPr().append(OxmlElement('w:tblHeader'))
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h; set_cell_shading(c,'D9EAF7')
        for r in c.paragraphs[0].runs: r.bold=True; r.font.size=Pt(9)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=str(v); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after=Pt(2)
                for r in p.runs: r.font.size=Pt(8.5)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)
    return t

def bullet(doc, text):
    doc.add_paragraph(text, style='List Bullet')

def make_annex():
    d=Document(); configure(d)
    p=d.add_paragraph(style='Title'); p.add_run('Anexo técnico C')
    d.add_paragraph('Registros canônicos de taxonomias, critérios, categorias e gêneros do TEXTOPIA_', style='Subtitle')
    d.add_paragraph('Versão normativa: registro-linguistico-1.0.0  |  Estado: publicada  |  Escopo: backend, curadoria e auditoria')
    d.add_heading('1. Finalidade e caráter normativo', level=1)
    d.add_paragraph('Este anexo define a fonte canônica e versionada dos dados linguístico-pedagógicos que o backend recupera antes de executar a revisão e a avaliação. Os prompts não consultam bancos nem escolhem critérios: recebem somente o contexto autorizado e resolvido pelo sistema.')
    d.add_paragraph('As propriedades version, status e os identificadores estáveis são obrigatórios em todos os registros e payloads. Um artefato publicado é imutável; qualquer alteração gera nova versão.')
    d.add_heading('2. Objetos canônicos', level=1)
    add_table(d,['Objeto','Função','Identificador'],[
        ('Taxonomia','Agrupa critérios e categorias reconhecidos pelo sistema.','version'),
        ('Critério','Define o aspecto linguístico-pedagógico observado.','code'),
        ('Categoria','Classifica o tipo de achado produzido pela revisão.','code'),
        ('Catálogo de gêneros','Versiona o repertório de gêneros textuais.','version'),
        ('Gênero','Descreve o tipo de texto selecionável na atividade.','id'),
        ('Mapeamento gênero–taxonomia','Autoriza o subconjunto aplicável a um gênero.','version + genreId'),
        ('Configuração da atividade','Fixa versões e escopo antes da submissão.','activityId + version'),
        ('Taxonomia autorizada','Payload resolvido e injetado nos prompts.','checksum + version'),
    ])
    d.add_heading('3. Esquemas JSON canônicos', level=1)
    d.add_paragraph('Os exemplos abaixo são contratos mínimos. Implementações podem acrescentar metadados, mas não podem remover campos obrigatórios nem alterar o significado dos códigos.')
    schemas = {
      '3.1 Registro de taxonomia': {
        'version':'1.2.0','status':'published','name':'Taxonomia linguístico-pedagógica',
        'criteria':[{'code':'COERENCIA_GLOBAL','name':'Coerência global','description':'Avalia a continuidade de sentidos no texto.','dimension':'textual','status':'active'}],
        'categories':[{'code':'RUPTURA_COERENCIA','name':'Ruptura de coerência','description':'Marca incompatibilidade de sentidos recuperável na sentença.','status':'active'}],
        'publishedAt':'2026-08-01T12:00:00Z'},
      '3.2 Catálogo de gêneros': {
        'version':'1.0.0','status':'published','genres':[{'id':'ARTIGO_OPINIAO','name':'Artigo de opinião','description':'Texto argumentativo voltado à defesa de um ponto de vista.','status':'active'}],'publishedAt':'2026-08-01T12:00:00Z'},
      '3.3 Mapeamento gênero–taxonomia': {
        'version':'1.0.0','status':'published','genreCatalogVersion':'1.0.0','genreId':'ARTIGO_OPINIAO','taxonomyVersion':'1.2.0','criterionCodes':['COERENCIA_GLOBAL'],'categoryCodes':['RUPTURA_COERENCIA'],'scope':'sentence','publishedAt':'2026-08-01T12:00:00Z'},
      '3.4 Configuração de análise da atividade': {
        'version':'1.0.0','activityId':'ATV-2026-001','genreId':'ARTIGO_OPINIAO','genreCatalogVersion':'1.0.0','taxonomyVersion':'1.2.0','mappingVersion':'1.0.0','revisionScope':'sentence','status':'published'},
      '3.5 Payload AUTHORIZED_TAXONOMY': {
        'version':'1.0.0','taxonomyVersion':'1.2.0','mappingVersion':'1.0.0','genreCatalogVersion':'1.0.0','genreId':'ARTIGO_OPINIAO','criteria':[{'code':'COERENCIA_GLOBAL','name':'Coerência global','description':'Avalia a continuidade de sentidos no texto.'}],'categories':[{'code':'RUPTURA_COERENCIA','name':'Ruptura de coerência','description':'Marca incompatibilidade de sentidos na sentença.'}],'checksum':'sha256:…'}
    }
    for title,obj in schemas.items():
        d.add_heading(title, level=2); add_code(d,obj)
    d.add_heading('4. Regras de validação', level=1)
    for x in [
        'version deve obedecer ao versionamento semântico MAJOR.MINOR.PATCH e nunca pode ser omitida.',
        'status aceita draft, in_review, approved, published, deprecated ou archived; apenas published pode ser usado em nova execução.',
        'code e id são estáveis, únicos no próprio registro, escritos em ASCII e nunca reutilizados com outro significado.',
        'criterionCodes e categoryCodes devem existir e estar ativos na taxonomyVersion declarada.',
        'genreId deve existir e estar ativo na genreCatalogVersion declarada.',
        'scope deve ser sentence para a revisão canônica atual; outro valor exige nova versão principal do contrato.',
        'publishedAt é obrigatório quando status é published, deprecated ou archived.',
        'checksum deve ser calculado sobre serialização determinística do payload resolvido, para auditoria.'
    ]: bullet(d,x)
    d.add_heading('5. Modelo de persistência', level=1)
    add_table(d,['Tabela/coleção','Campos principais','Restrições essenciais'],[
        ('taxonomy_versions','version, name, status, published_at','PK version; publicado imutável'),
        ('criteria','taxonomy_version, code, name, description, dimension, status','PK composta; FK taxonomy_versions'),
        ('categories','taxonomy_version, code, name, description, status','PK composta; FK taxonomy_versions'),
        ('genre_catalog_versions','version, status, published_at','PK version; publicado imutável'),
        ('genres','catalog_version, id, name, description, status','PK composta; FK catálogo'),
        ('genre_taxonomy_mappings','version, genre_catalog_version, genre_id, taxonomy_version, scope, status','unicidade por versão e gênero; FKs válidas'),
        ('genre_taxonomy_mapping_items','mapping_version, item_type, item_code','item_type criterion/category; item existente'),
        ('activity_analysis_configs','activity_id, version, genre_id, catalog_version, taxonomy_version, mapping_version, revision_scope, status','uma configuração publicada ativa por atividade'),
        ('publication_events','entity_type, entity_id, from_status, to_status, actor_id, timestamp','append-only; trilha de auditoria'),
    ])
    d.add_heading('6. Resolução do contexto de análise', level=1)
    for i,x in enumerate([
        'Carregar a configuração publicada da atividade e validar sua version.',
        'Recuperar o gênero na genreCatalogVersion fixada; não inferir o gênero a partir do texto.',
        'Recuperar o mapeamento publicado cuja mappingVersion, genreId, genreCatalogVersion e taxonomyVersion coincidam integralmente.',
        'Recuperar somente critérios e categorias ativos e listados no mapeamento.',
        'Construir AUTHORIZED_TAXONOMY, calcular checksum e registrar as versões na execução.',
        'Injetar TEXT_TYPE, TAXONOMY_VERSION e AUTHORIZED_TAXONOMY no prompt de revisão; disponibilizar o pacote validado ao prompt de avaliação.',
        'Interromper a execução em caso de ausência, incompatibilidade ou conjunto autorizado vazio.'
    ],1): d.add_paragraph(f'{i}. {x}')
    d.add_heading('7. Contratos de API', level=1)
    add_table(d,['Método e rota','Uso','Resposta mínima'],[
        ('GET /v1/taxonomies/{version}','Recuperar taxonomia publicada.','registro de taxonomia'),
        ('GET /v1/taxonomies/{version}/criteria?genreId={id}','Listar subconjunto aplicável.','critérios e categorias autorizados'),
        ('GET /v1/genre-catalogs/{version}','Recuperar catálogo publicado.','catálogo de gêneros'),
        ('GET /v1/genres/{genreId}?catalogVersion={version}','Recuperar gênero.','registro de gênero'),
        ('GET /v1/genre-taxonomy-mappings/{version}','Recuperar mapeamento.','registro de mapeamento'),
        ('POST /v1/analysis-contexts/resolve','Resolver contexto antes dos prompts.','textType, authorizedTaxonomy, checksum'),
        ('POST /v1/admin/{entity}/drafts','Criar versão em rascunho.','registro draft'),
        ('POST /v1/admin/{entity}/{id}/submit-review','Enviar à revisão.','registro in_review'),
        ('POST /v1/admin/{entity}/{id}/approve','Aprovar.','registro approved'),
        ('POST /v1/admin/{entity}/{id}/publish','Publicar versão imutável.','registro published'),
        ('POST /v1/admin/{entity}/{id}/deprecate','Descontinuar sem apagar.','registro deprecated'),
    ])
    d.add_heading('7.1 Exemplo de resolução', level=2)
    add_code(d,{'version':'1.0.0','activityId':'ATV-2026-001','submissionId':'SUB-889'})
    d.add_paragraph('Resposta 200:')
    add_code(d,{'version':'1.0.0','textType':{'genreId':'ARTIGO_OPINIAO','name':'Artigo de opinião','genreCatalogVersion':'1.0.0'},'taxonomyVersion':'1.2.0','mappingVersion':'1.0.0','authorizedTaxonomy':'<objeto conforme 3.5>','checksum':'sha256:…'})
    d.add_heading('8. Ciclo de publicação e permissões', level=1)
    d.add_paragraph('Fluxo obrigatório: draft → in_review → approved → published → deprecated → archived. Um estado não pode ser saltado. A descontinuação impede novos usos, mas preserva a reprodução de execuções anteriores.')
    add_table(d,['Papel','Permissões'],[
        ('Curador linguístico-pedagógico','Criar e editar rascunhos; propor códigos e descrições.'),
        ('Aprovador pedagógico','Revisar e aprovar; não publica a própria alteração.'),
        ('Administrador','Publicar, descontinuar e arquivar versões aprovadas.'),
        ('Professor','Selecionar gênero e configuração entre versões publicadas autorizadas.'),
        ('Backend','Ler registros publicados e resolver o contexto; não altera conteúdo.'),
        ('Auditor','Ler versões, eventos, checksums e vínculos de execução.'),
    ])
    d.add_heading('9. Compatibilidade e migração', level=1)
    for x in [
        'PATCH corrige forma sem alterar sentido, código nem contrato.',
        'MINOR acrescenta item compatível; atividades existentes permanecem fixadas à versão anterior.',
        'MAJOR altera sentido, remove item, muda escopo ou modifica estrutura obrigatória.',
        'Códigos removidos tornam-se deprecated e podem apontar deprecatedBy para um sucessor; nunca são apagados ou reutilizados.',
        'O backend rejeita mapeamento cuja taxonomyVersion ou genreCatalogVersion difira da configuração da atividade.',
        'Reprocessamento histórico usa exatamente as versões e o checksum registrados na execução original.'
    ]: bullet(d,x)
    d.add_heading('10. Contrato de erros', level=1)
    add_code(d,{'version':'1.0.0','error':{'code':'VERSION_MISMATCH','message':'As versões da configuração e do mapeamento não coincidem.','details':{'expectedTaxonomyVersion':'1.2.0','receivedTaxonomyVersion':'1.1.0'},'retryable':False}})
    add_table(d,['Código','Condição','Ação'],[
        ('TAXONOMY_VERSION_NOT_FOUND','Versão inexistente.','Interromper e corrigir configuração.'),
        ('TAXONOMY_NOT_PUBLISHED','Versão não publicada.','Impedir uso em nova execução.'),
        ('GENRE_NOT_FOUND','Gênero ausente/inativo.','Solicitar seleção válida.'),
        ('MAPPING_NOT_FOUND','Sem mapeamento compatível.','Publicar mapeamento antes da atividade.'),
        ('VERSION_MISMATCH','Versões divergentes.','Não adaptar silenciosamente.'),
        ('EMPTY_AUTHORIZED_TAXONOMY','Nenhum item autorizado.','Bloquear chamada ao modelo.'),
        ('INVALID_ACTIVITY_ANALYSIS_CONFIG','Configuração incompleta/inválida.','Corrigir e publicar nova versão.'),
    ])
    d.add_heading('11. Critérios de aceitação', level=1)
    for x in [
        'Toda execução registra version do contrato, gênero, catálogo, taxonomia, mapeamento e checksum.',
        'Testes recusam campos version ausentes em qualquer registro, resposta ou erro.',
        'Testes recusam códigos inexistentes, duplicados, inativos ou incompatíveis.',
        'Uma versão published não admite atualização ou exclusão física.',
        'O mesmo conjunto de versões produz AUTHORIZED_TAXONOMY idêntica e o mesmo checksum.',
        'O prompt recebe apenas o subconjunto autorizado e não possui instrução para consultar ou inventar taxonomias.',
        'Execuções históricas permanecem reprodutíveis após a publicação de versões novas.'
    ]: bullet(d,x)
    d.save(OUT/'Anexo técnico C - Registros canônicos de taxonomias, critérios, categorias e gêneros do Textopia_.docx')

def insert_after(paragraph, text, style=None):
    new_p=OxmlElement('w:p'); paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    p=Paragraph(new_p, paragraph._parent)
    if style: p.style=style
    p.add_run(text)
    return p

def update_ch9(source):
    d=Document(source)
    marker='Os dados linguístico-pedagógicos não são inferidos pelo modelo'
    if not any(marker in p.text for p in d.paragraphs):
        target=next(p for p in d.paragraphs if p.text.strip().startswith('A taxonomia é mantida em registro próprio'))
        insert_after(target, marker+' nem escritos diretamente no prompt. O backend os resolve a partir da configuração versionada da atividade, conforme o capítulo 6 e o Anexo técnico C. A resolução recupera o gênero no catálogo publicado, aplica o mapeamento gênero–taxonomia e injeta TEXT_TYPE, TAXONOMY_VERSION e AUTHORIZED_TAXONOMY. Ausência, incompatibilidade de versões ou conjunto autorizado vazio interrompem a execução antes da chamada ao modelo.')
    d.save(OUT/source.name)

make_annex()
src=Path('/Users/mariomartins2/Library/CloudStorage/GoogleDrive-mario.martins@ufersa.edu.br/Meu Drive/TEXTOPIA/Documentação técnico-pedagógica/9. Especificação técnica dos módulos de revisão e avaliação do Textopia_.docx')
update_ch9(src)
print(OUT)
