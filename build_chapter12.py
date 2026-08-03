from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SRC='/Users/mariomartins2/Downloads/12. Evolução futura e roteiro de desenvolvimento do TEXTOPIA_.docx'
OUT='/Users/mariomartins2/Library/CloudStorage/GoogleDrive-mario.martins@ufersa.edu.br/Meu Drive/Disciplinas_em_R/ANETE/12. Evolução futura e roteiro de desenvolvimento do Textopia_ - versão corrigida.docx'
doc=Document(SRC)
body=doc._element.body; sect=body.sectPr
for x in list(body):
    if x is not sect: body.remove(x)
for s in doc.sections:
    s.top_margin=Inches(1); s.bottom_margin=Inches(1); s.left_margin=Inches(1); s.right_margin=Inches(1)

def p(text='',style=None): return doc.add_paragraph(text,style=style)
def h(text,level=2): return doc.add_heading(text,level=level)
def bullets(items):
    for item in items:
        q=p(); q.paragraph_format.left_indent=Pt(18); q.paragraph_format.first_line_indent=Pt(-12)
        q.add_run('• '); q.add_run(item)
def shade(cell,fill):
    pr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); pr.append(shd)
def table(headers,rows,widths=None):
    t=doc.add_table(rows=1,cols=len(headers)); t.autofit=False
    borders=OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        e=OxmlElement('w:'+edge); e.set(qn('w:val'),'single'); e.set(qn('w:sz'),'4'); e.set(qn('w:color'),'B8C2CC'); borders.append(e)
    t._tbl.tblPr.append(borders)
    for i,x in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=x; shade(c,'E8ECEF'); c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.LEFT
        for r in c.paragraphs[0].runs:r.bold=True
    trPr=t.rows[0]._tr.get_or_add_trPr(); repeat=OxmlElement('w:tblHeader'); repeat.set(qn('w:val'),'true'); trPr.append(repeat)
    for row in rows:
        cells=t.add_row().cells
        for i,x in enumerate(row):
            cells[i].text=str(x); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; cells[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.LEFT
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    return t

title=p('12. Evolução futura e roteiro de desenvolvimento do Textopia_'); title.style=doc.styles['Heading 1']
p('Roteiro editorial, pedagógico e técnico — versão 1.0')

h('12.1 Finalidade, público e caráter do roteiro')
p('Este capítulo organiza a evolução do Textopia_ sem apresentar hipóteses, especificações ou pilotos como funcionalidades disponíveis. Ele serve à coordenação do produto, às equipes pedagógica, linguística, técnica e editorial e às instituições que acompanham a evolução da plataforma.')
p('O roteiro integra quatro perspectivas: produto, pedagogia, tecnologia e documentação. Uma iniciativa pode possuir maturidade, horizonte e estado de entrega diferentes. O capítulo registra intenções e critérios; somente evidência operacional validada autoriza declarar uma capacidade disponível.')
p('Data de referência editorial: 2 de agosto de 2026. Versão do roteiro: roadmap-1.0.0. A coordenação do produto mantém o conteúdo; a coordenação editorial verifica sua coerência com o livro. O histórico de decisões e versões permanece no Apêndice, enquanto este capítulo apresenta o estado vigente e as próximas hipóteses.')

h('12.2 Vocabulário de situação')
table(['Termo','Significado'],[
 ['Documentado','descrito nos capítulos ou anexos, sem comprovação de implantação'],
 ['Confirmado','evidência de funcionamento em ambiente identificado e data registrada'],
 ['Proposto','iniciativa submetida à análise, sem compromisso de entrega'],
 ['Aprovado','iniciativa autorizada, com responsável, capacidade e critérios de entrada'],
 ['Piloto','capacidade limitada a público, escopo e período definidos'],
 ['Disponível','liberação geral homologada, observável e documentada'],
 ['Exploratório','hipótese de pesquisa, sem versão-alvo ou compromisso de entrega'],
],[1.35,5.15])
p('O uso desses termos é obrigatório no livro, no painel do roteiro e na comunicação institucional. “Especificado” não equivale a “implementado”; “piloto” não equivale a “disponível”.')

h('12.3 Estado de referência')
p('A documentação técnico-pedagógica descreve uma arquitetura completa, mas este levantamento não incluiu auditoria do código, da infraestrutura ou de uma implantação produtiva. Assim, as capacidades abaixo têm estado editorial documentado e estado operacional a confirmar. A primeira iniciativa do roteiro é produzir evidência do estado real.')
table(['Capacidade documentada','Fonte principal','Estado operacional'],[
 ['Fundamentos e ação pedagógica','capítulos 1 e 2','a confirmar'],
 ['Visão e arquitetura técnica','capítulos 3 e 4','a confirmar'],
 ['Revisão pedagógica e linguística','capítulos 5 e 6; Anexos A e C','a confirmar'],
 ['Ambientes de estudante e professor','capítulos 7 e 8','a confirmar'],
 ['Revisão, avaliação e nota','capítulo 9; Anexos B, D e F','a confirmar'],
 ['Domínio, estados e segmentação','Anexo E','a confirmar'],
 ['Gestão e administração','capítulo 10','a confirmar'],
 ['Marca e design system','capítulo 11','parcialmente documentado; implementação a confirmar'],
],[2.6,2.3,1.6])
p('Até a conclusão da auditoria, o livro não deve afirmar que autenticação, revisão automática, avaliação, cálculo, relatórios, integrações, isolamento institucional ou publicação de nota estejam disponíveis em produção.')

h('12.4 Dimensões de dívida')
table(['Dimensão','Exemplos de dívida a verificar'],[
 ['Técnica','contratos executáveis, migrações, filas, idempotência, concorrência e testes'],
 ['Pedagógica','seletividade, evidências positivas, contestação e indicadores de aprendizagem'],
 ['Linguística','cobertura dos gêneros, critérios, variedades e fixtures'],
 ['Acessibilidade','teclado, leitores de tela, ampliação, contraste e linguagem clara'],
 ['Privacidade e segurança','inventário de dados, isolamento, retenção, segredos e incidentes'],
 ['Operacional','observabilidade, capacidade, custos, backup, recuperação e suporte'],
 ['Documental','divergências entre arquivos, exemplos, versões, hashes e estado real'],
],[1.7,4.8])

h('12.5 Horizontes')
table(['Horizonte','Janela após aprovação','Função'],[
 ['Fundação','0–6 meses','confirmar estado, eliminar bloqueadores e tornar a base verificável'],
 ['Consolidação','6–18 meses','entregar e homologar fluxos centrais e capacidades institucionais'],
 ['Exploração','sem data-alvo','investigar hipóteses de alto risco ou baixa evidência, sem compromisso'],
],[1.25,1.65,3.6])
p('A janela começa após aprovação formal e alocação de capacidade, não na data deste documento. Horizonte expressa planejamento temporal; maturidade expressa evidência; estado expressa andamento. A coordenação revisa os horizontes trimestralmente, limita trabalho simultâneo e registra movimentos e justificativas.')

h('12.6 Método de priorização')
p('Cada iniciativa recebe pontuação de 0 a 5 em valor pedagógico, gravidade do problema, alcance, redução de risco, acessibilidade, evidência, sustentabilidade e reversibilidade. Esforço, custo recorrente, complexidade de migração e dependência externa também recebem 0 a 5 e reduzem a prioridade relativa.')
p('Segurança crítica, integridade da nota, isolamento institucional, privacidade obrigatória e bloqueios graves de acessibilidade não competem apenas por pontuação: funcionam como requisitos de liberação. Demandas de grupos pequenos não podem ser descartadas por baixo alcance quando houver barreira de direitos, acesso ou equidade.')
p('A pontuação orienta, mas não automatiza a decisão. Produto coordena; pedagogia, linguística, tecnologia, segurança, privacidade, acessibilidade e editorial registram pareceres. Divergência não resolvida impede a aprovação.')

h('12.7 Estados e transições do item')
table(['Estado','Critério principal'],[
 ['proposto','problema e público identificados'],['em_descoberta','hipótese, evidência e alternativas em análise'],['em_revisao','pareceres e riscos em avaliação'],['aprovado','responsável, capacidade, dependências e critérios de entrada válidos'],['planejado','versão ou janela alocada'],['em_desenvolvimento','execução iniciada e observável'],['bloqueado','impedimento e responsável registrados'],['em_homologacao','implementação candidata submetida aos testes'],['piloto','público, prazo, métricas e reversão delimitados'],['liberacao_gradual','expansão monitorada por etapas'],['disponivel','homologação concluída e documentação publicada'],['retido','avanço interrompido por risco ou evidência insuficiente'],['cancelado','decisão de não prosseguir'],['revertido','liberação retirada e versão anterior restaurada'],['descontinuado','retirada concluída com histórico preservado'],
],[1.7,4.8])
p('Toda transição gera evento append-only com itemId, estado anterior, estado posterior, ator, data, justificativa, evidência e correlationId. Reabertura cria novo evento. Iniciativa substituída referencia a sucessora. Somente autoridades definidas na matriz RACI podem aprovar, liberar, cancelar ou descontinuar.')

h('12.8 Modelo canônico de item do roteiro')
bullets([
 'Identificação: itemId estável, título, versão, datas, autor e origem da demanda.',
 'Problema: evidência, público afetado, gravidade, linha de base e risco de não agir.',
 'Resultados: produto, pedagógico, técnico, operacional ou institucional, sem confundi-los com entregáveis.',
 'Escopo: objetivos, não objetivos, premissas, restrições e alternativas rejeitadas.',
 'Planejamento: horizonte, prioridade, estado, responsável, autoridade, equipe, esforço, custo e confiança da estimativa.',
 'Dependências: item, versão mínima, tipo, obrigatoriedade e responsável.',
 'Riscos: categoria, probabilidade, impacto, mitigação, contingência, proprietário e risco residual.',
 'Entrega: critérios de entrada, aceitação, definição de pronto, implantação, observabilidade, comunicação e reversão.',
 'Avaliação: métricas, fonte, linha de base, alvo, janela, amostra, segmentações, responsável e data de revisão.',
 'Rastreabilidade: capítulos, anexos, artefatos canônicos, decisões e migrações afetados.',
])

h('12.9 Roteiro de fundação — 0 a 6 meses após aprovação')
p('Todos os itens permanecem propostos até aprovação formal. A ordem abaixo representa dependência recomendada, não promessa de prazo.')
table(['ID','Iniciativa','Dependências','Saída verificável'],[
 ['F-01','Auditoria do estado real','acesso a código, ambientes e responsáveis','mapa de capacidades confirmado, dívidas e evidências'],
 ['F-02','Registro canônico e manifesto de compatibilidade','F-01; Anexos A–F','artefatos published, hashes e combinações testadas'],
 ['F-03','Contratos, estados e persistência','F-01 e F-02','schemas executáveis, migrações, invariantes e testes'],
 ['F-04','Identidade, autorização e isolamento institucional','F-01 e modelo do Anexo E','testes de acesso cruzado, papéis e auditoria'],
 ['F-05','Privacidade, segurança e retenção','F-01 e inventário de dados','matriz de tratamento, segredos, retenção e incidentes'],
 ['F-06','Base de acessibilidade','capítulo 11 e protótipos dos capítulos 7–8','fluxos críticos testados por teclado e tecnologias assistivas'],
 ['F-07','Observabilidade, idempotência e recuperação','F-03 a F-05','métricas, filas, tentativas, backup e testes de restauração'],
 ['F-08','Publicação automática da documentação','F-02 e fontes editoriais','livro compilado com versões, hashes e reproduções verificadas'],
],[0.55,2.0,1.5,2.45])

h('12.10 Roteiro de consolidação — 6 a 18 meses após aprovação')
table(['ID','Iniciativa','Dependências','Condição de avanço'],[
 ['C-01','Fluxo completo do estudante','F-03 a F-07','submissão, revisão, decisão e reescrita homologadas'],
 ['C-02','Fluxo completo do professor','F-03 a F-07','atividade, validação, avaliação e nota homologadas'],
 ['C-03','Cobertura inicial dos gêneros','F-02, C-01 e C-02','GSP, taxonomia, rubrica e fixtures por gênero'],
 ['C-04','Curadoria administrativa','F-02, F-04 e F-05','segregação de funções e publicação auditável'],
 ['C-05','Relatórios formativos e institucionais','C-01 a C-03; privacidade','finalidade, agregação, acessibilidade e proteção contra reidentificação'],
 ['C-06','Integrações institucionais prioritárias','F-04, F-05 e F-07','padrões, segurança, suporte, reversão e contrato de dados'],
 ['C-07','Piloto de efetividade pedagógica','C-01 a C-03','protocolo, amostra, métricas, vieses e critérios de interrupção'],
 ['C-08','Operação e suporte em escala','F-07 e pilotos concluídos','níveis de serviço, capacidade, custos e base de conhecimento'],
],[0.55,2.0,1.5,2.45])

h('12.11 Portfólio exploratório — sem compromisso de entrega')
table(['ID','Hipótese','Questões obrigatórias antes de planejar'],[
 ['E-01','Personalização formativa','quais dados, finalidade, explicabilidade, equidade, contestação e separação da nota'],
 ['E-02','Entrada e feedback multimodais','contratos de imagem/áudio, acessibilidade, direitos, privacidade e custo'],
 ['E-03','Escrita colaborativa','coautoria, permissões, concorrência, versões e avaliação individual'],
 ['E-04','Roteamento entre modelos','qualidade, compatibilidade, custo, latência, privacidade e reprodutibilidade'],
 ['E-05','Assistência longitudinal','retenção, perfilamento, autoria, benefício comprovado e direito de contestar'],
],[0.6,1.8,4.1])
p('Um item exploratório somente entra em descoberta mediante problema documentado e responsável. Não recebe versão-alvo até superar os critérios de pesquisa, risco e valor. Resultado negativo ou inconclusivo é registrado com a mesma integridade que resultado favorável.')

h('12.12 Dependências e caminho crítico')
p('F-01 antecede decisões de implementação. F-02 e F-03 formam a base contratual. F-04 e F-05 são bloqueadores de qualquer piloto com dados reais. F-06 e F-07 são bloqueadores de liberação geral. C-01 e C-02 dependem dessa fundação; C-03 antecede expansão de relatórios e a avaliação pedagógica. Integrações e escala não devem preceder isolamento, privacidade, observabilidade e recuperação.')
p('O quadro de dependências registra relações bloqueia, requer, recomenda ou substitui, com versão mínima. Trabalhos sem dependência entre si podem avançar em paralelo somente quando houver capacidade e responsáveis distintos.')

h('12.13 Pesquisa e experimentação')
p('Pesquisa acadêmica, descoberta de produto e experimento técnico são registrados separadamente. Todo estudo define hipótese, protocolo, população, amostra, critérios de inclusão e exclusão, acessibilidade, linha de base, métricas, análise, vieses, critérios de interrupção e plano de dados. Quando aplicável, obtém avaliação ética institucional antes da coleta.')
p('Textos reais não são usados fora da finalidade autorizada. Consentimento não é presumido como única base; privacidade, minimização, anonimização, retenção e acesso seguem o capítulo 10. Resultados negativos, danos e limitações são registrados. A promoção de uma hipótese a iniciativa exige revisão humana e não decorre automaticamente de uma métrica.')
p('Indicadores pedagógicos abrangem interpretação dos achados, decisões autorais, reescrita, autorregulação, evidências positivas e evolução entre versões. Nota, quantidade de erros, correções ou cliques não medem aprendizagem isoladamente. Análises devem considerar gênero, variedade linguística, falsos positivos, falsos negativos e diferenças entre grupos.')

h('12.14 Gestão de riscos')
p('O registro consolidado cobre riscos pedagógicos, linguísticos, acessibilidade, privacidade, segurança, operação, finanças, fornecedor, reputação e documentação. Cada risco possui probabilidade, impacto, exposição, proprietário, mitigação, contingência, indicador, revisão e risco residual. Riscos críticos bloqueiam aprovação ou liberação.')
p('Indícios de integridade acadêmica, mudanças na nota, personalização e uso de novos modelos exigem análise humana reforçada. A equipe registra também o risco de não executar uma iniciativa, evitando que apenas riscos da mudança sejam considerados.')

h('12.15 Métricas, aceitação e avaliação posterior')
p('Critérios de aceitação verificam a entrega; métricas observam comportamento e resultado; critérios de sucesso avaliam benefício após uma janela. Cada iniciativa define fonte, linha de base, alvo, unidade, período, amostra, segmentações, responsável e limites de interrupção. Conclusão técnica não equivale a sucesso pedagógico.')
p('Após piloto ou liberação, uma revisão em data definida decide expandir, ajustar, reter, reverter ou descontinuar. Benefício insuficiente, regressão de acessibilidade, viés, aumento de risco, custo insustentável ou perda de confiabilidade podem interromper a expansão.')

h('12.16 Migração e compatibilidade')
p('Cada migração identifica objetos, versões de origem e destino, compatibilidade, transformações, validações, lotes, idempotência, backup, ponto de restauração e plano de reversão. Simulação antecede dados reais; falha parcial não deixa registros utilizáveis em estado intermediário.')
p('Prompts, contratos, taxonomias, rubricas, mapeamentos, estados e configurações mantêm matriz de compatibilidade. A convivência entre versões tem prazo e critérios de desligamento. A migração atualiza testes, documentação, exemplos, base de conhecimento e comunicação dos públicos afetados.')

h('12.17 Descontinuação')
p('A retirada passa por proposta, análise, aprovação, aviso, congelamento, migração, retirada e arquivamento. A decisão registra motivo, usuários afetados, alternativa, prazo, exportação acessível, retenção, remoção de permissões, integrações e segredos, atualização documental e plano de reversão.')
p('Histórico necessário para interpretar atividades, revisões, avaliações e notas anteriores é preservado. A descontinuação não apaga silenciosamente dados nem converte versões históricas. O impacto é medido após a retirada.')

h('12.18 Governança do roteiro')
p('Produto coordena o portfólio; pedagogia e linguística avaliam valor formativo; tecnologia estima arquitetura e esforço; segurança, privacidade e acessibilidade exercem poder de bloqueio em seus requisitos críticos; editorial mantém rastreabilidade com o livro. Finanças e operação validam sustentabilidade.')
p('O roteiro é revisado trimestralmente e após incidente grave, mudança institucional ou alteração significativa de fornecedor. Cada revisão publica versão, data, decisões, movimentos, cancelamentos e justificativas. Alterações não reescrevem o histórico.')

h('12.19 Rastreabilidade editorial')
table(['Tipo de mudança','Documentos a revisar'],[
 ['Princípio pedagógico','capítulos 1, 2, 5, 7 e 8'],['Arquitetura ou domínio','capítulos 3, 4, 9 e Anexo E'],['Prompt ou contrato','capítulo 9 e Anexos A, B ou F'],['Gênero ou taxonomia','capítulo 6 e Anexo C'],['Rubrica ou cálculo','capítulos 8 e 9; Anexo D'],['Gestão, segurança ou privacidade','capítulo 10'],['Interface ou identidade','capítulos 7, 8 e 11'],['Capacidade disponível ou retirada','capítulos afetados, capítulo 12, base de conhecimento e Apêndice'],
],[2.0,4.5])
p('Toda iniciativa mantém uma matriz de impacto documental. Mudança canônica exige nova versão e hash. Exemplos incompatíveis são atualizados ou marcados como hipotéticos e não executáveis. O livro compilado distingue graficamente capacidade disponível, piloto e hipótese.')

h('12.20 Artefatos obrigatórios')
bullets([
 'mapa versionado de capacidades documentadas, confirmadas e planejadas;',
 'registro estruturado dos itens do roteiro;',
 'quadro de dependências e caminho crítico;',
 'registro consolidado de riscos e benefícios;',
 'registro de decisões arquiteturais, pedagógicas, linguísticas e editoriais;',
 'planos e resultados de experimentos;',
 'matriz de métricas, critérios de aceitação e revisões pós-lançamento;',
 'políticas distintas de versionamento, migração e descontinuação;',
 'calendário de decisões e revisões;',
 'matriz de rastreabilidade entre iniciativas, versões, capítulos e anexos.',
])

h('12.21 Critérios de qualidade do roteiro')
bullets([
 'Nenhuma capacidade é declarada disponível sem ambiente, versão, data e evidência.',
 'Todo item possui problema, público, resultados, responsável, riscos, dependências e critérios.',
 'Horizonte, maturidade e estado permanecem campos distintos.',
 'Bloqueadores de segurança, privacidade, acessibilidade e integridade da nota não são rebaixados por pontuação.',
 'Hipóteses exploratórias não recebem promessa pública de entrega.',
 'Migrações e descontinuações preservam dados, histórico e comunicação.',
 'Resultados negativos e limitações são registrados.',
 'Mudanças atualizam os capítulos e anexos afetados com versão e hash.',
])

doc.save(OUT)
print(OUT)
