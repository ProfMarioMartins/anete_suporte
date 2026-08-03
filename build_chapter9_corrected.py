from docx import Document
from docx.shared import Pt
import json

BASE = "/Users/mariomartins2/Library/CloudStorage/GoogleDrive-mario.martins@ufersa.edu.br/Meu Drive/Disciplinas_em_R/ANETE"
SRC = "/Users/mariomartins2/Library/CloudStorage/GoogleDrive-mario.martins@ufersa.edu.br/Meu Drive/TEXTOPIA/Documentação técnico-pedagógica/9. Especificação técnica dos módulos de revisão e avaliação do Textopia_.docx"
CHAPTER = BASE + "/9. Especificação técnica dos módulos de revisão e avaliação do Textopia_ - versão corrigida.docx"
ANNEX = BASE + "/Anexo técnico F - Contratos canônicos dos módulos de revisão, validação, avaliação e nota do Textopia_.docx"

def clear(doc):
    body=doc._element.body; sect=body.sectPr
    for x in list(body):
        if x is not sect: body.remove(x)

def setup(doc):
    for s in doc.styles:
        if hasattr(s,'font'):
            s.font.name='Arial'
    doc.styles['Normal'].font.size=Pt(11)

def add_helpers(doc):
    def p(text='', style=None): return doc.add_paragraph(text, style=style)
    def h(text, level=1): return doc.add_heading(text, level=level)
    def bullets(items):
        for item in items:
            q=p(style='List Bullet'); q.add_run(item)
    return p,h,bullets

doc=Document(SRC); clear(doc); setup(doc); p,h,bullets=add_helpers(doc)
title=p('9. Especificação técnica dos módulos de revisão, validação, avaliação e nota do Textopia_')
title.style=doc.styles['Title']

h('9.1 Escopo e responsabilidades')
p('Este capítulo define a integração entre revisão pedagógica automatizada, validação contratual e humana, avaliação rubricada, cálculo determinístico, decisão docente e publicação da nota. As operações possuem contratos, execuções e máquinas de estados independentes, relacionados por identificadores imutáveis. O estudante permanece autor do texto; o modelo produz saídas delimitadas; o backend valida, persiste e calcula; e o professor toma as decisões pedagógicas que lhe competem.')
p('As fontes normativas são: Anexo A, prompt de revisão; Anexo B, prompt de avaliação; Anexo C, gêneros e taxonomias; Anexo D, rubricas, mapeamento e cálculo; Anexo E, domínio, persistência, estados, achados globais e segmentação; e Anexo F, contratos dos módulos. Este capítulo não mantém cópias manuais desses artefatos. As reproduções destinadas ao livro são geradas na publicação mediante versão e SHA-256.')

h('9.2 Componentes e fronteiras')
bullets([
 'API de atividades: resolve a ActivityPublication imutável e sua configuração pedagógica.',
 'Serviço de segmentação: normaliza o texto em NFC, calcula textHash e produz paragraphId, sentenceId, offsets e conteúdo literal.',
 'Registro canônico: fornece prompts, contratos, gêneros, taxonomias, mapeamentos e rubricas publicados, versionados e imutáveis.',
 'Orquestrador de revisão: compõe o contexto, cria execução idempotente e coordena modelo, validação e persistência.',
 'Adaptador de modelo: envia instruções do sistema separadas dos dados não confiáveis e solicita saída estruturada.',
 'Validador contratual e semântico: verifica schema, versões, hashes, códigos, IDs, trechos, limites e invariantes.',
 'Serviço de validação: registra decisões automáticas e humanas sem sobrescrever a saída original.',
 'Serviço de avaliação: compõe o pacote rubricado a partir do texto completo, rubrica e achados em estados autorizados.',
 'Serviço de cálculo: converte níveis em pontuações e aplica pesos e arredondamento de modo determinístico.',
 'Serviço de decisão e publicação: registra confirmação, modificação ou rejeição docente e libera somente uma nota válida.',
 'Persistência e auditoria: mantém registros append-only, idempotência, correlação, proveniência e isolamento institucional.',
])

h('9.3 Resolução do contexto autorizado')
p('Antes de chamar qualquer modelo, o backend valida autenticação, autorização, institutionId, activityPublicationId, textVersionId, tamanho e tipo do conteúdo. A publicação da atividade determina gênero, objetivos, instruções, política de revisão, rubrica e critérios habilitados. O conjunto efetivo de critérios é calculado conforme o Anexo E: (mapeados para o gênero ∩ habilitados) − desabilitados.')
p('O backend recupera somente artefatos com estado published e combinação declarada como compatível. Ausência, hash divergente, código desconhecido, conjunto efetivo vazio ou conflito de precedência bloqueia a chamada. O sistema registra versões e hashes da atividade, texto, segmentação, prompts, contratos, catálogo de gêneros, taxonomia, mapeamentos, rubrica e modelo.')
p('Texto do estudante, instruções docentes e conteúdos recuperados são delimitados como dados não confiáveis. Eles não podem substituir instruções do sistema, ampliar a taxonomia, alterar schemas, solicitar segredos ou introduzir propriedades executáveis.')

h('9.4 Segmentação e integridade textual')
p('A revisão consome o contrato canônico de segmentação do Anexo E. O texto é normalizado em NFC; offsets usam intervalos semiabertos sobre pontos de código Unicode; e conteudoLiteral corresponde exatamente ao intervalo indicado. Cada nova TextVersion é segmentada novamente, sem transportar sentenceIds por similaridade.')
p('Divergência entre textVersionId, textHash, offsets, sentenceId e conteúdo literal interrompe a execução antes do modelo ou invalida sua saída. Intervenções citam exatamente uma sentença. Fenômenos apoiados por mais de uma sentença pertencem a achadosGlobais.')

h('9.5 Composição do prompt de revisão')
p('O orquestrador preenche exclusivamente as variáveis declaradas no Anexo A. A composição inclui identificadores, versões, gênero, objetivos, instruções, regras docentes autorizadas, escopo, limiar de confiança, taxonomia efetiva, mapa de sentenças e texto. Variáveis ausentes ou incompatíveis impedem a chamada.')
p('Limites de tamanho, quantidade de intervenções, agrupamento de recorrências e uso opcional de exemplos são obtidos da ActivityPublication. Exemplos nunca são obrigatórios. A seleção de contexto deve ser mínima, mas não pode omitir informações necessárias à interpretação do gênero, da atividade ou da rubrica.')

h('9.6 Fluxo da revisão')
bullets([
 'Validar requisição, permissão, idempotencyKey e máquina de estados.',
 'Resolver a publicação da atividade e a combinação compatível de artefatos.',
 'Segmentar ou recuperar segmentação verificada da TextVersion.',
 'Criar RevisionExecution imutável, com correlationId, versões, hashes e estado inicial.',
 'Compor o prompt e solicitar saída estruturada segundo o Anexo F.',
 'Validar JSON, schema e invariantes semânticos.',
 'Aplicar limiar de confiança e política explícita de disponibilização, retenção ou abstenção.',
 'Persistir saída original, eventos e limitações; nunca sobrescrever resultado anterior.',
 'Encaminhar a execução válida à máquina de validação do Anexo E.',
])
p('Falhas transitórias podem ser repetidas conforme política versionada, com espera progressiva, limite temporal, idempotência e registro de cada tentativa. Timeout, recusa, truncamento, indisponibilidade, saída inválida e falha de persistência possuem códigos distintos. Uma nova tentativa da mesma operação não cria nova versão textual nem duplica decisões.')

h('9.7 Unidades e regras semânticas da revisão')
bullets([
 'Intervenção sentencial: exatamente um sentenceId, trecho literal e diagnóstico observável.',
 'Achado global estruturado: um ou mais sentenceIds que sustentam fenômeno suprassentencial.',
 'Feedback global: síntese de pontos fortes, prioridades, comentário e limitações; não recebe diagnósticos suprassentenciais.',
 'Evidência, diagnóstico e orientação permanecem distinguíveis; orientação é coerente, executável e sustentada pelo texto e pela atividade.',
 'IDs são únicos em toda a saída; intervenções seguem a ordem textual e referências pertencem à execução.',
 'Prioridade usa baixa, media ou alta; confiança varia de 0 a 1 e não é juízo sobre o estudante.',
 'A saída distingue revisão concluída sem achados de abstenção, retenção, análise incompleta e falha.',
 'Nota, conceito, nível da rubrica e pontuação são proibidos na revisão.',
])

h('9.8 Contrato da revisão e exemplo editorial')
p('O schema normativo é SaidaRevisaoTextopia 2.0.0, definido no Anexo F. Ele inclui metadados de proveniência, status estruturado, limiar e política aplicados, intervencoes, achadosGlobais e feedbackGlobal. Restrições que JSON Schema não expressa são invariantes semânticos obrigatórios do mesmo anexo.')
p('Exemplos exibidos neste capítulo são gerados a partir dos fixtures válidos do Anexo F. O fixture de referência usa promptVersion revisao-1.1.0, códigos semânticos publicados, activityPublicationId, mappingVersion, genreCatalogVersion, segmentationVersion e textHash. Exemplos hipotéticos não executáveis devem ser marcados explicitamente como tais.')

h('9.9 Erros e observabilidade segura')
p('O contrato ErroModuloTextopia 2.0.0, no Anexo F, diferencia erro de entrada, autorização, resolução de versão, segmentação, provedor, contrato, semântica, persistência e concorrência. Ele inclui errorId, correlationId, stage, code, retryable, occurredAt e mensagem segura. executionId é opcional quando a falha antecede a execução.')
p('Logs comuns não armazenam texto integral nem saída bruta. Artefatos indispensáveis à investigação permanecem em repositório restrito, criptografado, auditado e submetido a retenção definida. Mensagens destinadas ao usuário são separadas de detalhes técnicos.')

h('9.10 Validação automática e humana')
p('A validação automática verifica contrato, versões, hashes, códigos, referências, conteúdo literal, limites, ordenação e política de confiança. A validação humana ocorre somente quando configurada ou exigida pela política. Cada decisão atua sobre uma unidade específica e preserva saída original, autor, data, justificativa, estado anterior e estado posterior.')
p('O pacote RevisaoValidadaTextopia 2.0.0, no Anexo F, possui validationId, activityPublicationId, revisionExecutionId, política aplicada e decisões sobre intervenções, achados globais e feedback global. Estados globais são derivados dos estados individuais; pendente não exige validatedAt, enquanto estados finais exigem ator e data.')
p('Editar significa produzir orientação validada separada; não altera silenciosamente evidência, trecho, critério, categoria ou diagnóstico. Solicitar nova análise cria nova execução relacionada. Contestação do estudante pertence a registro próprio e não modifica retroativamente a validação original.')

h('9.11 Preparação da avaliação rubricada')
p('A avaliação usa a mesma TextVersion e ActivityPublication, o texto completo, a rubrica publicada, o mapeamento canônico do Anexo D, evidências positivas, recorrências, achados em estados autorizados e limitações. Achados rejeitados, invalidados, pendentes ou abaixo do limiar aplicável não influenciam a avaliação.')
p('A revisão é fonte de evidências, não tabela de descontos. A quantidade de problemas não é convertida mecanicamente em perda de pontos. Quando a revisão se abstém ou está limitada, essa condição é propagada; o modelo não completa lacunas por inferência improvisada.')

h('9.12 Fluxo da avaliação')
bullets([
 'Validar estado, permissão, idempotência e compatibilidade entre texto, atividade, revisão, rubrica e mapeamento.',
 'Criar GradingExecution imutável e registrar versões, hashes e estados de achados autorizados.',
 'Compor o prompt do Anexo B com dados delimitados e rubrica autorizada.',
 'Solicitar níveis, evidências, justificativas, confiança e limitações, sem cálculo de nota.',
 'Validar schema, cobertura e unicidade dos critérios e pertencimento de cada levelId.',
 'Aplicar política de abstenção por critério ou avaliação indisponível.',
 'Persistir a proposta e encaminhá-la ao cálculo somente quando válida.',
])

h('9.13 Contrato da avaliação rubricada')
p('O schema AvaliacaoRubricadaTextopia 2.0.0 é definido no Anexo F. Ele registra activityPublicationId, textHash, versões e hashes dos artefatos, estado da execução, critérios únicos, evidências referenciáveis, recorrência, confiança, limitações e abstenção por critério. Regras condicionais exigem limitações quando a avaliação está indisponível e cobertura integral dos critérios aplicáveis quando há proposta válida.')

h('9.14 Cálculo determinístico')
p('Após validar a avaliação, o backend consulta exclusivamente a rubrica e a fórmula publicadas no Anexo D. Para cada critério aplicável, recupera levelId, pontos, peso e máximo; exclui critérios somente quando allowNotApplicable autoriza; calcula sem arredondamento intermediário; e aplica o arredondamento apenas ao resultado final.')
p('O registro do cálculo contém calculationId, gradingExecutionId, rubricVersion, mappingVersion, inputHash, nota bruta, resultado proposto, regra de arredondamento e estado. Repetir o cálculo com o mesmo inputHash produz o mesmo resultado. Nova avaliação ou rubrica gera cálculo distinto e não sobrescreve o anterior.')

h('9.15 Decisão docente e nota publicada')
p('O professor confirma, modifica ou rejeita o resultado calculado. A decisão não altera o cálculo original e registra decisionId, calculationId, ator, data, tipo, valor anterior, valor decidido e justificativa. Modificação ou rejeição exige justificativa; confirmação preserva a identidade entre resultado calculado e decidido.')
p('A nota publicada é um registro separado, liberado somente após decisão docente válida. Ela contém publishedGradeId, decisionId, activityPublicationId, textVersionId, valor, escala, data, responsável e estado. Reabertura, retificação e anulação seguem a máquina de estados do Anexo E, produzem novos eventos e nunca apagam o histórico.')

h('9.16 Integridade acadêmica, diversidade e agência')
p('Indícios de integridade acadêmica exigem linguagem cautelosa, acesso restrito e decisão humana; não constituem acusação, prova ou desconto automático. Diferenças de variedade linguística não são tratadas automaticamente como deficiência. Prescrições dependem de gênero, situação comunicativa, objetivos e interlocutores.')
p('O estudante pode aceitar, adaptar, rejeitar ou contestar orientações. Essas decisões não alteram automaticamente a avaliação ou a nota. Uma correção somente existe quando o autor incorpora uma alteração em nova TextVersion.')

h('9.17 Segurança e privacidade')
bullets([
 'Classificar e minimizar dados por finalidade, sem omitir contexto indispensável à análise.',
 'Separar instruções do sistema, configuração autorizada e dados não confiáveis.',
 'Impedir que conteúdo recuperado amplie permissões, taxonomias ou schemas.',
 'Aplicar isolamento por institutionId e autorização em todas as consultas e relacionamentos.',
 'Criptografar artefatos restritos, limitar acesso e aplicar retenção verificável.',
 'Mascarar telemetria e separar mensagens públicas de detalhes técnicos.',
])

h('9.18 Idempotência, concorrência e transações')
p('Operações de publicação, revisão, validação, avaliação, cálculo, decisão e nota exigem idempotencyKey quando capazes de duplicar efeitos. Restrições de unicidade e controle otimista impedem execuções ou decisões concorrentes incompatíveis. Persistência do resultado e transição de estado ocorrem atomicamente; falha parcial não publica artefato incompleto.')

h('9.19 Testes e critérios operacionais')
bullets([
 'Testes de contrato para payloads válidos, inválidos e todas as combinações de versões admitidas.',
 'Testes de propriedade para normalização, offsets, hashes e conteúdo literal.',
 'Testes semânticos para códigos, IDs, ordenação, limites, abstenção e estados condicionais.',
 'Testes de idempotência, concorrência, repetição, timeout e recuperação de falhas.',
 'Testes de segurança contra injeção direta e indireta, vazamento entre instituições e conteúdo executável.',
 'Testes pedagógico-linguísticos por gênero, variedade, critério e grupo, incluindo falsos positivos e linguagem acusatória.',
 'Testes de cálculo, reabertura, retificação, anulação e publicação da nota.',
])
p('A observabilidade mede latência, custo, tokens, falhas, tentativas, abstenções, retenções, divergências, alterações docentes e vieses. Limites operacionais e orçamentos são configurados por política; regressões acima dos limites bloqueiam a publicação de nova combinação de artefatos.')

h('9.20 Versionamento, compatibilidade e publicação')
p('Cada artefato possui versão semântica independente. Mudança incompatível exige nova versão principal; mudança compatível de capacidade exige versão secundária; correção sem alteração de contrato exige versão de correção. O registro de compatibilidade declara combinações aceitas e seus hashes. Migradores são explícitos, testados e nunca misturam objetos sem identificação.')
p('Na publicação do livro, o processo extrai dos Anexos A, B e F os artefatos marcados para reprodução, normaliza LF, calcula SHA-256 e compara versão, hash e status. A divergência interrompe a publicação. O capítulo-fonte contém apenas marcadores editoriais; o artefato compilado recebe reproduções identificadas como geradas e verificadas.')

h('9.21 Critérios de aceitação da implementação')
bullets([
 'Nenhuma chamada ocorre com artefatos ausentes, incompatíveis ou sem hash verificado.',
 'Toda saída é vinculada à ActivityPublication e TextVersion exatas.',
 'Saída original, validação, avaliação, cálculo, decisão e nota permanecem imutáveis e auditáveis.',
 'Estados e transições correspondem às máquinas canônicas do Anexo E.',
 'Abstenção, retenção, falha e ausência de achados são estados distinguíveis.',
 'Achados não autorizados ou invalidados nunca influenciam a avaliação.',
 'O modelo não calcula nota e nenhuma nota é publicada sem decisão docente válida.',
 'Operações repetidas ou concorrentes não produzem duplicidades.',
 'Os fixtures dos anexos validam contra os schemas publicados.',
])

doc.save(CHAPTER)

# Anexo F: fonte normativa única dos contratos do módulo.
a=Document(); setup(a); p,h,bullets=add_helpers(a)
t=p('Anexo técnico F — Contratos canônicos dos módulos de revisão, validação, avaliação e nota do Textopia_')
t.style=a.styles['Title']
p('Versão normativa: contratos-modulos-2.0.0  |  Estado: published  |  Fonte canônica')
h('1. Finalidade')
p('Este anexo é a fonte normativa dos schemas de troca dos módulos. O capítulo 9 contém explicações e reproduções geradas, nunca cópias editadas manualmente. Todos os registros exigem version, identificador estável e createdAt quando constituem persistência.')
h('2. Regras comuns')
bullets(['additionalProperties é false em todos os objetos normativos.','IDs são strings não vazias e únicos no escopo indicado.','Datas usam RFC 3339 em UTC.','Hashes usam o prefixo sha256:.','Campos condicionais e invariantes semânticos são obrigatórios mesmo quando não exprimíveis apenas por JSON Schema.','activityPublicationId, textVersionId e textHash devem referir-se ao mesmo texto e à mesma configuração.'])

schemas={
'SaidaRevisaoTextopia 2.0.0': {
 '$id':'https://textopia.local/schemas/revisao/2.0.0','required':['version','revisionExecutionId','activityId','activityPublicationId','textVersionId','textHash','segmentationVersion','promptVersion','contractVersion','modelVersion','genreCatalogVersion','taxonomyVersion','mappingVersion','confidenceThreshold','deliveryPolicy','status','createdAt','intervencoes','achadosGlobais','feedbackGlobal'],
 'status':['concluida','sem_achados','retida','abstencao','incompleta'],
 'intervencaoRequired':['id','sentenceId','paragraphId','inicio','fim','trecho','criterioCodigo','categoriaCodigo','diagnostico','orientacao','prioridade','confianca'],
 'achadoGlobalRef':'https://textopia.local/schemas/achado-global/1.0.0',
 'feedbackGlobalRequired':['pontosFortes','prioridadeIds','comentarioFinal','limitacoes'],
 'invariantes':['trecho == conteudoLiteral da segmentação','uma intervenção referencia exatamente uma sentença','IDs únicos entre todas as unidades','retida, abstencao e incompleta exigem limitacoes estruturadas','quantidades respeitam ActivityPublication']},
'ErroModuloTextopia 2.0.0': {
 '$id':'https://textopia.local/schemas/erro-modulo/2.0.0','required':['version','errorId','correlationId','stage','code','retryable','occurredAt','userMessage'],'optional':['executionId','retryAfterSeconds','safeDetails'],
 'stage':['request','authorization','resolution','segmentation','provider','contract','semantic','persistence','concurrency']},
'RevisaoValidadaTextopia 2.0.0': {
 '$id':'https://textopia.local/schemas/revisao-validada/2.0.0','required':['version','validationId','activityPublicationId','textVersionId','revisionExecutionId','revisionOutputRef','policyVersion','status','createdAt','interventionDecisions','globalFindingDecisions','feedbackDecision'],
 'status':['pendente','parcialmente_validada','validada','rejeitada','retida'],
 'decisionRequired':['decisionId','targetId','status','decidedBy','decidedAt','justification'],
 'invariantes':['pendente não exige validatedAt','estados finais exigem validatedBy e validatedAt','estado global é derivado das decisões','edição preserva original e registra apenas orientação validada']},
'AvaliacaoRubricadaTextopia 2.0.0': {
 '$id':'https://textopia.local/schemas/avaliacao-rubricada/2.0.0','required':['version','gradingExecutionId','activityId','activityPublicationId','textVersionId','textHash','revisionExecutionId','validatedRevisionRef','gradingPromptVersion','gradingContractVersion','revisionPromptVersion','revisionContractVersion','rubricVersion','mappingVersion','modelVersion','authorizedFindingStates','status','createdAt','criterios','justificativaGlobal','limitacoes'],
 'status':['em_execucao','proposta_valida','parcialmente_abstida','avaliacao_indisponivel','falha'],
 'criterionRequired':['criterioId','status','nivelId','evidenciasPositivas','intervencoesRelacionadas','achadosGlobaisRelacionados','recorrencias','justificativa','confianca','limitacoes'],
 'criterionStatus':['avaliado','nao_aplicavel','abstencao'],
 'invariantes':['criterioId único e cobertura integral da rubrica','nivelId pertence ao criterioId quando avaliado','somente achados em authorizedFindingStates','indisponível e abstenção exigem limitações']},
'DecisaoDocenteNota 1.0.0': {
 '$id':'https://textopia.local/schemas/decisao-docente-nota/1.0.0','required':['version','decisionId','calculationId','activityPublicationId','textVersionId','decisionType','calculatedScore','decidedBy','decidedAt','justification'],
 'decisionType':['confirmar','modificar','rejeitar'],'optional':['decidedScore'],
 'invariantes':['confirmar exige decidedScore igual a calculatedScore','modificar exige decidedScore válido','rejeitar não publica nota','modificar e rejeitar exigem justificativa não vazia']},
'NotaPublicadaTextopia 1.0.0': {
 '$id':'https://textopia.local/schemas/nota-publicada/1.0.0','required':['version','publishedGradeId','decisionId','activityPublicationId','textVersionId','score','scaleMinimum','scaleMaximum','state','publishedBy','publishedAt','createdAt'],
 'state':['publicada','retificada','anulada'],'optional':['supersedesPublishedGradeId','reason'],
 'invariantes':['somente decisão confirmar ou modificar produz nota','retificação referencia nota anterior','anulação preserva histórico']}
}

for i,(name,schema) in enumerate(schemas.items(),3):
    h(f'{i}. {name}')
    q=p(json.dumps(schema,ensure_ascii=False,indent=2))
    q.style=a.styles['Normal']
    for run in q.runs: run.font.name='Courier New'; run.font.size=Pt(8)

h('9. Fixture válido de revisão')
fixture={'version':'2.0.0','revisionExecutionId':'rev-1042','activityId':'ATV-2026-001','activityPublicationId':'ATV-2026-001-p1','textVersionId':'texto-7-v3','textHash':'sha256:fixture','segmentationVersion':'1.0.0','promptVersion':'revisao-1.1.0','contractVersion':'2.0.0','modelVersion':'modelo-2026-07','genreCatalogVersion':'1.0.0','taxonomyVersion':'1.0.0','mappingVersion':'1.0.0','confidenceThreshold':0.8,'deliveryPolicy':'validacao_humana','status':'concluida','createdAt':'2026-08-02T12:00:00Z','intervencoes':[{'id':'i-1','sentenceId':'s-12','paragraphId':'p-3','inicio':91,'fim':126,'trecho':'Os resultados demonstra a hipótese.','criterioCodigo':'CONCORDANCIA','categoriaCodigo':'DESVIO_CONCORDANCIA','diagnostico':'O verbo não concorda com o núcleo plural do sujeito.','orientacao':'Revise a forma verbal para que concorde com “resultados”.','prioridade':'alta','confianca':0.98}],'achadosGlobais':[],'feedbackGlobal':{'pontosFortes':['O objetivo está explicitado.'],'prioridadeIds':['i-1'],'comentarioFinal':'Examine a concordância identificada e decida como reescrever.','limitacoes':[]}}
q=p(json.dumps(fixture,ensure_ascii=False,indent=2));
for run in q.runs: run.font.name='Courier New'; run.font.size=Pt(8)
p('O valor sha256:fixture é marcador de fixture e deve ser substituído pelo hash calculado no teste executável. Os códigos devem existir na versão publicada da taxonomia usada pelo ambiente de teste; caso contrário, o fixture é recusado.')
h('10. Compatibilidade e publicação')
p('O manifesto de compatibilidade relaciona cada versão deste anexo às versões aceitas dos Anexos A a E. O processo editorial calcula SHA-256 do conteúdo normativo normalizado em LF e só reproduz um artefato no capítulo 9 quando versão, hash, estado published e testes dos fixtures forem válidos.')
a.save(ANNEX)
print(CHAPTER); print(ANNEX)
