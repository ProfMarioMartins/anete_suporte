#!/usr/bin/env python3
"""Gera a cópia editorial do capítulo 9 a partir dos Anexos A e B.

O capítulo-fonte não contém cópias manuais dos prompts. Este programa valida
os SHA-256 declarados no próprio capítulo e insere reproduções somente no
artefato compilado de publicação.
"""
from copy import deepcopy
from hashlib import sha256
from pathlib import Path
from docx import Document
import argparse, re

REV_HEAD='PROMPT CANÔNICO DE REVISÃO PEDAGÓGICA DO TEXTOPIA_'
GRADE_HEAD='PROMPT CANÔNICO DE AVALIAÇÃO RUBRICADA E PROPOSTA DE NOTA DO TEXTOPIA_'

def canonical_paragraphs(path, heading):
    doc=Document(path); paragraphs=doc.paragraphs
    start=next(i for i,p in enumerate(paragraphs) if p.text==heading)
    selected=paragraphs[start:]
    text='\n'.join(p.text for p in selected).strip().replace('\r\n','\n')
    return selected, sha256(text.encode('utf-8')).hexdigest()

def declared_hash(doc, section):
    heading=next(i for i,p in enumerate(doc.paragraphs) if p.text.startswith(section))
    text=doc.paragraphs[heading+1].text
    match=re.search(r'([0-9a-f]{64})',text)
    if not match: raise ValueError(f'SHA-256 não declarado após {section}')
    return match.group(1), doc.paragraphs[heading+1]

def insert_before(anchor, paragraphs):
    for p in paragraphs: anchor._p.addprevious(deepcopy(p._p))

def build(chapter, annex_a, annex_b, output):
    doc=Document(chapter)
    rev_paras, rev_hash=canonical_paragraphs(annex_a,REV_HEAD)
    grade_paras, grade_hash=canonical_paragraphs(annex_b,GRADE_HEAD)
    expected_rev, rev_note=declared_hash(doc,'9.7 Prompt canônico')
    expected_grade, grade_note=declared_hash(doc,'9.15 Prompt canônico')
    if rev_hash!=expected_rev: raise ValueError(f'Hash do Anexo A divergente: esperado {expected_rev}, recebido {rev_hash}')
    if grade_hash!=expected_grade: raise ValueError(f'Hash do Anexo B divergente: esperado {expected_grade}, recebido {grade_hash}')
    h98=next(p for p in doc.paragraphs if p.text.startswith('9.8 Exemplo válido'))
    h916=next(p for p in doc.paragraphs if p.text.startswith('9.16 Cálculo determinístico'))
    insert_before(h98,rev_paras)
    insert_before(h916,grade_paras)
    rev_note.text+=' [REPRODUÇÃO GERADA E VERIFICADA]'
    grade_note.text+=' [REPRODUÇÃO GERADA E VERIFICADA]'
    doc.save(output)

def main():
    ap=argparse.ArgumentParser()
    ap.add_argument('--chapter',required=True,type=Path)
    ap.add_argument('--annex-a',required=True,type=Path)
    ap.add_argument('--annex-b',required=True,type=Path)
    ap.add_argument('--output',required=True,type=Path)
    a=ap.parse_args(); build(a.chapter,a.annex_a,a.annex_b,a.output)

if __name__=='__main__': main()
