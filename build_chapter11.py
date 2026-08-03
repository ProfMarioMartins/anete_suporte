from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SRC='/Users/mariomartins2/Downloads/11. Identidade visual e diretrizes de marca do TEXTOPIA_.docx'
OUT='/Users/mariomartins2/Library/CloudStorage/GoogleDrive-mario.martins@ufersa.edu.br/Meu Drive/Disciplinas_em_R/ANETE/11. Identidade visual e diretrizes de marca do Textopia_ - versão corrigida.docx'
doc=Document(SRC)
body=doc._element.body; sect=body.sectPr
for x in list(body):
    if x is not sect: body.remove(x)
for section in doc.sections:
    section.top_margin=Inches(1); section.bottom_margin=Inches(1)
    section.left_margin=Inches(1); section.right_margin=Inches(1)

def p(text='',style=None): return doc.add_paragraph(text,style=style)
def h(text,level=2): return doc.add_heading(text,level=level)
def bullets(items):
    for item in items:
        q=p(); q.paragraph_format.left_indent=Pt(18); q.paragraph_format.first_line_indent=Pt(-12)
        q.add_run('• '); q.add_run(item)
def shade(cell,fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill.replace('#','')); tcPr.append(shd)
def table(headers,rows,widths=None):
    t=doc.add_table(rows=1,cols=len(headers)); t.autofit=False
    tblPr=t._tbl.tblPr; borders=OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        e=OxmlElement('w:'+edge); e.set(qn('w:val'),'single'); e.set(qn('w:sz'),'4'); e.set(qn('w:color'),'B8C2CC'); borders.append(e)
    tblPr.append(borders)
    for i,x in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=x; shade(c,'E8ECEF')
        c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.LEFT
        for run in c.paragraphs[0].runs: run.bold=True
    for row in rows:
        cells=t.add_row().cells
        for i,x in enumerate(row):
            cells[i].text=str(x); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[i].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.LEFT
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    return t

title=p('11. Identidade visual e diretrizes de marca do Textopia_'); title.style=doc.styles['Heading 1']
p('Manual de marca e design system — versão 1.0')
pic=p(); pic.alignment=WD_ALIGN_PARAGRAPH.CENTER
pic.add_run().add_picture('/tmp/textopia-brand/image1.png',width=Inches(5.4))

h('11.1 Papel e posicionamento da identidade')
p('A identidade do Textopia_ expressa uma plataforma acadêmica que combina rigor, clareza e acolhimento para apoiar processos de escrita sem substituir a autoria humana. A marca deve parecer formativa, confiável, contemporânea e sóbria; não deve parecer punitiva, infantilizada, burocrática ou tecnologicamente intimidante.')
table(['Princípio','Expressão visual','Evitar'],[
 ['Clareza','hierarquia nítida, textos legíveis, poucos elementos simultâneos','ornamento sem função e densidade excessiva'],
 ['Acolhimento','espaço respirável, linguagem respeitosa e cantos moderados','tom acusatório, cores agressivas e infantilização'],
 ['Rigor acadêmico','alinhamento, consistência, evidências e estados explícitos','aparência improvisada ou ambígua'],
 ['Agência humana','ações reversíveis quando possível e decisões claramente atribuídas','automação que pareça decidir pelo estudante ou professor'],
],[1.25,2.65,2.6])
p('Públicos principais: estudantes, professores, equipes pedagógicas, administradores institucionais e equipes técnicas. A identidade institucional regula reconhecimento e comunicação; o design system regula interfaces; o projeto editorial regula o livro e documentos. Os três sistemas compartilham tokens e princípios, mas possuem regras próprias.')

h('11.2 Nome, grafia e voz')
p('A forma canônica em texto é Textopia_, com T maiúsculo, demais letras minúsculas e sublinhado final. A assinatura gráfica oficial emprega a estilização em caixa alta TEXTOPIA_, conforme o arquivo mestre, e constitui exceção deliberada. A forma preferencial nos títulos, interfaces e textos novos é Textopia_.')
table(['Contexto','Forma'],[
 ['Texto e interface','Textopia_'],['Nome acessível','Textopia'],['Pronúncia','tecs-TÓ-pi-a'],['URL e domínio','forma registrada pelo serviço, sem simular o sublinhado quando não permitido'],['Código e pacote','convenção técnica do ambiente, documentada sem redefinir a marca'],
],[2.0,4.5])
p('O sublinhado integra a assinatura visual, mas não deve produzir ruído em leitores de tela. Componentes visuais usam nome acessível “Textopia”, e a descrição longa pode informar que a marca possui sublinhado final. A voz é clara, específica, respeitosa, não acusatória e orientada à ação. Mensagens não culpabilizam, não presumem deficiência e distinguem sistema, professor e estudante.')

h('11.3 Sistema de assinatura visual')
h('11.3.1 Assinatura principal',3)
p('A assinatura principal é a composição horizontal apresentada na abertura: símbolo à esquerda, palavra TEXTOPIA_ em azul-marinho e verde-água e sublinhado amarelo. O sublinhado funciona como sinal de continuidade e espaço para a próxima decisão do autor. A assinatura não deve ser redigitada em peças finais; devem ser usados os arquivos mestres.')
h('11.3.2 Símbolo',3)
p('O símbolo corresponde ao módulo situado à esquerda da assinatura: inicial T enquadrada por marcas de foco, arco verde-água e sublinhado amarelo. É reservado a favicon, ícone de aplicativo, avatar e áreas em que a assinatura completa seja ilegível. Sempre que houver espaço, a primeira ocorrência em uma peça deve usar a assinatura completa.')
h('11.3.3 Versões e fundos',3)
table(['Versão','Uso'],[
 ['Principal','grafite e verde sobre branco ou fundo suave'],['Negativa','branco com sublinhado claro sobre verde escuro ou grafite'],['Monocromática positiva','grafite ou preto em impressão de uma cor'],['Monocromática negativa','branco em fundos escuros'],['Símbolo','áreas quadradas e reduzidas'],
],[2.2,4.3])
h('11.3.4 Proteção e tamanho mínimo',3)
p('A unidade x corresponde à altura do sublinhado da assinatura. A área livre mínima é 4x em todos os lados. Nenhum texto, borda ou imagem deve entrar nessa área. A assinatura horizontal mede no mínimo 120 CSS pixels em tela e 30 mm em impressão; o símbolo mede no mínimo 24 CSS pixels ou 8 mm. Em telas de alta densidade, o ativo rasterizado deve ser fornecido em 2× ou 3×, sem alterar o tamanho CSS.')
h('11.3.5 Usos proibidos',3)
bullets(['não distorcer, inclinar, rotacionar ou recompor;','não alterar tipografia, proporções ou posição do sublinhado;','não aplicar sombra, contorno, gradiente ou transparência;','não usar sobre fundo sem contraste ou imagem complexa sem área de respiro;','não substituir o verde ou o grafite por cores não aprovadas;','não usar o símbolo isolado como primeira identificação quando houver espaço para a assinatura.'])

h('11.4 Arquivos mestres e nomenclatura')
p('Os mestres vetoriais usam SVG para produtos digitais e PDF vetorial para impressão. PNG transparente é derivado de distribuição. Cada arquivo contém versão, variante, espaço de cor e tamanho quando rasterizado.')
table(['Padrão','Exemplo'],[['Nome','textopia_<ativo>_<variante>_<versao>.<ext>'],['Assinatura','textopia_wordmark_principal_v1.0.svg'],['Símbolo','textopia_symbol_principal_v1.0.svg'],['Raster','textopia_wordmark_principal_v1.0_2x.png']], [1.5,5.0])
p('O repositório oficial mantém manifesto com SHA-256, licença, responsável, estado e data de publicação. Arquivos publicados são imutáveis; correções recebem nova versão. Ativos descontinuados permanecem identificados, mas não são oferecidos como padrão.')

h('11.5 Paleta cromática')
colors=[
 ('Azul-marinho 900','#0D1B3D','assinatura, títulos e superfícies institucionais'),
 ('Verde 700','#176B63','ação primária acessível e foco sobre fundo claro'),
 ('Verde-água 500','#4AA39A','marca, acentos, gráficos e superfícies; não usar como texto normal sobre branco'),
 ('Amarelo 500','#F4B400','destaque e atenção com texto grafite; nunca texto normal sobre branco'),
 ('Grafite 900','#1F2933','texto principal e assinatura'),
 ('Cinza 700','#52606D','texto secundário'),
 ('Cinza 100','#F5F7FA','fundo suave'),
 ('Branco','#FFFFFF','fundo principal e texto sobre verde 700'),
 ('Vermelho 700','#B42318','erro e ação destrutiva'),
 ('Azul 700','#175CD3','informação e links'),
]
t=table(['Cor','Valor','Função'],colors,[1.4,1.2,3.9])
for row,c in zip(t.rows[1:],colors): shade(row.cells[0],c[1]); row.cells[0].paragraphs[0].runs[0].font.color.rgb=RGBColor(255,255,255) if c[1] not in ('#F4B400','#F5F7FA','#FFFFFF') else RGBColor(31,41,51)
h('11.5.1 Contraste aprovado',3)
table(['Combinação','Razão aproximada','Uso'],[
 ['Grafite 900 / branco','14,76:1','texto normal e títulos'],['Cinza 700 / branco','6,46:1','texto secundário'],['Branco / verde 700','6,33:1','texto e ícones essenciais'],['Grafite 900 / amarelo 500','7,99:1','alerta de atenção'],['Verde 500 / branco','3,00:1','gráficos e componentes; não texto normal'],['Amarelo 500 / branco','1,85:1','combinação proibida para informação'],
],[2.6,1.4,2.5])
p('Texto normal exige pelo menos 4,5:1; texto grande, componentes, bordas essenciais e foco exigem 3:1. Texto grande significa pelo menos 24 CSS pixels em peso normal ou aproximadamente 18,66 CSS pixels em negrito. Estados combinam cor, texto e ícone. As mesmas verificações devem ser repetidas para modo escuro e qualquer nova cor.')

h('11.6 Tipografia')
p('Inter é a família principal das interfaces. A pilha é Inter, Arial, sans-serif; deve ser hospedada segundo licença e política de privacidade aplicáveis. Arial é a família editorial e alternativa de sistema para assegurar compatibilidade nos documentos. Código, JSON e identificadores usam ui-monospace, SFMono-Regular, Consolas, Liberation Mono, monospace.')
table(['Token','Tamanho/entrelinha','Peso','Uso'],[
 ['display','40/48 px','700','aberturas excepcionais'],['heading-1','32/40 px','700','título de página'],['heading-2','24/32 px','600','seção'],['heading-3','20/28 px','600','subseção'],['body','16/24 px','400','texto de interface'],['body-strong','16/24 px','600','ênfase funcional'],['small','14/20 px','400','metadados não essenciais'],['label','14/20 px','600','rótulos e controles'],['code','14/20 px','400','código e JSON'],
],[1.5,1.7,1.0,2.3])
p('Tamanhos de interface são implementados em rem, preservando preferências do usuário. Caixa alta é reservada a siglas. Sublinhado fica reservado a links ou à própria assinatura. Não se usa peso como único indicador de estado. Em impressão, o corpo editorial parte de 11 pt com entrelinha de aproximadamente 1,35, ajustada após prova de leitura.')

h('11.7 Tokens do design system')
p('A fonte canônica dos tokens é um registro versionado compartilhado pelo código e pela biblioteca de design. Tokens primitivos guardam valores; tokens semânticos exprimem função. Nomes seguem a estrutura categoria.papel.estado.')
table(['Categoria','Exemplos'],[
 ['Cor','color.text.primary, color.action.primary, color.feedback.warning'],['Espaço','space.0, space.1=4px, space.2=8px, space.3=12px, space.4=16px, space.6=24px, space.8=32px, space.12=48px'],['Raio','radius.none=0, radius.field=4px, radius.card=8px, radius.pill=999px'],['Borda','border.default=1px, border.focus=2px'],['Elevação','elevation.1 e elevation.2, somente para hierarquia funcional'],['Movimento','motion.fast=120ms, motion.normal=200ms, motion.slow=320ms'],['Layout','content.max=1200px e breakpoints definidos no pacote do produto'],
],[1.5,5.0])
p('O pacote é publicado em JSON e transformado automaticamente em CSS e formatos das plataformas usadas. Nenhum valor manual deve substituir um token existente. Alterações recebem versão, changelog, testes de contraste e validação visual.')

h('11.8 Layout e responsividade')
p('A interface usa grade fluida de 4 colunas em telas estreitas, 8 em telas médias e 12 em telas largas. Margens mínimas são 16, 24 e 32 pixels, respectivamente; o conteúdo principal possui largura máxima de 1200 pixels. A ordem de leitura e de foco acompanha a ordem visual e não depende de reposicionamento apenas por CSS.')
p('Em ampliação e telas estreitas, painéis se tornam fluxo vertical; ações essenciais permanecem próximas ao objeto; tabelas usam cartões, rolagem identificada ou colunas prioritárias; nomes longos quebram sem truncar informação essencial. A localização deve admitir expansão textual de pelo menos 30%.')

h('11.9 Componentes e estados')
p('Botões, campos, seletores, tabelas, cartões, abas, menus, diálogos, alertas e notificações possuem estados default, hover, focus-visible, active, selected, disabled, loading, success e error quando aplicáveis. Cada estado combina forma, texto, ícone e cor. Áreas interativas têm pelo menos 44 × 44 CSS pixels, salvo controles compactos com alvo equivalente.')
p('Campos conservam rótulo visível, instrução associada e mensagem de erro que identifica problema e ação. Diálogos movem o foco para o título, contêm descrição e ações claras e devolvem o foco ao acionador. Ações destrutivas informam objeto, consequência e possibilidade de recuperação antes da confirmação.')
h('11.9.1 Estados do sistema',3)
table(['Estado','Representação'],[
 ['Em processamento','rótulo textual, progresso quando mensurável e cancelamento quando seguro'],['Retido','motivo, responsável pela próxima ação e ausência de prazo falso'],['Abstenção','indicação de que não houve conclusão sustentada e encaminhamento disponível'],['Falha','problema, preservação dos dados e ação possível'],['Sem achados','confirmação de conclusão, sem sugerir perfeição absoluta'],['Indisponível','razão comunicável e alternativa ou tentativa posterior'],
],[1.5,5.0])

h('11.10 Padrões da revisão pedagógica')
p('A interface representa separadamente intervenção sentencial, achado global estruturado e feedback global sintetizador. Evidência, diagnóstico e orientação não são fundidos. O texto original permanece identificável e não é alterado automaticamente.')
table(['Unidade','Elementos visuais obrigatórios'],[
 ['Intervenção sentencial','trecho literal, critério, categoria, diagnóstico, orientação, prioridade, estado de validação e decisões'],['Achado global','diagnóstico suprassentencial, sentenças relacionadas, orientação, prioridade e estado'],['Feedback global','pontos fortes, prioridades referenciadas, comentário e limitações; sem diagnósticos novos'],
],[1.6,4.9])
p('Prioridade usa rótulo e ordem, não somente cor. Confiança aparece como “confiança do processamento”, acompanhada de explicação e política aplicada, nunca como julgamento do estudante. Evidências positivas têm presença visual equivalente às limitações, sem transformar revisão em inventário de falhas.')
p('As decisões aceitar, adaptar, rejeitar e contestar usam verbos claros e explicam seus efeitos. Aceitar uma orientação não altera o texto. Comparações entre versões oferecem visão linear acessível além da marcação visual de inserções e exclusões. Padrões recorrentes são agrupados com expansão controlada.')

h('11.11 Avaliação, nota e integridade acadêmica')
p('Revisão, validação, avaliação, cálculo, decisão docente e nota publicada possuem títulos e estados distintos. Uma proposta calculada nunca recebe o mesmo tratamento visual da nota publicada. Alterações e retificações preservam o histórico e informam responsável e data.')
p('Indícios de integridade acadêmica usam linguagem cautelosa, neutra e não acusatória, ícone informativo e indicação de análise humana necessária. Não se usa vermelho de erro para acusar o estudante, nem linguagem de certeza para inferências automatizadas.')

h('11.12 Iconografia, imagens e visualização de dados')
p('Ícones usam desenho linear de 2 pixels em grade de 24 pixels, terminais arredondados e proporções consistentes. Ícones interativos possuem nome acessível; ícone nunca substitui sozinho o rótulo de uma ação crítica. Símbolos de alerta, erro, sucesso e informação mantêm significados estáveis.')
p('Fotografias e ilustrações representam diversidade sem estereótipos e devem possuir finalidade informativa. Imagens decorativas recebem alternativa vazia; imagens informativas recebem texto alternativo que comunica função. Diagramas usam formas e rótulos além de cor.')
p('Gráficos incluem título, unidade, legenda, valores recuperáveis e alternativa tabular ou textual. Paletas são testadas para diferenças de percepção cromática. Não se usam gráficos para inferir aprendizagem exclusivamente por nota, quantidade de erros ou cliques.')

h('11.13 Movimento')
p('Movimento confirma causalidade, continuidade ou mudança de estado; nunca é meramente decorativo em fluxos acadêmicos. Transições usam os tokens de 120 a 320 ms. O sistema respeita prefers-reduced-motion, remove deslocamentos não essenciais e conserva feedback textual. Não há flashes ou animações repetitivas que prejudiquem leitura e atenção.')

h('11.14 Projeto editorial do livro')
p('O livro usa página Carta ou A4 conforme a edição, margens de aproximadamente 25 mm, Arial 11 pt no corpo e hierarquia numerada consistente. Títulos não ficam órfãos; blocos de código podem continuar em outra página somente com identificação; tabelas repetem cabeçalho e não usam linhas de altura fixa.')
table(['Elemento','Padrão'],[
 ['Título do capítulo','Arial 23–26 pt, preto, alinhamento consistente, sem regra decorativa automática'],['Seção','Arial 16–18 pt, negrito, numeração hierárquica'],['Subseção','Arial 13–15 pt, negrito'],['Corpo','Arial 11 pt, entrelinha aproximada 1,35, espaçamento posterior consistente'],['Código e JSON','fonte monoespaçada 8,5–9,5 pt, fundo suave, continuação identificada'],['Tabela','cabeçalho destacado, contraste adequado, largura explícita e repetição em nova página'],['Figura','número, título informativo, texto alternativo e fonte ou autoria'],
],[1.6,4.9])
p('Sumário, referências cruzadas e numeração de figuras e tabelas são gerados, não digitados manualmente. Diagramas entidade-relacionamento e máquinas de estados possuem legenda e alternativa textual. O projeto editorial evita páginas compostas apenas por fragmentos de schema sem título de continuação.')

h('11.15 Acessibilidade e inclusão')
bullets([
 'Semântica correta de títulos, regiões, listas, tabelas e controles.',
 'Navegação completa por teclado, foco visível e retorno previsível do foco.',
 'Redimensionamento de texto e ampliação até 400% sem perda de conteúdo ou funcionalidade essencial.',
 'Contraste verificado em todos os estados, inclusive disabled quando ainda informativo.',
 'Alternativas não visuais para imagens, diagramas, gráficos, comparação e estado.',
 'Legendas e transcrições para mídia temporal.',
 'Linguagem clara, mensagens específicas e ausência de pressupostos capacitistas.',
 'Testes automatizados combinados com avaliação manual por pessoas e tecnologias assistivas.',
])

h('11.16 Templates e artefatos')
p('A identidade deve distribuir: assinaturas SVG e PNG; símbolo e favicon; manifesto de ativos; biblioteca de cores e tipografia; pacote versionado de tokens; biblioteca de componentes; templates de documento, apresentação e relatório; padrões de e-mail e redes; modelos de diagramas e gráficos; checklist de contraste; e exemplos de usos corretos e incorretos.')

h('11.17 Governança da marca e do design system')
p('O responsável pela identidade mantém estratégia e ativos; o responsável por acessibilidade homologa fluxos; design e frontend mantêm componentes e tokens; comunicação mantém templates; pedagogia e linguística revisam padrões que afetam feedback e avaliação. Criação, aprovação e publicação de ativos sensíveis seguem segregação de funções.')
p('Ativos e tokens usam estados draft, in_review, approved, published, deprecated e archived. Exceções registram solicitante, finalidade, escopo, justificativa, aprovador e data de expiração. Mudanças publicadas recebem versão, data de vigência, changelog, testes, plano de migração e comunicação.')
p('A revisão da identidade ocorre ao menos a cada ciclo maior do produto ou quando métricas, acessibilidade, nova plataforma ou mudança institucional justificarem. Ativos antigos são descontinuados sem apagar o histórico. Licenças de fontes, imagens e ícones são registradas no manifesto.')

h('11.18 Critérios de aceitação')
bullets([
 'Toda aplicação usa ativo oficial, versão e proporções corretas.',
 'Combinações de cor atendem aos contrastes definidos e nunca comunicam estado apenas por cor.',
 'Componentes usam tokens publicados e possuem estados, teclado e nomes acessíveis.',
 'Revisão, achados globais, feedback global, avaliação e nota permanecem visualmente distintos.',
 'A interface comunica autoria, responsável pela próxima ação e limitações do sistema.',
 'Documentos seguem hierarquia, tipografia, tabelas, figuras e blocos técnicos acessíveis.',
 'Mudanças de marca ou design passam por revisão, testes, aprovação e publicação versionada.',
])

h('11.19 Relação com a documentação canônica')
p('Este capítulo rege identidade visual, design system e aplicação editorial. Não redefine contratos pedagógicos ou técnicos. Quando uma decisão visual afetar significado, estado, autoria, privacidade, avaliação ou acessibilidade, prevalecem os capítulos especializados e os anexos canônicos vigentes. Conflitos não resolvidos bloqueiam a publicação do componente ou do material.')

doc.save(OUT)
print(OUT)
