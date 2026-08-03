from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SRC=Path('/Users/mariomartins2/Library/CloudStorage/GoogleDrive-mario.martins@ufersa.edu.br/Meu Drive/TEXTOPIA/Documentação técnico-pedagógica/Anexo técnico C - Registros canônicos de taxonomias, critérios, categorias e gêneros do Textopia_.docx')
OUT=Path('/private/tmp/textopia_catalogs')
OUT.mkdir(parents=True,exist_ok=True)

genres=[
('RESUMO_ACADEMICO','Resumo acadêmico','Síntese autônoma e concisa de um texto, estudo ou trabalho.'),
('RESENHA_CRITICA','Resenha crítica','Apresentação, síntese e avaliação fundamentada de uma obra.'),
('FICHAMENTO','Fichamento','Registro organizado de ideias, dados e passagens relevantes de uma fonte.'),
('ARTIGO_OPINIAO','Artigo de opinião','Defesa pública de um ponto de vista sobre tema controverso.'),
('ENSAIO_ACADEMICO','Ensaio acadêmico','Discussão autoral e argumentativa de uma questão acadêmica.'),
('RESPOSTA_DISSERTATIVA','Resposta dissertativa','Resposta desenvolvida a uma questão, com tese ou foco explícito.'),
('RELATORIO_TECNICO','Relatório técnico','Descrição e análise de procedimento, situação ou intervenção técnica.'),
('RELATORIO_CIENTIFICO','Relatório científico','Registro de pesquisa com método, resultados e interpretação.'),
('PROJETO_PESQUISA','Projeto de pesquisa','Proposição fundamentada de investigação futura.'),
('PLANO_PESQUISA','Plano de pesquisa','Planejamento operacional de etapas, recursos e prazos de uma investigação.'),
('ARTIGO_CIENTIFICO','Artigo científico','Comunicação completa de pesquisa para circulação acadêmica.'),
('RELATO_EXPERIENCIA','Relato de experiência','Descrição contextualizada e reflexão crítica sobre uma prática vivida.'),
('ESTUDO_CASO','Estudo de caso','Análise aprofundada e contextualizada de um caso delimitado.'),
('MONOGRAFIA_TCC','Monografia ou TCC','Trabalho acadêmico extenso que desenvolve uma investigação ou estudo.'),
]

criteria=[
('ADEQUACAO_GENERO','Adequação ao gênero','discursiva','Atende à finalidade e à organização esperadas para o gênero selecionado.'),
('ATENDIMENTO_TAREFA','Atendimento à tarefa','discursiva','Cumpre objetivos, instruções e limites explicitados na atividade.'),
('PROPOSITO_COMUNICATIVO','Propósito comunicativo','discursiva','Mantém clara a ação comunicativa pretendida.'),
('ADEQUACAO_INTERLOCUTOR','Adequação ao interlocutor','discursiva','Ajusta explicitação, tom e escolhas ao leitor previsto.'),
('DELIMITACAO_TEMA','Delimitação do tema','conteudo','Define e mantém um recorte temático controlável.'),
('TESE_OU_FOCO','Tese ou foco central','conteudo','Formula e sustenta a posição ou o foco que organiza o texto.'),
('PERTINENCIA_CONTEUDO','Pertinência do conteúdo','conteudo','Seleciona informações relevantes para o objetivo.'),
('DESENVOLVIMENTO_IDEIAS','Desenvolvimento das ideias','conteudo','Explica e articula suficientemente as ideias apresentadas.'),
('FIDELIDADE_FONTE','Fidelidade à fonte','conteudo','Representa corretamente ideias e relações atribuídas a uma fonte fornecida.'),
('ESTRUTURA_COMPOSICIONAL','Estrutura composicional','textual','Organiza partes e seções conforme a função de cada uma.'),
('PROGRESSAO_TEMATICA','Progressão temática','textual','Faz o texto avançar sem saltos, estagnações ou desvios.'),
('COERENCIA_GLOBAL','Coerência global','textual','Mantém compatibilidade entre as informações e o sentido global.'),
('COESAO_REFERENCIAL','Coesão referencial','textual','Permite recuperar referentes com clareza.'),
('COESAO_SEQUENCIAL','Coesão sequencial','textual','Explicita relações lógico-semânticas adequadas.'),
('ARTICULACAO_PARAGRAFOS','Articulação entre parágrafos','textual','Constrói transições e relações funcionais entre parágrafos.'),
('ARGUMENTACAO','Argumentação','argumentativa','Apresenta razões pertinentes e articuladas em favor de uma posição.'),
('EVIDENCIAS_SUSTENTACAO','Evidências e sustentação','argumentativa','Sustenta afirmações com dados, exemplos ou razões verificáveis fornecidos.'),
('CONTRA_ARGUMENTACAO','Contra-argumentação','argumentativa','Considera objeções ou posições alternativas quando pertinentes.'),
('ANALISE_INTERPRETACAO','Análise e interpretação','analitica','Vai além da descrição e explica relações, implicações e significados.'),
('PROBLEMA_OBJETIVOS','Problema e objetivos de pesquisa','pesquisa','Mantém coerência entre problema, questão e objetivos.'),
('JUSTIFICATIVA_RELEVANCIA','Justificativa e relevância','pesquisa','Explicita a importância e a contribuição esperada do trabalho.'),
('ADEQUACAO_METODO','Adequação metodológica','pesquisa','Alinha procedimentos, dados e análise aos objetivos declarados.'),
('RESULTADOS_DISCUSSAO','Resultados e discussão','pesquisa','Distingue resultados e os interpreta em relação aos objetivos.'),
('CONCLUSAO','Conclusão','pesquisa','Retoma objetivos e sintetiza respostas sem introduzir afirmações desconectadas.'),
('VIABILIDADE_PLANEJAMENTO','Viabilidade e planejamento','pesquisa','Compatibiliza etapas, recursos, prazos e produtos.'),
('REFLEXAO_EXPERIENCIA','Reflexão sobre a experiência','reflexiva','Relaciona descrição, aprendizagem e análise crítica da experiência.'),
('CLAREZA_PRECISAO','Clareza e precisão','estilo','Evita formulações vagas, ambíguas ou imprecisas.'),
('ADEQUACAO_REGISTRO','Adequação do registro','estilo','Mantém registro compatível com a situação acadêmica e o gênero.'),
('CONCISAO','Concisão','estilo','Evita repetição e extensão sem função comunicativa.'),
('CONSTRUCAO_SENTENCA','Construção da sentença','gramatical','Produz sentenças completas, legíveis e sintaticamente organizadas.'),
('CONCORDANCIA','Concordância','gramatical','Mantém concordância nominal e verbal.'),
('REGENCIA','Regência e complementação','gramatical','Emprega complementos e preposições de modo adequado.'),
('PONTUACAO','Pontuação','gramatical','Usa sinais de pontuação de acordo com a estrutura e o sentido.'),
('ORTOGRAFIA','Ortografia e acentuação','gramatical','Aplica convenções ortográficas e de acentuação.'),
('ESCOLHA_LEXICAL','Escolha lexical','lexical','Seleciona palavras precisas e compatíveis com o contexto.'),
]

categories=[
('DESVIO_GENERO','Desvio de gênero','ADEQUACAO_GENERO','A forma ou a finalidade não corresponde ao gênero selecionado.'),
('DESCUMPRIMENTO_INSTRUCAO','Descumprimento de instrução','ATENDIMENTO_TAREFA','Uma exigência explícita da atividade não foi atendida.'),
('PROPOSITO_INDEFINIDO','Propósito indefinido','PROPOSITO_COMUNICATIVO','A ação comunicativa do trecho ou do texto não fica clara.'),
('INTERLOCUTOR_INADEQUADO','Inadequação ao interlocutor','ADEQUACAO_INTERLOCUTOR','O grau de explicitação ou o tom não atende ao leitor previsto.'),
('TEMA_INDEFINIDO','Tema insuficientemente delimitado','DELIMITACAO_TEMA','O recorte temático está amplo, instável ou indeterminado.'),
('TESE_FOCO_AUSENTE','Tese ou foco ausente','TESE_OU_FOCO','Falta uma posição ou foco organizador recuperável.'),
('TESE_FOCO_INSTAVEL','Tese ou foco instável','TESE_OU_FOCO','A posição ou o foco muda sem desenvolvimento justificável.'),
('INFORMACAO_IRRELEVANTE','Informação irrelevante','PERTINENCIA_CONTEUDO','A informação não contribui para o objetivo do texto.'),
('DESENVOLVIMENTO_INSUFICIENTE','Desenvolvimento insuficiente','DESENVOLVIMENTO_IDEIAS','Uma ideia importante é afirmada sem explicação suficiente.'),
('DISTORCAO_FONTE','Distorção da fonte','FIDELIDADE_FONTE','A representação não corresponde ao material-fonte fornecido.'),
('ATRIBUICAO_AMBIGUA','Atribuição ambígua','FIDELIDADE_FONTE','Não fica claro se a ideia pertence ao autor do texto ou à fonte.'),
('ORGANIZACAO_INADEQUADA','Organização inadequada','ESTRUTURA_COMPOSICIONAL','Uma parte está ausente, deslocada ou exerce função incompatível.'),
('RUPTURA_PROGRESSAO','Ruptura de progressão','PROGRESSAO_TEMATICA','O desenvolvimento apresenta salto, desvio ou estagnação.'),
('RUPTURA_COERENCIA','Ruptura de coerência','COERENCIA_GLOBAL','Há incompatibilidade de sentidos no contexto recuperável.'),
('REFERENCIA_AMBIGUA','Referência ambígua','COESAO_REFERENCIAL','O referente de uma expressão não pode ser identificado com segurança.'),
('REFERENCIA_INCONSISTENTE','Referência inconsistente','COESAO_REFERENCIAL','A cadeia referencial muda ou se rompe indevidamente.'),
('CONECTOR_INADEQUADO','Conector inadequado','COESAO_SEQUENCIAL','O conector sinaliza relação incompatível com o sentido.'),
('RELACAO_NAO_EXPLICITADA','Relação não explicitada','COESAO_SEQUENCIAL','A relação lógico-semântica necessária não está suficientemente marcada.'),
('TRANSICAO_INSUFICIENTE','Transição insuficiente','ARTICULACAO_PARAGRAFOS','A passagem entre parágrafos não explicita sua relação funcional.'),
('RAZAO_INSUFICIENTE','Razão insuficiente','ARGUMENTACAO','Uma posição é apresentada sem razão pertinente suficiente.'),
('FALHA_INFERENCIAL','Falha inferencial','ARGUMENTACAO','A conclusão não decorre adequadamente das razões apresentadas.'),
('AFIRMACAO_NAO_SUSTENTADA','Afirmação não sustentada','EVIDENCIAS_SUSTENTACAO','Uma afirmação relevante carece de sustentação fornecida no contexto.'),
('EVIDENCIA_INADEQUADA','Evidência inadequada','EVIDENCIAS_SUSTENTACAO','A evidência não sustenta a afirmação a que foi associada.'),
('OBJECAO_NAO_RESPONDIDA','Objeção não respondida','CONTRA_ARGUMENTACAO','Uma objeção pertinente é mencionada, mas não integrada ao raciocínio.'),
('ANALISE_SUPERFICIAL','Análise superficial','ANALISE_INTERPRETACAO','O trecho descreve dados ou fatos sem interpretar relações relevantes.'),
('INCOERENCIA_PROBLEMA_OBJETIVO','Incoerência entre problema e objetivo','PROBLEMA_OBJETIVOS','Problema, questão e objetivo não se correspondem.'),
('JUSTIFICATIVA_INSUFICIENTE','Justificativa insuficiente','JUSTIFICATIVA_RELEVANCIA','A relevância ou contribuição esperada não está explicada.'),
('METODO_INCOMPATIVEL','Método incompatível','ADEQUACAO_METODO','O procedimento não permite atender ao objetivo declarado.'),
('RESULTADO_INTERPRETADO_COMO_DADO','Confusão entre resultado e interpretação','RESULTADOS_DISCUSSAO','Dado observado e interpretação não são distinguidos.'),
('CONCLUSAO_DESVINCULADA','Conclusão desvinculada','CONCLUSAO','A conclusão não responde aos objetivos ou introduz afirmação desconectada.'),
('PLANEJAMENTO_INVIAVEL','Planejamento inviável','VIABILIDADE_PLANEJAMENTO','Etapas, recursos ou prazos são incompatíveis.'),
('REFLEXAO_INSUFICIENTE','Reflexão insuficiente','REFLEXAO_EXPERIENCIA','A experiência é narrada sem análise das aprendizagens ou implicações.'),
('IMPRECISAO','Imprecisão','CLAREZA_PRECISAO','A formulação é vaga ou não permite interpretação suficientemente determinada.'),
('AMBIGUIDADE','Ambiguidade','CLAREZA_PRECISAO','A sentença admite leituras concorrentes não resolvidas pelo contexto.'),
('REGISTRO_INADEQUADO','Registro inadequado','ADEQUACAO_REGISTRO','A escolha linguística é incompatível com o gênero ou a situação acadêmica.'),
('REDUNDANCIA','Redundância','CONCISAO','A sentença repete conteúdo sem nova função discursiva.'),
('FRAGMENTO_SENTENCIAL','Fragmento sentencial','CONSTRUCAO_SENTENCA','A unidade apresentada como sentença está sintaticamente incompleta.'),
('PERIODO_MAL_ESTRUTURADO','Período mal estruturado','CONSTRUCAO_SENTENCA','A organização sintática prejudica a interpretação da sentença.'),
('DESVIO_CONCORDANCIA','Desvio de concordância','CONCORDANCIA','Há incompatibilidade de número, pessoa ou gênero em relação sintática.'),
('DESVIO_REGENCIA','Desvio de regência','REGENCIA','A relação de complementação ou a preposição está inadequada.'),
('DESVIO_PONTUACAO','Desvio de pontuação','PONTUACAO','A pontuação contraria a estrutura ou altera indevidamente o sentido.'),
('DESVIO_ORTOGRAFIA','Desvio de ortografia','ORTOGRAFIA','Há grafia ou acentuação incompatível com a convenção adotada.'),
('ESCOLHA_LEXICAL_INADEQUADA','Escolha lexical inadequada','ESCOLHA_LEXICAL','A palavra é imprecisa, incompatível ou produz combinação inadequada.'),
]

profiles={
'BASE':'ADEQUACAO_GENERO, ATENDIMENTO_TAREFA, PROPOSITO_COMUNICATIVO, ADEQUACAO_INTERLOCUTOR, DELIMITACAO_TEMA, PERTINENCIA_CONTEUDO, DESENVOLVIMENTO_IDEIAS, ESTRUTURA_COMPOSICIONAL, PROGRESSAO_TEMATICA, COERENCIA_GLOBAL, COESAO_REFERENCIAL, COESAO_SEQUENCIAL, ARTICULACAO_PARAGRAFOS, CLAREZA_PRECISAO, ADEQUACAO_REGISTRO, CONSTRUCAO_SENTENCA, CONCORDANCIA, REGENCIA, PONTUACAO, ORTOGRAFIA, ESCOLHA_LEXICAL',
'SINTESE':'FIDELIDADE_FONTE, CONCISAO',
'ARGUMENTATIVO':'TESE_OU_FOCO, ARGUMENTACAO, EVIDENCIAS_SUSTENTACAO, CONTRA_ARGUMENTACAO, ANALISE_INTERPRETACAO',
'PESQUISA':'PROBLEMA_OBJETIVOS, JUSTIFICATIVA_RELEVANCIA, ADEQUACAO_METODO, RESULTADOS_DISCUSSAO, CONCLUSAO',
'PROJETO':'PROBLEMA_OBJETIVOS, JUSTIFICATIVA_RELEVANCIA, ADEQUACAO_METODO, VIABILIDADE_PLANEJAMENTO',
'REFLEXIVO':'ANALISE_INTERPRETACAO, REFLEXAO_EXPERIENCIA',
}

genre_profiles=[
('RESUMO_ACADEMICO','BASE + SINTESE'),('RESENHA_CRITICA','BASE + SINTESE + ARGUMENTATIVO'),('FICHAMENTO','BASE + SINTESE'),
('ARTIGO_OPINIAO','BASE + ARGUMENTATIVO'),('ENSAIO_ACADEMICO','BASE + ARGUMENTATIVO'),('RESPOSTA_DISSERTATIVA','BASE + ARGUMENTATIVO'),
('RELATORIO_TECNICO','BASE + PESQUISA'),('RELATORIO_CIENTIFICO','BASE + PESQUISA'),('PROJETO_PESQUISA','BASE + PROJETO'),
('PLANO_PESQUISA','BASE + PROJETO'),('ARTIGO_CIENTIFICO','BASE + ARGUMENTATIVO + PESQUISA'),('RELATO_EXPERIENCIA','BASE + REFLEXIVO'),
('ESTUDO_CASO','BASE + ARGUMENTATIVO + PESQUISA'),('MONOGRAFIA_TCC','BASE + ARGUMENTATIVO + PESQUISA')]

def shade(cell,fill='D9EAF7'):
    p=cell._tc.get_or_add_tcPr(); s=OxmlElement('w:shd'); s.set(qn('w:fill'),fill); p.append(s)

def table(doc,headers,rows):
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    trpr=t.rows[0]._tr.get_or_add_trPr(); trpr.append(OxmlElement('w:tblHeader'))
    for i,h in enumerate(headers):
        t.rows[0].cells[i].text=h; shade(t.rows[0].cells[i])
        for r in t.rows[0].cells[i].paragraphs[0].runs: r.bold=True; r.font.size=Pt(8.5)
    for row in rows:
        new_row=t.add_row()
        cant=OxmlElement('w:cantSplit'); new_row._tr.get_or_add_trPr().append(cant)
        cells=new_row.cells
        for i,v in enumerate(row):
            cells[i].text=str(v); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after=Pt(1)
                for r in p.runs: r.font.size=Pt(8)
    doc.add_paragraph()
    return t

def move_since(doc,start_index,before):
    elems=list(doc._body._body)[start_index:]
    for e in elems: before._p.addprevious(e)

d=Document(SRC)
for p in d.paragraphs:
    if 'registro-linguistico-1.0.0' in p.text:
        for r in p.runs: r.text=r.text.replace('registro-linguistico-1.0.0','registro-linguistico-1.1.0')
    if 'Os exemplos abaixo são contratos mínimos' in p.text:
        p.text='Os exemplos abaixo ilustram os contratos mínimos. As listas canônicas completas e publicadas constam nas seções 3.6 a 3.9; implementações não podem ampliar essas listas sem nova versão.'

before=next(p for p in d.paragraphs if p.text.strip().startswith('4. Regras de validação'))
start=len(list(d._body._body))
d.add_heading('3.6 Catálogo canônico de gêneros — gênero-1.0.0',level=2)
d.add_paragraph('Os gêneros abaixo constituem a lista inicial publicada. A interface deve apresentar o nome, mas persistir o id e a genreCatalogVersion.')
table(d,['id','Nome','Definição'],genres)
d.add_heading('3.7 Catálogo canônico de critérios — taxonomia-1.0.0',level=2)
d.add_paragraph('Todos os critérios possuem status active nesta versão. O código é estável e a descrição delimita o que pode ser diagnosticado.')
table(d,['code','Nome','Dimensão','Definição'],criteria)
d.add_heading('3.8 Catálogo canônico de categorias — taxonomia-1.0.0',level=2)
d.add_paragraph('Cada categoria está vinculada a exatamente um critério. O modelo só pode emitir uma categoria quando seu critério também estiver autorizado.')
table(d,['code','Nome','criterionCode','Definição operacional'],categories)
d.add_heading('3.9 Perfis e mapeamento canônico por gênero — mapeamento-1.0.0',level=2)
d.add_paragraph('O perfil BASE é aplicado a todos os gêneros. Os demais perfis acrescentam critérios específicos. As categorias autorizadas são exatamente as vinculadas aos critérios resultantes da união dos perfis.')
table(d,['Perfil','criterionCodes'],list(profiles.items()))
table(d,['genreId','Perfis aplicáveis'],genre_profiles)
d.add_paragraph('Regra determinística: criterionCodes = união sem duplicatas dos perfis do gênero, preservando a ordem BASE → perfil específico. categoryCodes = todas as categorias da seção 3.8 cujo criterionCode pertença a criterionCodes. O backend registra genreCatalogVersion = 1.0.0, taxonomyVersion = 1.0.0 e mappingVersion = 1.0.0.')
move_since(d,start,before)
# The section property occupies the final body position; explicitly place the
# first appended heading before its introductory paragraph as well.
h36=next(p for p in d.paragraphs if p.text.strip().startswith('3.6 Catálogo canônico de gêneros'))
intro36=next(p for p in d.paragraphs if p.text.strip().startswith('Os gêneros abaixo constituem'))
intro36._p.addprevious(h36._p)

out=OUT/SRC.name
d.save(out)
print(out)
