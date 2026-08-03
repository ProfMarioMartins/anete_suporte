from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont
import json, hashlib, re

BASE=Path('/Users/mariomartins2/Library/CloudStorage/GoogleDrive-mario.martins@ufersa.edu.br/Meu Drive/TEXTOPIA/Documentação técnico-pedagógica')
OUT=Path('/private/tmp/textopia_normative'); OUT.mkdir(parents=True,exist_ok=True)

def shade(cell,fill='D9EAF7'):
    pr=cell._tc.get_or_add_tcPr(); s=OxmlElement('w:shd'); s.set(qn('w:fill'),fill); pr.append(s)
def no_split(row): row._tr.get_or_add_trPr().append(OxmlElement('w:cantSplit'))
def config(doc,title,subtitle,version):
    sec=doc.sections[0]; sec.top_margin=sec.bottom_margin=Inches(.8); sec.left_margin=sec.right_margin=Inches(.85)
    n=doc.styles['Normal']; n.font.name='Calibri'; n.font.size=Pt(10.5); n.paragraph_format.space_after=Pt(5); n.paragraph_format.line_spacing=1.18
    for nm,sz,col,bef,aft in [('Heading 1',16,'2E74B5',16,8),('Heading 2',13,'2E74B5',12,6),('Heading 3',11.5,'1F4D78',9,4)]:
        s=doc.styles[nm]; s.font.name='Calibri'; s.font.size=Pt(sz); s.font.color.rgb=RGBColor.from_string(col); s.paragraph_format.space_before=Pt(bef); s.paragraph_format.space_after=Pt(aft); s.paragraph_format.keep_with_next=True
    p=doc.add_paragraph(); r=p.add_run(title); r.font.name='Calibri'; r.font.size=Pt(22); r.font.color.rgb=RGBColor.from_string('1F4D78')
    p.paragraph_format.space_after=Pt(4)
    p=doc.add_paragraph(); r=p.add_run(subtitle); r.italic=True; r.font.size=Pt(11); r.font.color.rgb=RGBColor.from_string('5B6573'); p.paragraph_format.space_after=Pt(12)
    doc.add_paragraph(version)
    for s in doc.sections:
        f=s.footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.CENTER; f.add_run('TEXTOPIA_ — documentação técnico-pedagógica').font.size=Pt(8)
def table(doc,headers,rows,sizes=None):
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=True
    t.rows[0]._tr.get_or_add_trPr().append(OxmlElement('w:tblHeader'))
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h; shade(c)
        for r in c.paragraphs[0].runs: r.bold=True; r.font.size=Pt(8.5)
    for row in rows:
        rr=t.add_row(); no_split(rr)
        for i,v in enumerate(row):
            c=rr.cells[i]; c.text=str(v); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in c.paragraphs:
                p.paragraph_format.space_after=Pt(1)
                for r in p.runs:r.font.size=Pt(8)
    doc.add_paragraph()
    return t
def code(doc,obj):
    txt=obj if isinstance(obj,str) else json.dumps(obj,ensure_ascii=False,indent=2)
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.18); p.paragraph_format.space_after=Pt(7)
    for line in txt.splitlines():
        r=p.add_run(line+'\n'); r.font.name='Consolas'; r.font.size=Pt(7.7); r.font.color.rgb=RGBColor.from_string('25364A')
def bullet(doc,x): doc.add_paragraph(x,style='List Bullet')

def make_d():
    d=Document(); config(d,'Anexo técnico D','Registro canônico de rubricas, níveis, pesos e cálculo da nota do TEXTOPIA_','Versão normativa: rubricas-1.0.0  |  Estado: publicada  |  Fonte canônica')
    d.add_heading('1. Finalidade e princípios',1)
    d.add_paragraph('Este anexo define os registros que permitem ao backend transformar níveis de desempenho propostos pelo modelo em uma nota reproduzível. O modelo seleciona níveis e apresenta evidências; o backend calcula; o professor confirma, modifica ou rejeita.')
    for x in ['Rubricas publicadas são imutáveis e integralmente versionadas.','Pesos, pontos, escala e arredondamento nunca são inferidos pelo modelo.','A revisão fornece evidências, mas a quantidade de intervenções não se converte mecanicamente em desconto.','Critérios não aplicáveis só podem ser excluídos quando a rubrica os declarar como allowNotApplicable.','Toda alteração docente preserva a proposta original e registra justificativa.']: bullet(d,x)
    d.add_heading('2. Esquema canônico da rubrica',1)
    schema={'$schema':'https://json-schema.org/draft/2020-12/schema','$id':'https://textopia.local/schemas/rubrica/1.0.0','type':'object','additionalProperties':False,'required':['version','rubricId','status','name','scale','criteria','publishedAt'],'properties':{'version':{'const':'1.0.0'},'rubricId':{'type':'string','minLength':1},'status':{'enum':['draft','in_review','approved','published','deprecated','archived']},'name':{'type':'string','minLength':1},'scale':{'type':'object','required':['minimum','maximum','roundingMode','decimalPlaces'],'properties':{'minimum':{'type':'number'},'maximum':{'type':'number'},'roundingMode':{'const':'half_up'},'decimalPlaces':{'type':'integer','minimum':0,'maximum':2}}},'criteria':{'type':'array','minItems':1,'items':{'type':'object','required':['criterionId','name','description','weight','allowNotApplicable','levels'],'properties':{'criterionId':{'type':'string'},'name':{'type':'string'},'description':{'type':'string'},'weight':{'type':'number','exclusiveMinimum':0},'allowNotApplicable':{'type':'boolean'},'levels':{'type':'array','minItems':2,'items':{'type':'object','required':['levelId','label','points','descriptor'],'properties':{'levelId':{'type':'string'},'label':{'type':'string'},'points':{'type':'number','minimum':0},'descriptor':{'type':'string'}}}}}}},'publishedAt':{'type':'string','format':'date-time'}}}
    code(d,schema)
    d.add_heading('3. Rubrica canônica inicial — escrita-academica-1.0.0',1)
    d.add_paragraph('A rubrica inicial usa escala de 0 a 10 e cinco critérios. A soma dos pesos deve ser exatamente 1,000000.')
    crit=[
      ('ADEQUACAO_TAREFA_GENERO','Adequação à tarefa e ao gênero','0,15','Atendimento às instruções, ao propósito, ao interlocutor e ao gênero.'),
      ('CONTEUDO_DESENVOLVIMENTO','Conteúdo e desenvolvimento','0,20','Delimitação, pertinência, foco e desenvolvimento suficiente das ideias.'),
      ('ORGANIZACAO_TEXTUAL','Organização, coerência e coesão','0,20','Estrutura, progressão, coerência e relações coesivas.'),
      ('ARGUMENTACAO_ANALISE','Argumentação, análise e evidências','0,25','Razões, evidências, interpretação e adequação metodológica quando aplicável.'),
      ('LINGUAGEM_CLAREZA','Linguagem, clareza e correção','0,20','Clareza, registro, construção sentencial, léxico e convenções linguísticas.')]
    table(d,['criterionId','Nome','Peso','Escopo'],crit)
    d.add_heading('3.1 Níveis comuns',2)
    levels=[('N4','Excelente','4','Atende plenamente ao critério; decisões são consistentes, precisas e sustentadas.'),('N3','Adequado','3','Atende ao critério com limitações localizadas que não comprometem o resultado global.'),('N2','Em desenvolvimento','2','Atende parcialmente; limitações recorrentes reduzem a eficácia do texto.'),('N1','Inicial','1','Atendimento restrito; problemas relevantes comprometem o critério.'),('N0','Não demonstrado','0','Não há evidência suficiente de atendimento ao critério no texto avaliado.')]
    table(d,['levelId','Rótulo','Pontos','Descritor'],levels)
    d.add_heading('4. Mapeamento canônico entre taxonomia e rubrica',1)
    maps=[
      ('ADEQUACAO_TAREFA_GENERO','ADEQUACAO_GENERO, ATENDIMENTO_TAREFA, PROPOSITO_COMUNICATIVO, ADEQUACAO_INTERLOCUTOR'),
      ('CONTEUDO_DESENVOLVIMENTO','DELIMITACAO_TEMA, TESE_OU_FOCO, PERTINENCIA_CONTEUDO, DESENVOLVIMENTO_IDEAS, FIDELIDADE_FONTE'),
      ('ORGANIZACAO_TEXTUAL','ESTRUTURA_COMPOSICIONAL, PROGRESSAO_TEMATICA, COERENCIA_GLOBAL, COESAO_REFERENCIAL, COESAO_SEQUENCIAL, ARTICULACAO_PARAGRAFOS'),
      ('ARGUMENTACAO_ANALISE','ARGUMENTACAO, EVIDENCIAS_SUSTENTACAO, CONTRA_ARGUMENTACAO, ANALISE_INTERPRETACAO, PROBLEMA_OBJETIVOS, JUSTIFICATIVA_RELEVANCIA, ADEQUACAO_METODO, RESULTADOS_DISCUSSAO, CONCLUSAO, VIABILIDADE_PLANEJAMENTO, REFLEXAO_EXPERIENCIA'),
      ('LINGUAGEM_CLAREZA','CLAREZA_PRECISAO, ADEQUACAO_REGISTRO, CONCISAO, CONSTRUCAO_SENTENCA, CONCORDANCIA, REGENCIA, PONTUACAO, ORTOGRAFIA, ESCOLHA_LEXICAL')]
    # fix a catalog code in presentation
    maps=[(a,b.replace('DESENVOLVIMENTO_IDEAS','DESENVOLVIMENTO_IDEIAS')) for a,b in maps]
    table(d,['rubricCriterionId','taxonomyCodes'],maps)
    d.add_paragraph('O mapeamento possui version = 1.0.0, taxonomyVersion = 1.0.0 e rubricVersion = 1.0.0. Códigos ausentes, duplicados ou não autorizados invalidam a avaliação.')
    d.add_heading('5. Cálculo determinístico',1)
    d.add_paragraph('Para cada critério aplicável i, o backend recupera pontos_i do levelId publicado, peso_i e maxPontos_i. A nota bruta é calculada sem arredondamento intermediário:')
    code(d,'notaBruta = escalaMaxima × [Σ(peso_i × pontos_i / maxPontos_i)] / Σ(peso_i aplicável)')
    for x in ['A escala inicial possui mínimo 0 e máximo 10.','O arredondamento é half_up, aplicado somente ao resultado final, com uma casa decimal.','Se allowNotApplicable = false, N/A invalida a proposta.','Se allowNotApplicable = true, o peso é retirado do numerador e do denominador.','Se nenhum critério permanecer aplicável, a avaliação fica indisponível.','O backend usa aritmética decimal, nunca ponto flutuante binário.']: bullet(d,x)
    d.add_heading('5.1 Exemplo',2)
    table(d,['Critério','Peso','Nível','Pontos/4','Contribuição'],[('Adequação','0,15','N3','0,75','0,1125'),('Conteúdo','0,20','N3','0,75','0,1500'),('Organização','0,20','N2','0,50','0,1000'),('Argumentação','0,25','N3','0,75','0,1875'),('Linguagem','0,20','N4','1,00','0,2000')])
    d.add_paragraph('Soma = 0,75; notaBruta = 10 × 0,75 = 7,5; nota proposta = 7,5.')
    d.add_heading('6. Persistência e estados',1)
    table(d,['Tabela','Campos essenciais','Restrições'],[('rubric_versions','rubric_id, version, status, name, scale_min, scale_max, rounding_mode, decimal_places, published_at','PK rubric_id+version; published imutável'),('rubric_criteria','rubric_id, rubric_version, criterion_id, name, description, weight, allow_na','PK composta; soma dos pesos = 1'),('rubric_levels','rubric_id, rubric_version, criterion_id, level_id, label, points, descriptor','levelId único por critério'),('taxonomy_rubric_mappings','version, taxonomy_version, rubric_id, rubric_version, status','combinação publicada única'),('taxonomy_rubric_mapping_items','mapping_version, taxonomy_code, rubric_criterion_id','FKs válidas'),('grade_calculations','calculation_id, version, evaluation_id, rubric_version, input_hash, raw_score, final_score, state','idempotência por input_hash'),('grade_decisions','decision_id, calculation_id, actor_id, action, previous_score, decided_score, justification, decided_at','append-only')])
    d.add_heading('7. Contrato da nota calculada',1)
    code(d,{'version':'1.0.0','calculationId':'calc-001','evaluationExecutionId':'eval-001','rubricVersion':'1.0.0','mappingVersion':'1.0.0','inputHash':'sha256:…','rawScore':'7.500000','proposedScore':'7.5','rounding':{'mode':'half_up','decimalPlaces':1},'state':'aguardando_decisao_docente','teacherDecision':None})
    d.add_heading('8. Critérios de aceitação',1)
    for x in ['Reprocessar a mesma entrada e versões produz o mesmo inputHash e a mesma nota.','Pesos não somam 1: rubrica não pode ser publicada.','levelId inexistente ou duplicado invalida a avaliação.','O modelo nunca envia pontos, pesos ou nota.','Alteração docente não sobrescreve proposta ou cálculo.','Mudança de fórmula exige nova versão principal do contrato.']: bullet(d,x)
    d.save(OUT/'Anexo técnico D - Registro canônico de rubricas, níveis, pesos e cálculo da nota do Textopia_.docx')

def er_png(path):
    W,H=2200,1450; im=Image.new('RGB',(W,H),'white'); dr=ImageDraw.Draw(im)
    try: f=ImageFont.truetype('/System/Library/Fonts/Supplemental/Arial.ttf',28); fb=ImageFont.truetype('/System/Library/Fonts/Supplemental/Arial Bold.ttf',31)
    except: f=ImageFont.load_default(); fb=f
    boxes={
      'Instituição':(80,80,480,210),'Usuário':(80,330,480,460),'Atividade':(660,80,1080,210),'Config. análise':(660,330,1080,460),
      'Submissão':(1260,80,1660,210),'Versão textual':(1260,330,1660,460),'Execução revisão':(1260,580,1660,710),
      'Intervenção':(1760,500,2140,630),'Achado global':(1760,690,2140,820),'Validação':(1260,830,1660,960),
      'Avaliação':(660,830,1080,960),'Avaliação critério':(660,1080,1080,1210),'Cálculo nota':(80,830,480,960),'Decisão docente':(80,1080,480,1210),
      'Rubrica/versão':(660,580,1080,710)}
    def center(b):return((b[0]+b[2])//2,(b[1]+b[3])//2)
    links=[('Instituição','Usuário','1:N'),('Instituição','Atividade','1:N'),('Usuário','Atividade','professor'),('Atividade','Config. análise','1:1'),('Atividade','Submissão','1:N'),('Submissão','Versão textual','1:N'),('Versão textual','Execução revisão','1:N'),('Execução revisão','Intervenção','1:N'),('Execução revisão','Achado global','1:N'),('Execução revisão','Validação','1:N'),('Validação','Avaliação','0..1:N'),('Rubrica/versão','Avaliação','1:N'),('Avaliação','Avaliação critério','1:N'),('Avaliação','Cálculo nota','1:1'),('Cálculo nota','Decisão docente','1:N')]
    for a,b,l in links:
        x1,y1=center(boxes[a]); x2,y2=center(boxes[b]); dr.line((x1,y1,x2,y2),fill='#718096',width=4); dr.text(((x1+x2)//2,(y1+y2)//2),l,font=f,fill='#334E68',anchor='mm')
    for name,b in boxes.items():
        dr.rounded_rectangle(b,18,fill='#EAF3F9',outline='#2E74B5',width=5); dr.text(center(b),name,font=fb,fill='#1F4D78',anchor='mm')
    im.save(path)

def make_e():
    img=OUT/'modelo-dominio-er.png'; er_png(img)
    d=Document(); config(d,'Anexo técnico E','Modelo canônico do domínio, máquinas de estados e segmentação textual do TEXTOPIA_','Versão normativa: dominio-1.0.0  |  Estado: publicada  |  Fonte canônica')
    d.add_heading('1. Escopo',1); d.add_paragraph('Este anexo define as entidades centrais, seus vínculos, os estados operacionais e o contrato determinístico que transforma uma versão textual em parágrafos e sentenças identificáveis.')
    d.add_heading('2. Regras comuns',1)
    for x in ['Todos os registros possuem version, createdAt e identificador estável.','Conteúdo textual e artefatos publicados são imutáveis.','Transições de estado geram eventos append-only com ator, origem e correlationId.','Operações capazes de duplicar processamento exigem idempotencyKey.','Relacionamentos entre instituições usam institutionId e são validados no backend.']: bullet(d,x)
    d.add_heading('3. Diagrama entidade-relacionamento',1); d.add_picture(str(img),width=Inches(6.45)); d.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    d.add_paragraph('Figura E.1 — Modelo lógico central. As cardinalidades representam o fluxo normativo; tabelas auxiliares de autenticação, anexos e telemetria podem ser acrescentadas sem alterar os contratos.')
    d.add_heading('4. Dicionário do domínio',1)
    ents=[('institutions','Organização responsável pelo isolamento de dados.','institution_id'),('users','Identidade; papéis são atribuídos por contexto.','user_id'),('activities','Configuração pedagógica publicada pelo professor.','activity_id'),('activity_analysis_configs','Gênero, versões e critérios habilitados.','activity_id+version'),('submissions','Entrega lógica do estudante.','submission_id'),('text_versions','Conteúdo imutável e hash.','text_version_id'),('text_segments','Parágrafos e sentenças com offsets.','text_version_id+sentence_id'),('revision_executions','Chamada versionada ao mecanismo de revisão.','revision_execution_id'),('interventions','Achados restritos a uma sentença.','intervention_id'),('global_findings','Achados suprassentenciais estruturados.','global_finding_id'),('validation_runs','Envelope de decisões humanas.','validation_id'),('evaluation_executions','Proposta de níveis da rubrica.','evaluation_execution_id'),('criterion_assessments','Nível e evidências por critério.','evaluation_id+criterion_id'),('grade_calculations','Cálculo determinístico da nota.','calculation_id'),('grade_decisions','Confirmação, modificação ou rejeição docente.','decision_id'),('domain_events','Histórico de transições e auditoria.','event_id')]
    table(d,['Entidade','Responsabilidade','Chave'],ents)
    d.add_heading('5. Esquemas de persistência',1)
    tables=[
      ('activities','activity_id UUID PK; institution_id UUID FK; owner_user_id UUID FK; version VARCHAR; title TEXT; status VARCHAR; published_at TIMESTAMPTZ; row_version BIGINT','UNIQUE(activity_id,version); estado publicado imutável'),
      ('activity_analysis_configs','activity_id UUID; version VARCHAR; genre_id VARCHAR; genre_catalog_version VARCHAR; taxonomy_version VARCHAR; mapping_version VARCHAR; enabled_criterion_codes JSONB; disabled_criterion_codes JSONB; revision_scope VARCHAR; checksum VARCHAR','PK(activity_id,version); disabled ⊆ enabled/mapeamento'),
      ('submissions','submission_id UUID PK; activity_id UUID FK; student_user_id UUID FK; status VARCHAR; current_text_version_id UUID; created_at TIMESTAMPTZ','uma submissão ativa por estudante quando configurado'),
      ('text_versions','text_version_id UUID PK; submission_id UUID FK; ordinal INT; content_uri TEXT; text_hash VARCHAR; char_length INT; normalization VARCHAR; created_at TIMESTAMPTZ','UNIQUE(submission_id,ordinal); UNIQUE(text_hash,submission_id)'),
      ('text_segments','text_version_id UUID FK; paragraph_id VARCHAR; sentence_id VARCHAR; paragraph_index INT; sentence_index INT; start_offset INT; end_offset INT; literal_text TEXT; segment_hash VARCHAR','PK(text_version_id,sentence_id); offsets não sobrepostos'),
      ('revision_executions','revision_execution_id UUID PK; text_version_id UUID FK; state VARCHAR; prompt_version VARCHAR; contract_version VARCHAR; model_version VARCHAR; taxonomy_version VARCHAR; context_checksum VARCHAR; idempotency_key VARCHAR; created_at TIMESTAMPTZ','UNIQUE(idempotency_key)'),
      ('interventions','intervention_id UUID PK; revision_execution_id UUID FK; sentence_id VARCHAR; criterion_code VARCHAR; category_code VARCHAR; excerpt TEXT; diagnosis TEXT; guidance TEXT; priority VARCHAR; confidence NUMERIC','FK sentence; confidence 0..1'),
      ('global_findings','global_finding_id UUID PK; revision_execution_id UUID FK; criterion_code VARCHAR; category_code VARCHAR; sentence_ids JSONB; diagnosis TEXT; guidance TEXT; priority VARCHAR; confidence NUMERIC','sentence_ids ≥ 2 ou justificativa suprassentencial'),
      ('validation_runs','validation_id UUID PK; revision_execution_id UUID FK; state VARCHAR; actor_id UUID; decisions JSONB; created_at TIMESTAMPTZ','append-only'),
      ('evaluation_executions','evaluation_execution_id UUID PK; text_version_id UUID FK; revision_execution_id UUID FK; rubric_version VARCHAR; mapping_version VARCHAR; state VARCHAR; input_hash VARCHAR','UNIQUE(input_hash)'),
      ('criterion_assessments','evaluation_execution_id UUID FK; rubric_criterion_id VARCHAR; level_id VARCHAR; evidence_refs JSONB; confidence NUMERIC; limitations JSONB','PK(evaluation_execution_id,rubric_criterion_id)'),
      ('grade_calculations','calculation_id UUID PK; evaluation_execution_id UUID FK; formula_version VARCHAR; raw_score NUMERIC; proposed_score NUMERIC; state VARCHAR; input_hash VARCHAR','UNIQUE(input_hash)'),
      ('grade_decisions','decision_id UUID PK; calculation_id UUID FK; actor_id UUID FK; action VARCHAR; previous_score NUMERIC; decided_score NUMERIC; justification TEXT; decided_at TIMESTAMPTZ','append-only'),
      ('domain_events','event_id UUID PK; aggregate_type VARCHAR; aggregate_id UUID; from_state VARCHAR; to_state VARCHAR; actor_id UUID; correlation_id UUID; payload JSONB; occurred_at TIMESTAMPTZ','índice aggregate+occurred_at')]
    table(d,['Tabela','Campos','Restrições'],tables)
    d.add_heading('6. Seleção efetiva de critérios',1)
    d.add_paragraph('A configuração da atividade deve declarar enabledCriterionCodes e disabledCriterionCodes. Ambos são códigos explícitos e versionados. disabledCriterionCodes é opcional e existe para permitir desativação pedagógica sem publicar novo mapeamento de gênero.')
    code(d,{'version':'1.0.0','activityId':'ATV-001','genreId':'ARTIGO_OPINIAO','genreCatalogVersion':'1.0.0','taxonomyVersion':'1.0.0','mappingVersion':'1.0.0','enabledCriterionCodes':['ADEQUACAO_GENERO','TESE_OU_FOCO','ARGUMENTACAO','COERENCIA_GLOBAL','CLAREZA_PRECISAO'],'disabledCriterionCodes':['CONTRA_ARGUMENTACAO'],'revisionScope':'sentence','status':'published'})
    code(d,'mapped = códigos publicados para o gênero\nrequested = enabledCriterionCodes, ou mapped quando ausente\neffective = (mapped ∩ requested) − disabledCriterionCodes\nAUTHORIZED_TAXONOMY = critérios effective + categorias vinculadas')
    d.add_paragraph('O backend rejeita códigos habilitados fora de mapped, códigos simultaneamente habilitados e desabilitados e effective vazio. O professor nunca amplia o mapeamento publicado por simples configuração.')
    d.add_heading('7. Contrato de achados globais',1)
    gf={'$schema':'https://json-schema.org/draft/2020-12/schema','$id':'https://textopia.local/schemas/achado-global/1.0.0','type':'object','additionalProperties':False,'required':['id','criterioCodigo','categoriaCodigo','sentenceIds','diagnostico','orientacao','prioridade','confianca'],'properties':{'id':{'type':'string','minLength':1},'criterioCodigo':{'type':'string','minLength':1},'categoriaCodigo':{'type':'string','minLength':1},'sentenceIds':{'type':'array','minItems':1,'uniqueItems':True,'items':{'type':'string','minLength':1}},'diagnostico':{'type':'string','minLength':1},'orientacao':{'type':'string','minLength':1},'prioridade':{'enum':['baixa','media','alta']},'confianca':{'type':'number','minimum':0,'maximum':1}}}
    code(d,gf)
    d.add_paragraph('sentenceIds contém todas as sentenças que sustentam o achado. Um achado pode usar uma sentença quando o diagnóstico depender também da posição, da seção ou da função global; essa condição deve ser explicada no diagnóstico. Códigos devem pertencer a AUTHORIZED_TAXONOMY.')
    d.add_heading('8. Contrato de segmentação',1)
    seg={'$schema':'https://json-schema.org/draft/2020-12/schema','$id':'https://textopia.local/schemas/segmentacao/1.0.0','type':'object','additionalProperties':False,'required':['version','textVersionId','textHash','normalization','offsetUnit','paragraphs'],'properties':{'version':{'const':'1.0.0'},'textVersionId':{'type':'string'},'textHash':{'type':'string','pattern':'^sha256:'},'normalization':{'const':'NFC'},'offsetUnit':{'const':'unicode_code_point'},'paragraphs':{'type':'array','items':{'type':'object','required':['paragraphId','paragraphIndex','inicio','fim','conteudoLiteral','sentences'],'properties':{'paragraphId':{'type':'string'},'paragraphIndex':{'type':'integer','minimum':0},'inicio':{'type':'integer','minimum':0},'fim':{'type':'integer','minimum':0},'conteudoLiteral':{'type':'string'},'sentences':{'type':'array','items':{'type':'object','required':['sentenceId','sentenceIndex','paragraphId','inicio','fim','conteudoLiteral','segmentHash'],'properties':{'sentenceId':{'type':'string'},'sentenceIndex':{'type':'integer','minimum':0},'paragraphId':{'type':'string'},'inicio':{'type':'integer','minimum':0},'fim':{'type':'integer','minimum':0},'conteudoLiteral':{'type':'string'},'segmentHash':{'type':'string','pattern':'^sha256:'}}}}}}}}}
    code(d,seg)
    for x in ['Offsets usam intervalo semiaberto [inicio,fim) sobre pontos de código Unicode após normalização NFC.','conteudoLiteral deve ser exatamente text[inicio:fim], sem trim ou substituição.','paragraphId = p-{paragraphIndex}; sentenceId = s-{índice global}; IDs são imutáveis dentro de textVersionId.','Cada nova versão textual é segmentada novamente; IDs não são transportados por similaridade.','Abreviaturas, listas, títulos e citações devem integrar um conjunto versionado de testes de segmentação.','Divergência entre hash, offsets e conteúdo bloqueia a chamada ao modelo.']: bullet(d,x)
    d.add_heading('9. Máquinas de estados',1)
    machines={
      'Atividade':[('rascunho','em_revisao','submeter','professor'),('em_revisao','publicada','publicar','professor autorizado'),('publicada','encerrada','encerrar','professor'),('publicada','arquivada','arquivar','administrador')],
      'Submissão':[('rascunho','submetida','submeter','estudante'),('submetida','em_processamento','enfileirar revisão','sistema'),('em_processamento','revisao_disponivel','publicar revisão','sistema/professor'),('revisao_disponivel','reescrita_em_andamento','iniciar reescrita','estudante'),('reescrita_em_andamento','nova_versao_submetida','submeter versão','estudante'),('qualquer não terminal','cancelada','cancelar','estudante/professor')],
      'Revisão':[('criada','em_processamento','iniciar','worker'),('em_processamento','validacao_automatica','receber saída','worker'),('validacao_automatica','aguardando_validacao','reter','sistema'),('validacao_automatica','disponivel','publicar direto','sistema'),('validacao_automatica','invalidada','rejeitar saída','validador'),('qualquer transitório','falha','falhar','sistema'),('aguardando_validacao','disponivel','aprovar','professor'),('disponivel','encerrada','encerrar ciclo','sistema')],
      'Validação':[('pendente','em_analise','assumir','professor'),('em_analise','validada','aprovar tudo','professor'),('em_analise','parcialmente_validada','decidir itens','professor'),('em_analise','rejeitada','rejeitar','professor'),('pendente','expirada','expirar','sistema')],
      'Avaliação':[('criada','em_processamento','iniciar','worker'),('em_processamento','proposta_disponivel','validar saída','sistema'),('em_processamento','indisponivel','abster/falhar','sistema'),('proposta_disponivel','calculada','calcular nota','backend'),('qualquer transitório','falha','falhar','sistema')],
      'Nota':[('calculada','aguardando_decisao_docente','publicar proposta','backend'),('aguardando_decisao_docente','confirmada','confirmar','professor'),('aguardando_decisao_docente','modificada','modificar+justificar','professor'),('aguardando_decisao_docente','rejeitada','rejeitar+justificar','professor'),('confirmada/modificada','publicada','publicar','professor/sistema')],
    }
    for name,rows in machines.items(): d.add_heading('9.'+str(list(machines).index(name)+1)+' '+name,2); table(d,['De','Para','Evento','Autoridade'],rows)
    d.add_heading('10. Relações e invariantes entre estados',1)
    for x in ['Uma revisão só começa quando atividade está publicada e submissão está submetida.','Avaliação só começa para textVersionId idêntico ao da revisão validada.','Avaliação indisponível nunca cria cálculo de nota.','Nota publicada exige decisão docente confirmada ou modificada.','Nova versão textual não altera revisão, avaliação ou nota da versão anterior.','Falha de revisão não altera a imutabilidade da versão textual.','Repetição com a mesma idempotencyKey devolve o agregado existente.']: bullet(d,x)
    d.add_heading('11. Critérios de aceitação',1)
    for x in ['Todos os FKs respeitam institutionId.','Nenhuma transição não listada é aceita.','Offsets recompõem literalmente o texto normalizado.','Achados globais referenciam sentenceIds existentes.','Critérios efetivos são reproduzíveis a partir das quatro versões e da seleção docente.','Eventos permitem reconstruir a sequência de estados sem consultar logs de aplicação.']: bullet(d,x)
    d.save(OUT/'Anexo técnico E - Modelo canônico do domínio, máquinas de estados e segmentação textual do Textopia_.docx')

def replace_runs(p,old,new):
    if old not in p.text:return
    text=p.text.replace(old,new)
    for r in p.runs:r.text=''
    p.runs[0].text=text if p.runs else ''
    if not p.runs:p.add_run(text)

def update_a():
    src=BASE/'Anexo técnico A - Prompt canônico de revisão do Textopia_.docx'; d=Document(src)
    for p in d.paragraphs:
        replace_runs(p,'revisao-1.0.0','revisao-1.1.0'); replace_runs(p,'Contrato de saída: 1.0.0','Contrato de saída: 1.1.0')
        replace_runs(p,'Registre a limitação em feedbackGlobal.limitacoes.','Registre a limitação em feedbackGlobal.limitacoes; não invente código.')
        if p.text.startswith('Produza intervenções localizadas para problemas'):
            p.text='Produza intervencoes para problemas identificáveis em uma única sentença, achadosGlobais estruturados para fenômenos suprassentenciais e feedbackGlobal apenas como síntese pedagógica.'
        if p.text.startswith('Use feedbackGlobal para aspectos'):
            p.text='Use achadosGlobais para todo fenômeno que dependa de relações entre sentenças, parágrafos, seções ou partes maiores. Cada achado deve conter id, criterioCodigo, categoriaCodigo, sentenceIds, diagnostico, orientacao, prioridade e confianca. Use exclusivamente códigos de AUTHORIZED_TAXONOMY e sentenceIds de SENTENCE_MAP.'
        if p.text.startswith('Não transforme problema suprassentencial'):
            p.text='Não transforme problema suprassentencial em intervenção localizada. Use feedbackGlobal para sintetizar pontos fortes, prioridades, comentário final e limitações, sem substituir os achadosGlobais estruturados.'
        if p.text=='Registre a razão em feedbackGlobal.limitacoes.': p.text='Registre a razão em feedbackGlobal.limitacoes. Se houver diagnóstico suprassentencial autorizado e sustentado, registre-o também em achadosGlobais.'
        if p.text.startswith('a) todos os campos obrigatórios'):
            p.text=p.text.replace('h) problemas globais não foram apresentados como intervenções;','h) problemas globais estão em achadosGlobais com códigos autorizados e sentenceIds existentes;')
        if p.text.startswith('Qualquer alteração textual exige nova promptVersion'):
            p.text='Qualquer alteração textual exige nova promptVersion e teste de conformidade. O capítulo 9 não mantém cópia manual: a reprodução editorial é gerada a partir deste anexo e validada por versão e SHA-256.'
    d.save(OUT/src.name)

def update_b():
    src=BASE/'Anexo técnico B - Prompt canônico de avaliação rubricada e proposta de nota do Textopia_.docx'; d=Document(src)
    for p in d.paragraphs:
        if p.text.startswith('Qualquer alteração textual exige nova'):
            p.text='Qualquer alteração textual exige nova gradingPromptVersion e teste de conformidade. O capítulo 9 não mantém cópia manual: a reprodução editorial é gerada a partir deste anexo e validada por versão e SHA-256.'
    d.save(OUT/src.name)

def update_c():
    src=BASE/'Anexo técnico C - Registros canônicos de taxonomias, critérios, categorias e gêneros do Textopia_.docx'; d=Document(src)
    for p in d.paragraphs:
        # Examples and canonical catalog now use the same taxonomy version.
        replace_runs(p,'"taxonomyVersion": "1.2.0"','"taxonomyVersion": "1.0.0"')
        replace_runs(p,'"version": "1.2.0"','"version": "1.0.0"')
        if p.text.strip().startswith('{') and '"activityId": "ATV-2026-001"' in p.text:
            obj=json.loads(p.text); obj['enabledCriterionCodes']=['ADEQUACAO_GENERO','TESE_OU_FOCO','ARGUMENTACAO','COERENCIA_GLOBAL','CLAREZA_PRECISAO']; obj['disabledCriterionCodes']=[]; p.text=json.dumps(obj,ensure_ascii=False,indent=2)
        if p.text.strip().startswith('4. Recuperar somente critérios'):
            p.text='4. Calcular effectiveCriterionCodes = (mapped ∩ enabledCriterionCodes) − disabledCriterionCodes; quando enabledCriterionCodes estiver ausente, usar mapped.'
        if p.text.strip().startswith('5. Construir AUTHORIZED_TAXONOMY'):
            p.text='5. Recuperar categorias vinculadas aos critérios efetivos, construir AUTHORIZED_TAXONOMY, calcular checksum e registrar as versões e a seleção docente.'
    d.save(OUT/src.name)

def prompt_hash(doc_path,heading):
    d=Document(doc_path); ps=[p.text for p in d.paragraphs]; i=next(i for i,t in enumerate(ps) if t==heading)
    body='\n'.join(ps[i:]).strip().replace('\r\n','\n')
    return hashlib.sha256(body.encode()).hexdigest()

def update_9():
    src=BASE/'9. Especificação técnica dos módulos de revisão e avaliação do Textopia_.docx'; d=Document(src)
    ps=d.paragraphs
    for p in ps:
        if p.text.strip() == 'Problemas que dependam de mais de uma sentença devem integrar feedbackGlobal.':
            p.text = ('Problemas que dependam de mais de uma sentença devem integrar achadosGlobais, '
                      'com os sentenceIds que sustentam o diagnóstico; feedbackGlobal conserva apenas a síntese pedagógica.')
    # Canonical revision schema becomes 1.1.0 with structured global findings.
    first=next(p for p in ps if p.text.strip().startswith('{') and 'SaidaRevisaoTextopia' in p.text)
    obj=json.loads(first.text); obj['$id']=obj['$id'].replace('/1.0.0','/1.1.0'); obj['properties']['version']={'const':'1.1.0'}
    if 'achadosGlobais' not in obj['required']: obj['required'].insert(-1,'achadosGlobais')
    obj['properties']['achadosGlobais']={'type':'array','items':{'$ref':'https://textopia.local/schemas/achado-global/1.0.0'}}
    first.text=json.dumps(obj,ensure_ascii=False,indent=2)
    # Update valid example to the new output contract.
    valid=next(p for p in d.paragraphs if p.text.strip().startswith('{') and '"executionId": "rev-1042"' in p.text)
    vobj=json.loads(valid.text); vobj['version']='1.1.0'; vobj['taxonomyVersion']='1.0.0'; vobj['achadosGlobais']=[]; valid.text=json.dumps(vobj,ensure_ascii=False,indent=2)
    # Validated package points to revision 1.1.0 and validates global finding IDs directly.
    vp=next(p for p in d.paragraphs if p.text.strip().startswith('{') and 'revisao-validada' in p.text)
    pobj=json.loads(vp.text); pobj['$id']=pobj['$id'].replace('/1.0.0','/1.1.0'); pobj['properties']['version']={'const':'1.1.0'}
    pobj['properties']['revisionOutput']['$ref']='https://textopia.local/schemas/revisao/1.1.0'
    gf=pobj['properties']['validation']['properties']['globalFindings']['items']
    gf['required']=['globalFindingId','status']; gf['properties']={'globalFindingId':{'type':'string','minLength':1},'status':gf['properties']['status'],'editedText':gf['properties']['editedText']}
    vp.text=json.dumps(pobj,ensure_ascii=False,indent=2)
    # Remove manually maintained prompt bodies from 9.7 and 9.15.
    def collapse(start,end,replacement):
        paras=d.paragraphs; a=next(i for i,p in enumerate(paras) if p.text.startswith(start)); b=next(i for i,p in enumerate(paras) if i>a and p.text.startswith(end))
        paras[a+1].text=replacement
        for p in paras[a+2:b]: p._element.getparent().remove(p._element)
    ha=prompt_hash(OUT/'Anexo técnico A - Prompt canônico de revisão do Textopia_.docx','PROMPT CANÔNICO DE REVISÃO PEDAGÓGICA DO TEXTOPIA_')
    hb=prompt_hash(OUT/'Anexo técnico B - Prompt canônico de avaliação rubricada e proposta de nota do Textopia_.docx','PROMPT CANÔNICO DE AVALIAÇÃO RUBRICADA E PROPOSTA DE NOTA DO TEXTOPIA_')
    collapse('9.7 Prompt canônico','9.8 Exemplo válido',f'A fonte normativa é o Anexo técnico A, versão revisao-1.1.0. Na publicação deste livro, a reprodução do prompt deve ser gerada automaticamente a partir do anexo, nunca editada neste capítulo, e aceita somente se o SHA-256 canônico for {ha}.')
    collapse('9.15 Prompt canônico','9.16 Cálculo determinístico',f'A fonte normativa é o Anexo técnico B, versão avaliacao-nota-1.0.0. Na publicação deste livro, a reprodução do prompt deve ser gerada automaticamente a partir do anexo, nunca editada neste capítulo, e aceita somente se o SHA-256 canônico for {hb}. A rubrica e o cálculo obedecem ao Anexo técnico D.')
    # Add normative cross-references after 9.18.
    d.add_heading('9.19 Fontes normativas complementares',level=2)
    d.add_paragraph('O Anexo técnico D define rubricas, níveis, pesos, mapeamento e cálculo da nota. O Anexo técnico E define o modelo canônico do domínio, persistência, achadosGlobais, segmentação e máquinas de estados. Em conflito, prevalece o anexo especializado de maior versão compatível.')
    d.add_heading('9.20 Regra de publicação documental',level=2)
    d.add_paragraph('O processo editorial extrai o texto canônico dos Anexos A e B, normaliza quebras de linha para LF, calcula SHA-256, compara versão e hash declarados nesta seção e somente então gera as reproduções destinadas ao livro. Divergência interrompe a publicação. Os arquivos de origem permanecem imutáveis e são os únicos editáveis.')
    d.save(OUT/src.name)

make_d(); make_e(); update_a(); update_b(); update_c(); update_9()
print('\n'.join(str(p) for p in OUT.glob('*.docx')))
